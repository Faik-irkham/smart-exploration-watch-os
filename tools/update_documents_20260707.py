# -*- coding: utf-8 -*-
"""Apply the 7 July 2026 technical corrections to both manuscripts and
create a new dated correction note without modifying the historical notes.

The manuscript generators in this repository still contain an older,
single-session draft, so this migration intentionally updates the current
DOCX files in place while preserving their figures and document structure.
"""

from pathlib import Path
import re

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]


def read_session5_metrics():
    """Parse tx metrics and the following ACK from the 5 July watch log."""
    log_path = ROOT / "watch_log_2026-07-05_23-37-57.txt"
    pattern = re.compile(
        r"tx_batch,(\d+),(\d+),(\d+),(\d+),([0-9.]+),(\d+)"
    )
    rows = []
    for line in log_path.read_text(encoding="utf-8").splitlines():
        match = pattern.search(line)
        if match:
            records, payload, frames, mtu, duration, throughput = match.groups()
            timestamp = re.search(r"(\d{2}:\d{2}:\d{2}\.\d{3})", line)
            rows.append({
                "time": timestamp.group(1) if timestamp else "—",
                "records": int(records),
                "payload": int(payload),
                "frames": int(frames),
                "mtu": int(mtu),
                "duration": float(duration),
                "throughput_bps": int(throughput),
                "throughput_kib": int(throughput) / 1024,
                "ack": None,
            })
        ack = re.search(r"ACK diterima: (\d+) record", line)
        if ack and rows:
            rows[-1]["ack"] = int(ack.group(1))
    if len(rows) != 7:
        raise ValueError(f"Expected 7 metric rows, found {len(rows)}")
    return rows


def generate_session5_metric_figure():
    """Convert the seven HR-METRIC log rows into a publication-ready chart."""
    out_path = ROOT / "figures" / "fig_session5_batch_metrics.png"
    rows = read_session5_metrics()

    batches = list(range(1, len(rows) + 1))
    durations = [r["duration"] for r in rows]
    throughputs = [r["throughput_kib"] for r in rows]
    mean_duration = sum(durations) / len(durations)
    mean_throughput = sum(throughputs) / len(throughputs)

    plt.rcParams.update({"font.size": 10, "font.family": "DejaVu Sans"})
    fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(9, 7), sharex=True)
    fig.suptitle(
        "Kinerja Pengiriman Batch Sesi 5 (5 Juli 2026)",
        fontsize=15,
        fontweight="bold",
    )

    bars1 = ax1.bar(batches, durations, color="#2474B5", width=0.65)
    ax1.axhline(mean_duration, color="#D1495B", linestyle="--", linewidth=1.5,
                label=f"Rata-rata {mean_duration:.1f} ms")
    ax1.set_ylabel("Durasi transfer (ms)")
    ax1.set_ylim(0, max(durations) * 1.28)
    ax1.grid(axis="y", alpha=0.22)
    ax1.legend(loc="upper left", frameon=False)
    ax1.bar_label(bars1, labels=[f"{v:.1f}" for v in durations], padding=3, fontsize=9)

    bars2 = ax2.bar(batches, throughputs, color="#35A77C", width=0.65)
    ax2.axhline(mean_throughput, color="#D1495B", linestyle="--", linewidth=1.5,
                label=f"Rata-rata {mean_throughput:.1f} KiB/dtk")
    ax2.set_ylabel("Throughput (KiB/detik)")
    ax2.set_xlabel("Nomor batch")
    ax2.set_xticks(batches)
    ax2.set_ylim(0, max(throughputs) * 1.28)
    ax2.grid(axis="y", alpha=0.22)
    ax2.legend(loc="upper left", frameon=False)
    ax2.bar_label(bars2, labels=[f"{v:.1f}" for v in throughputs], padding=3, fontsize=9)

    ack_counts = [r["ack"] for r in rows]
    ack_text = "7/7 batch menerima ACK 180 record" if ack_counts == [180] * 7 else (
        f"ACK tercatat: {sum(a is not None for a in ack_counts)}/7 batch"
    )
    fig.text(
        0.5,
        0.012,
        "Setiap batch: 180 record • 8.461 byte • 19 frame • MTU 512 • " + ack_text,
        ha="center",
        fontsize=9.5,
        color="#333333",
    )
    fig.tight_layout(rect=(0.03, 0.055, 0.98, 0.94))
    fig.savefig(out_path, dpi=220, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return out_path


def replace_paragraph(paragraph, text):
    """Replace paragraph text while retaining its paragraph-level style."""
    paragraph.text = text
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)


def update_manuscript(path, language):
    doc = Document(path)

    if language == "id":
        replacements = {
            "Kontribusi penelitian ini adalah:": (
                "Kontribusi penelitian ini adalah: (1) rancangan dan implementasi sistem dua "
                "aplikasi Wear OS–Android untuk akuisisi dan pengiriman data detak jantung; "
                "(2) protokol pengiriman batch di atas notifikasi BLE dengan framing "
                "ber-opcode dan kendali aliran; (3) mekanisme keandalan berupa "
                "store-and-forward, konfirmasi penerimaan (ACK) tingkat aplikasi, dan "
                "penerima idempoten (anti-duplikat); (4) eksekusi latar belakang melalui "
                "foreground service; serta (5) evaluasi empiris lintas lima sesi terhadap "
                "rasio keberhasilan pengiriman, fidelitas nilai, kualitas data, dan kinerja "
                "transfer—termasuk temuan keterbatasan backfill pada versi implementasi "
                "awal—beserta perangkat bantu pengolahan data yang mendukung reproduksibilitas."
            ),
            "Setiap pembacaan disimpan di smartwatch": (
                "Setiap pembacaan disimpan di smartwatch dengan penanda status terkirim "
                "(synced = 0). Secara berkala (interval dapat dipilih, misalnya 3 atau 5 "
                "menit), seluruh record yang belum terkirim diambil dan dikirim sebagai satu "
                "batch. Setelah transaksi basis data selesai, smartphone menulis ACK yang "
                "memuat jumlah record dalam batch yang berhasil diproses. Smartwatch kemudian "
                "menandai record sebagai terkirim (synced = 1) setelah ACK diterima. Apabila "
                "ACK tidak diterima dalam 30 detik, record tetap belum terkirim dan akan dicoba "
                "kembali pada interval berikutnya. Penerima bersifat idempoten melalui indeks "
                "unik pada atribut time dan INSERT OR IGNORE, sehingga pengiriman ulang tidak "
                "menghasilkan baris basis data ganda. ACK pada versi saat ini mengonfirmasi "
                "selesainya pemrosesan batch, tetapi belum membawa pengenal batch dan nilainya "
                "belum divalidasi terhadap jumlah record yang sedang ditunggu; karena itu, "
                "mekanisme ini meningkatkan keandalan, bukan jaminan formal tanpa kegagalan."
            ),
            "Pengujian dilakukan pada perangkat fisik": (
                "Pengujian dilakukan pada perangkat fisik karena komunikasi BLE tidak dapat "
                "diemulasikan. Spesifikasi perangkat uji ditunjukkan pada Tabel 1. Aplikasi "
                "dipasang dalam mode rilis agar pengukuran kinerja tidak terbias oleh mode "
                "debug. Pengujian mencakup empat sesi pada 23–28 Juni 2026 yang menggunakan "
                "versi implementasi awal, serta satu sesi pada 5 Juli 2026 yang menggunakan "
                "versi revisi dengan penantian ACK sebelum markSynced. Data gabungan lima sesi "
                "berjumlah 25.191 pembacaan pada periode pengukuran. Karena dataset mencakup "
                "dua versi perangkat lunak, hasil agregat bersifat deskriptif historis dan "
                "tidak boleh ditafsirkan sebagai estimasi kinerja satu versi implementasi."
            ),
            "Evaluasi dilakukan pada lima sesi pengujian": (
                "Evaluasi mencakup lima sesi—empat sesi versi awal pada 23–28 Juni 2026 dan "
                "satu sesi versi ACK-revisi pada 5 Juli 2026—dengan hasil pada Tabel 2. Secara "
                "gabungan, smartwatch merekam 25.191 pembacaan dan 22.429 di antaranya tercatat "
                "di smartphone (89,04%) tanpa duplikat. Angka agregat ini mendeskripsikan "
                "seluruh riwayat eksperimen, bukan rasio keberhasilan satu versi perangkat "
                "lunak. Pola temporal menunjukkan bahwa kehilangan pada empat sesi Juni "
                "terkonsentrasi pada periode smartphone belum terhubung; pola tersebut "
                "merupakan asosiasi dan tidak dengan sendirinya membuktikan satu-satunya sebab. "
                "Seluruh record yang berhasil dicocokkan memiliki bpm dan accuracy identik "
                "dengan catatan smartwatch (fidelitas nilai 100%). Sesi 5 versi revisi "
                "menghasilkan 1.974 dari 1.976 record; dua record terakhir masih pending dan "
                "seluruh tujuh batch yang dikirim memperoleh ACK."
            ),
            "Dari tabel terlihat rentang delivery": (
                "Tabel menunjukkan rentang delivery 0,00–99,90%. Sesi 5 mencapai 99,90%; "
                "dua pembacaan terakhir masih berstatus synced = 0 karena direkam setelah "
                "jadwal kirim terakhir. Sesi 3 turun ke 87,01% dan Sesi 2 tidak menerima "
                "record. Perbandingan langsung antarsesi perlu berhati-hati karena empat sesi "
                "Juni menggunakan versi awal, sedangkan Sesi 5 menggunakan revisi ACK."
            ),
            "Gambar tersebut menampilkan sinyal detak jantung": (
                "Gambar tersebut menampilkan sinyal detak jantung sesi utama beserta record "
                "yang tidak ditemukan di smartphone. Kehilangan terkonsentrasi sebelum "
                "smartphone tercatat terhubung sekitar pukul 15.00 dan praktis tidak muncul "
                "setelahnya. Pola ini menunjukkan asosiasi kuat dengan periode tanpa koneksi, "
                "tetapi bukan bukti eksperimental bahwa tidak ada faktor lain yang berperan."
            ),
            "Kinerja transfer diukur melalui": None,
            "Analisis silang antara penanda status terkirim": (
                "Analisis silang empat sesi Juni—yang dijalankan dengan versi implementasi "
                "awal—menemukan bahwa dari 2.760 record yang tidak ditemukan pada smartphone, "
                "2.755 (99,8%) sudah berstatus synced = 1 dan 5 masih pending. Temuan ini "
                "menunjukkan penandaan prematur pada versi awal ketika smartphone tidak "
                "terhubung. Temuan tersebut tidak menggambarkan kode terbaru: sebelum Sesi 5, "
                "implementasi direvisi agar markSynced hanya dijalankan setelah ACK diterima. "
                "Pada Sesi 5 tidak ditemukan false-sent; tujuh batch masing-masing menerima "
                "ACK 180 record dan dua record terakhir tetap synced = 0. Namun, Sesi 5 tidak "
                "mencakup skenario putus–sambung, sehingga ketahanan versi revisi terhadap "
                "disconnect–reconnect masih harus diuji secara khusus."
            ),
            "Jika dihitung pada seluruh rekaman smartwatch": (
                "Pada seluruh rekaman Juni (45.446 baris termasuk periode tidak dipakai), "
                "terdapat 2.879 record berstatus synced = 1 yang tidak ditemukan di smartphone. "
                "Angka utama 2.755 adalah bagian yang berada pada periode pengukuran valid dan "
                "karena itu konsisten dengan tabel hasil Juni. Keduanya merupakan hasil versi "
                "awal, bukan diagnosis terhadap implementasi ACK terbaru."
            ),
            "Hasil menunjukkan bahwa kombinasi": (
                "Hasil menunjukkan bahwa framing dan kendali aliran mampu membawa batch besar "
                "dengan fidelitas nilai 100% pada record yang diterima. Empat sesi versi awal "
                "mengungkap kelemahan penandaan synced saat koneksi tidak tersedia. Revisi ACK "
                "yang digunakan pada Sesi 5 menghilangkan false-sent pada kondisi tersambung, "
                "tetapi belum dievaluasi dalam eksperimen putus–sambung terkontrol. Oleh karena "
                "itu, klaim ketahanan store-and-forward versi terbaru dibatasi pada kondisi "
                "yang telah diuji."
            ),
            "Penelitian ini memiliki beberapa keterbatasan": (
                "Penelitian ini memiliki beberapa keterbatasan. Pertama, pengujian hanya "
                "menggunakan satu pasangan perangkat dan satu lingkungan. Kedua, dataset "
                "gabungan mencakup dua versi perangkat lunak, sehingga evaluasi komparatif "
                "versi terbaru memerlukan eksperimen ulang dengan skenario normal dan "
                "putus–sambung yang terkontrol. Ketiga, ACK terbaru belum memiliki batchId dan "
                "belum memvalidasi jumlah record ACK. Keempat, operasi latar belakang belum "
                "menjamin kelanjutan setelah force-close atau reboot. Kelima, komunikasi belum "
                "menambahkan autentikasi/enkripsi tingkat aplikasi. Keenam, sensor smartwatch "
                "konsumer tidak dinilai sebagai perangkat klinis; kontribusi penelitian ini "
                "berfokus pada pengiriman data, bukan validitas medis nilai BPM."
            ),
            "Penelitian ini merancang dan mengimplementasikan": (
                "Penelitian ini merancang dan mengimplementasikan sistem pengiriman data detak "
                "jantung berbasis BLE dengan framing batch, penyimpanan lokal, ACK tingkat "
                "aplikasi, penerima idempoten, dan foreground service. Lima sesi historis "
                "menghasilkan 25.191 record, fidelitas nilai 100% pada data yang diterima, dan "
                "delivery agregat deskriptif 89,04%. Empat sesi versi awal mengungkap kegagalan "
                "backfill akibat penandaan synced prematur; versi revisi pada Sesi 5 menunggu "
                "ACK dan mencapai 99,90%, dengan dua record terakhir masih pending. Transfer "
                "tujuh batch 180 record pada MTU 512 memerlukan rata-rata 121,5 ms. Pengujian "
                "berikutnya harus mengevaluasi versi revisi secara terpisah dalam skenario "
                "disconnect–reconnect, menambahkan identitas batch pada ACK, dan memperluas "
                "replikasi lintas perangkat serta kondisi radio."
            ),
        }
    else:
        replacements = {
            "The contributions of this work are:": (
                "The contributions of this work are: (1) the design and implementation of a "
                "two-application Wear OS–Android system for heart-rate acquisition and "
                "delivery; (2) an opcode-framed batch protocol over BLE notifications with "
                "flow control; (3) store-and-forward, application-level ACK, and idempotent "
                "receiver mechanisms; (4) background execution using foreground services; "
                "and (5) an empirical evaluation across five sessions, including the backfill "
                "limitation found in the initial implementation, with reproducible analysis "
                "tools."
            ),
            "Each reading is stored on the smartwatch": (
                "Each reading is stored on the smartwatch with synced = 0. At each selected "
                "sending interval, all unsent records are transmitted as one batch. After the "
                "database transaction completes, the smartphone writes an ACK containing the "
                "number of records processed in that batch. The smartwatch marks the records "
                "synced = 1 only after receiving an ACK; otherwise, they remain unsent and are "
                "retried at the next interval. A unique index on time with INSERT OR IGNORE "
                "makes database persistence idempotent. The current ACK confirms completion "
                "of batch processing, but it does not yet carry a batch identifier and its "
                "count is not validated against the pending batch. It therefore improves "
                "reliability rather than providing a formal failure-free guarantee."
            ),
            "Testing was conducted on physical devices": (
                "Testing was conducted on physical devices. Four sessions on 23–28 June 2026 "
                "used the initial implementation, while one session on 5 July 2026 used the "
                "revised implementation that waits for an ACK before markSynced. The combined "
                "dataset contains 25,191 measurement-period readings. Because it spans two "
                "software versions, the five-session aggregate is a descriptive historical "
                "summary and must not be interpreted as the performance estimate of one "
                "implementation version."
            ),
            "The evaluation was carried out over five test sessions": (
                "The evaluation covers four initial-version sessions from 23–28 June 2026 and "
                "one ACK-revised session from 5 July 2026. Combined, 22,429 of 25,191 watch "
                "records were found on the smartphone (89.04%) with no database duplicates. "
                "This aggregate describes the experiment history rather than one software "
                "version. The June losses were temporally concentrated before the smartphone "
                "was connected; this is an association, not proof of a sole cause. All matched "
                "records had identical bpm and accuracy values (100% value fidelity). In "
                "Session 5, the revised version delivered 1,974 of 1,976 records; the final two "
                "remained pending and all seven transmitted batches received ACKs."
            ),
            "The table shows a wide range": (
                "The table shows delivery ranging from 0.00% to 99.90%. Session 5 reached "
                "99.90%; its final two readings remained synced = 0 because they were recorded "
                "after the last scheduled flush. Direct comparisons require caution because "
                "the four June sessions used the initial implementation whereas Session 5 "
                "used the ACK revision."
            ),
            "The figure plots the heart-rate signal": (
                "The figure plots the main-session heart-rate signal and records not found on "
                "the smartphone. Missing records are concentrated before the smartphone was "
                "observed to connect at approximately 15:00 and are practically absent "
                "afterward. This pattern shows a strong association with the disconnected "
                "period but does not, by itself, establish a single cause."
            ),
            "Transfer performance was measured through": None,
            "A cross-analysis between the sent-status marker": (
                "A cross-analysis of the four June sessions, which used the initial software "
                "version, found that 2,755 of 2,760 records absent from the smartphone (99.8%) "
                "were already marked synced = 1; five were legitimately pending. This shows "
                "premature marking in the initial implementation while the smartphone was "
                "disconnected. It does not describe the latest code: before Session 5, the "
                "implementation was revised so that markSynced runs only after an ACK. Session "
                "5 had zero false-sent records, seven ACKed batches of 180 records, and two "
                "final records that remained synced = 0. However, Session 5 did not include a "
                "controlled disconnect–reconnect event, so the revised version's backfill "
                "resilience still requires dedicated testing."
            ),
            "When computed over the entire smartwatch recording": (
                "Across the full June recording (45,446 rows including the not-worn period), "
                "2,879 records marked synced = 1 were absent from the smartphone. The primary "
                "figure of 2,755 is the subset within the valid measurement period and is "
                "therefore consistent with the June results table. Both figures describe the "
                "initial version rather than the current ACK implementation."
            ),
            "The results show that the combination": (
                "The results show that opcode framing and flow control delivered received "
                "records with 100% value fidelity. Four initial-version sessions exposed a "
                "premature synced-marking failure when no connection was available. The ACK "
                "revision used in Session 5 eliminated false-sent records under continuous "
                "connection, but it has not yet been evaluated in a controlled "
                "disconnect–reconnect experiment. Claims about current store-and-forward "
                "resilience are therefore limited to tested conditions."
            ),
            "This study has several limitations": (
                "This study has several limitations. It uses one device pair in one "
                "environment, and the combined dataset spans two software versions. The "
                "latest version therefore requires a separate replicated evaluation under "
                "normal and controlled disconnect–reconnect conditions. Its ACK does not yet "
                "include a batchId or validate the ACK count. Background execution does not "
                "guarantee continuation after force-close or reboot, communication lacks "
                "application-level authentication/encryption, and the consumer smartwatch "
                "sensor was not assessed for clinical validity."
            ),
            "This study designed and implemented": (
                "This study designed and implemented a BLE heart-rate delivery system using "
                "batch framing, local persistence, application-level ACKs, an idempotent "
                "receiver, and foreground services. Across five historical sessions, 22,429 "
                "of 25,191 records were matched (a descriptive aggregate of 89.04%) with 100% "
                "value fidelity. Four initial-version sessions revealed failed backfill caused "
                "by premature synced marking. The revised Session 5 implementation waited for "
                "ACKs and reached 99.90%, with its final two readings still pending. Seven "
                "180-record batches at MTU 512 transferred in 121.5 ms on average. Future work "
                "should evaluate the revised version separately under controlled disconnection, "
                "add batch identity to ACKs, and replicate across devices and radio conditions."
            ),
        }

    for paragraph in doc.paragraphs:
        if paragraph.text.startswith(
            "Pemantauan detak jantung secara berkelanjutan menggunakan perangkat wearable"
        ):
            replace_paragraph(paragraph, (
                "Pemantauan detak jantung berkelanjutan menggunakan wearable memerlukan "
                "pengiriman yang menjaga integritas data meskipun ukuran notifikasi BLE "
                "terbatas dan koneksi dapat terputus. Penelitian ini merancang sistem dua "
                "aplikasi: smartwatch Wear OS merekam satu record per detik ke SQLite dan "
                "mengirim batch sebagai peripheral/GATT server, sedangkan smartphone Android "
                "merangkai frame, menyimpan record secara idempoten, dan mengirim ACK tingkat "
                "aplikasi. Evaluasi historis mencakup empat sesi versi awal pada 23–28 Juni "
                "2026 dan satu sesi versi ACK-revisi pada 5 Juli 2026, dengan total 25.191 "
                "record. Secara gabungan, 22.429 record cocok di smartphone (89,04%) dan semua "
                "record yang cocok memiliki bpm serta accuracy identik tanpa duplikat. Empat "
                "sesi versi awal mengungkap bahwa 2.755 record periode pengukuran telah keliru "
                "ditandai synced meskipun tidak ada di smartphone. Setelah markSynced diubah "
                "agar menunggu ACK, Sesi 5 mencapai 1.974 dari 1.976 record (99,90%); dua record "
                "terakhir masih pending dan seluruh tujuh batch memperoleh ACK. Pada MTU 512, "
                "batch 180 record (8.461 byte, 19 frame) terkirim dalam 121,5 ± 25,5 ms. "
                "Hasil menunjukkan integritas transfer yang tinggi, tetapi ketahanan versi "
                "revisi terhadap putus–sambung masih memerlukan eksperimen terkontrol."
            ))
            continue
        if paragraph.text.startswith(
            "Continuous heart-rate monitoring using wearable devices is a key component"
        ):
            replace_paragraph(paragraph, (
                "Continuous wearable heart-rate monitoring requires data delivery that "
                "preserves integrity despite limited BLE notification size and intermittent "
                "connectivity. This study implements two applications: a Wear OS smartwatch "
                "records one SQLite row per second and sends opcode-framed batches as a "
                "peripheral/GATT server, while an Android smartphone reassembles frames, "
                "persists records idempotently, and returns an application-level ACK. The "
                "historical evaluation comprises four initial-version sessions from 23–28 "
                "June 2026 and one ACK-revised session on 5 July 2026, totaling 25,191 records. "
                "Combined, 22,429 records were matched on the smartphone (89.04%); every "
                "matched record had identical bpm and accuracy values and no database "
                "duplicates. The initial-version sessions revealed 2,755 measurement-period "
                "records prematurely marked synced although absent from the smartphone. After "
                "markSynced was revised to wait for an ACK, Session 5 delivered 1,974 of 1,976 "
                "records (99.90%); the final two remained pending and all seven transmitted "
                "batches were acknowledged. At MTU 512, 180-record batches (8,461 bytes, 19 "
                "frames) transferred in 121.5 ± 25.5 ms. The results demonstrate high transfer "
                "integrity, while the revised version's disconnect–reconnect resilience still "
                "requires controlled evaluation."
            ))
            continue
        for prefix, replacement in replacements.items():
            if paragraph.text.startswith(prefix):
                if replacement is None:
                    paragraph.text = paragraph.text.replace(
                        "70,3 ± 13,2 KB/detik" if language == "id" else "70.3 ± 13.2 KB/s",
                        "70,3 ± 13,2 KiB/detik" if language == "id" else "70.3 ± 13.2 KiB/s",
                    )
                else:
                    replace_paragraph(paragraph, replacement)
                break

        if language == "id":
            paragraph.text = paragraph.text.replace(
                "0,1% pada tingkat sedang", "0,1% pada status tidak dapat dipercaya"
            ).replace(
                "0,1% pada nilai sedang", "0,1% pada status tidak dapat dipercaya"
            ).replace(
                "tingkat sedang hanya 0,1%", "status tidak dapat dipercaya hanya 0,1%"
            )
        else:
            paragraph.text = paragraph.text.replace(
                "0.1% at the medium level", "0.1% at the unreliable status"
            )

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if language == "id" and cell.text.strip() == "0 (sedang)":
                    cell.text = "0 (tidak dapat dipercaya)"
                if language == "en" and cell.text.strip() == "0 (medium)":
                    cell.text = "0 (unreliable)"

    doc.save(path)


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)


def create_note(path):
    # Start from the complete 6 July note so the new document is standalone
    # and retains all four embedded diagrams and five data tables.
    source = ROOT / "docs" / "26-07-06-Catatan-Faik-BLE-SmartWatch-sederhana.docx"
    doc = Document(source)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)

    # Correct a few statements in the inherited text that are known to be
    # technically inaccurate, while retaining the original experimental story.
    replace_paragraph(
        doc.paragraphs[0],
        "UJI PENGIRIMAN DATA DETAK JANTUNG — CATATAN LENGKAP DAN KOREKSI 7 JULI 2026",
    )
    for paragraph in doc.paragraphs:
        corrected = paragraph.text.replace(
            "3 berarti kontak baik, 0 berarti sedang, dan -1 berarti tidak ada kontak",
            "3 berarti akurasi tinggi, 0 berarti tidak dapat dipercaya, dan -1 berarti tidak ada kontak",
        )
        if corrected != paragraph.text:
            replace_paragraph(paragraph, corrected)
        if paragraph.text.startswith(
            "Artinya, pada implementasi saat ini smartwatch menandai data"
        ):
            replace_paragraph(paragraph, (
                "Artinya, pada implementasi awal yang menghasilkan empat sesi Juni, "
                "smartwatch menandai data sebagai terkirim sebelum ACK tingkat aplikasi "
                "benar-benar dikonfirmasi pada kondisi tanpa koneksi. Kode kemudian direvisi "
                "sebelum Sesi 5 agar markSynced hanya dijalankan setelah ACK diterima. "
                "Dengan demikian, temuan false-sent menjelaskan versi awal dan tidak boleh "
                "disebut sebagai perilaku implementasi terbaru."
            ))
        if paragraph.text.startswith("Saran perbaikan: (1) data tidak ditandai"):
            replace_paragraph(paragraph, (
                "Status perbaikan: mekanisme utama tersebut SUDAH diterapkan sebelum Sesi 5. "
                "Watch sekarang memanggil markSynced hanya setelah ACK diterima; bila tidak "
                "ada subscriber, pengiriman gagal, atau ACK tidak tiba dalam 30 detik, record "
                "tetap synced = 0 dan akan dicoba lagi pada interval berikutnya. Database phone "
                "juga sudah memakai indeks unik time dengan INSERT OR IGNORE agar retransmisi "
                "tidak membuat duplikat. Hasil Sesi 5 menunjukkan tujuh batch × 180 record "
                "seluruhnya memperoleh ACK, 1.974 dari 1.976 record cocok di phone, tidak ada "
                "duplikat maupun false-sent, dan dua record terakhir tetap pending (synced = 0). "
                "Yang belum dilakukan bukan penerapan ACK, melainkan pengujian terkontrol "
                "disconnect–reconnect serta penguatan opsional berupa batchId dan validasi "
                "jumlah ACK."
            ))
        if paragraph.text.startswith(
            "Terlihat bahwa MTU hanya memengaruhi banyaknya frame"
        ):
            replace_paragraph(paragraph, (
                "MTU memengaruhi kapasitas chunk dan jumlah frame. Nilai MTU 512 telah "
                "diamati pada perangkat uji, sedangkan jumlah frame pada MTU 23 adalah "
                "perhitungan matematis. Karena tidak tersedia log transfer nyata pada MTU 23, "
                "kelengkapan dan durasinya tidak boleh dinyatakan sebagai hasil empiris."
            ))
        if paragraph.text.startswith(
            "Pada gambar terlihat batch yang sama dipotong dengan dua cara"
        ):
            replace_paragraph(paragraph, (
                "Gambar membandingkan hasil terukur pada MTU 512 dengan proyeksi matematis "
                "pada MTU 23. Pada MTU 512, 24 frame dan durasi sekitar 0,32 detik berasal dari "
                "log. Pada MTU 23, 567 frame merupakan hasil perhitungan, bukan pengujian "
                "langsung; kelengkapan transfer pada kondisi tersebut belum diverifikasi."
            ))

    doc.add_page_break()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("KOREKSI TEKNIS DAN PEMISAHAN VERSI IMPLEMENTASI")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    subtitle = doc.add_paragraph("Tambahan tanggal 7 Juli 2026")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "Tujuan")
    add_para(doc, "Catatan ini memperbaiki interpretasi pada catatan sebelumnya dan kedua manuskrip. Catatan lama tetap dipertahankan sebagai rekam kronologis, tetapi hasil eksperimen harus dikaitkan dengan versi perangkat lunak yang benar. Perbaikan ini tidak mengubah data mentah maupun angka hasil; yang diperbaiki adalah pelabelan versi, terminologi, kekuatan klaim, dan unit pengukuran.")

    add_heading(doc, "1. Pemisahan Versi Implementasi")
    add_para(doc, "Empat sesi tanggal 23–28 Juni 2026 dijalankan menggunakan implementasi awal. Pada versi tersebut, analisis menemukan record dapat berstatus synced = 1 walaupun tidak ditemukan di basis data smartphone. Dari 2.760 record yang tidak ditemukan dalam periode pengukuran valid, 2.755 berstatus synced = 1 dan lima masih pending. Temuan ini sah untuk versi awal.")
    add_para(doc, "Sebelum pengujian 5 Juli 2026, kode direvisi agar smartwatch menunggu ACK tingkat aplikasi sebelum menjalankan markSynced. Pada Sesi 5, tujuh batch masing-masing berisi 180 record menerima ACK, tidak ditemukan false-sent, dan dua record terakhir tetap synced = 0 karena direkam setelah flush terakhir. Karena Sesi 5 berlangsung dengan koneksi penuh, hasil ini belum membuktikan ketahanan versi revisi dalam skenario putus–sambung. Uji disconnect–reconnect terkontrol masih diperlukan.")
    add_para(doc, "Akibat perbedaan versi tersebut, agregat lima sesi—25.191 record watch, 22.429 record cocok di phone, delivery 89,04%—adalah ringkasan historis deskriptif. Angka itu tidak boleh dipakai sebagai estimasi kinerja satu versi aplikasi.")

    add_heading(doc, "2. Interval Pengukuran dan Pengiriman")
    add_para(doc, "Sensor didengarkan secara kontinu oleh kode native. Lapisan aplikasi menyimpan nilai BPM terbaru yang valid sekitar satu kali per detik. Pada aplikasi utama, batch dikirim setiap 3 atau 5 menit. Pada varian basic_sensor_heart_rate_fastflush_sqflite_ble, proses pengukuran dan penyimpanan tetap satu kali per detik; hanya interval flush yang berubah menjadi 15 atau 30 detik.")

    add_heading(doc, "3. Arti Nilai Accuracy")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.rows[0].cells[0].text = "Nilai"
    table.rows[0].cells[1].text = "Arti Android SensorManager"
    for value, meaning in [("-1", "No contact"), ("0", "Unreliable / tidak dapat dipercaya"), ("1", "Low"), ("2", "Medium"), ("3", "High")]:
        cells = table.add_row().cells
        cells[0].text = value
        cells[1].text = meaning
    add_para(doc, "Dengan demikian, 15 record bernilai accuracy = 0 pada tabel manuskrip harus disebut unreliable, bukan medium.")

    add_heading(doc, "4. Batas Makna ACK Saat Ini")
    add_para(doc, "Receiver menulis ACK setelah transaksi penyimpanan selesai, tetapi payload ACK saat ini adalah jumlah record hasil decode, bukan jumlah baris baru yang benar-benar ditambahkan setelah INSERT OR IGNORE. Smartwatch juga belum memeriksa bahwa nilai ACK sama dengan jumlah record batch aktif dan ACK belum membawa batchId. Oleh karena itu, istilah yang tepat adalah ACK meningkatkan keandalan dan mengonfirmasi pemrosesan batch; belum tepat menyebutnya sebagai jaminan formal tanpa kegagalan. Penguatan yang disarankan adalah batchId, expectedCount, validasi ACK, dan pencatatan status batch.")

    add_heading(doc, "5. MTU, Frame, dan Kinerja")
    add_para(doc, "Pada ATT notification, kapasitas atribut adalah ATT_MTU − 3 byte. Karena protokol aplikasi memakai satu byte opcode, chunk DATA maksimum adalah MTU − 4. Pada MTU 512, batch 8.461 byte memerlukan 17 frame DATA ditambah START dan END, sehingga total 19 frame. Tujuh batch Sesi 5 memiliki durasi rata-rata 121,5 ± 25,5 ms.")
    add_para(doc, "Nilai throughput 70,3 diperoleh dengan konversi 1.024 byte per unit, sehingga unit yang tepat adalah KiB/detik. Jika memakai satuan SI kB/detik, rata-ratanya sekitar 72,0 kB/detik. Perhitungan 567 frame pada MTU 23 adalah proyeksi matematis untuk batch 10.717 byte, bukan hasil eksperimen yang telah dibuktikan pada perangkat. Estimasi backlog satu jam sekitar lima detik juga merupakan ekstrapolasi dan harus dilabeli sebagai perkiraan, bukan hasil ukur.")

    add_heading(doc, "5.1 Mengapa MTU 512 Dapat Mengirim Batch 10.717 Byte")
    add_para(doc, "MTU 512 bukan kapasitas seluruh batch, melainkan batas ukuran satu paket ATT/notifikasi BLE. Karena itu, batch JSON 10.717 byte tidak dimasukkan ke satu notifikasi. Watch memotongnya menjadi beberapa frame DATA, mengirim frame tersebut berurutan, lalu phone menggabungkannya kembali menjadi JSON utuh.")
    mtu_table = doc.add_table(rows=1, cols=3)
    mtu_table.style = "Table Grid"
    mtu_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, heading_text in enumerate(("Komponen", "Perhitungan", "Hasil")):
        mtu_table.rows[0].cells[index].text = heading_text
    mtu_rows = [
        ("ATT MTU hasil negosiasi", "—", "512 byte per notifikasi"),
        ("Header ATT", "512 − 3", "509 byte tersisa"),
        ("Opcode aplikasi", "509 − 1", "508 byte JSON per frame DATA"),
        ("Ukuran batch JSON", "—", "10.717 byte"),
        ("Jumlah frame DATA", "ceil(10.717 ÷ 508)", "22 frame DATA"),
        ("Frame kontrol", "1 START + 1 END", "2 frame"),
        ("Total frame", "22 + 2", "24 frame"),
    ]
    for component, calculation, result in mtu_rows:
        cells = mtu_table.add_row().cells
        cells[0].text = component
        cells[1].text = calculation
        cells[2].text = result
    add_para(doc, "Rinciannya: 21 frame DATA pertama dapat membawa paling banyak 21 × 508 = 10.668 byte. Sisa payload adalah 10.717 − 10.668 = 49 byte dan dikirim pada frame DATA ke-22. Urutan transportnya menjadi START → DATA 1 → DATA 2 → … → DATA 22 → END. Dengan demikian, jumlah yang tercatat di log adalah 24 frame: satu START, 22 DATA, dan satu END.")
    add_para(doc, "Di phone, frame START mengosongkan buffer penerimaan. Setiap frame DATA menambahkan byte JSON setelah opcode ke buffer tersebut. Ketika frame END diterima, phone mendekode seluruh 10.717 byte sebagai satu array JSON berisi 228 record, menyimpannya ke SQLite, lalu menulis ACK. Jadi, MTU 512 membatasi setiap potongan, sedangkan framing dan reassembly memungkinkan ukuran batch total melampaui 512 byte.")
    add_para(doc, "Baris log tx_batch,228,10717,24,512,323.1,33168 berarti: 228 record dengan payload 10.717 byte dikirim sebagai 24 frame pada MTU 512; seluruh pengiriman frame membutuhkan 323,1 ms dengan throughput 33.168 byte/detik (sekitar 32,4 KiB/detik). Kepastian batch telah diproses phone ditunjukkan terpisah oleh log ACK diterima: 228 record.")

    metric_figure = generate_session5_metric_figure()
    picture = doc.add_paragraph()
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture.add_run().add_picture(str(metric_figure), width=Inches(6.4))
    caption = doc.add_paragraph(
        "Gambar 5. Visualisasi tujuh baris HR-METRIC pada "
        "watch_log_2026-07-05_23-37-57.txt. Panel atas menunjukkan durasi "
        "transfer dan panel bawah menunjukkan throughput setiap batch."
    )
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in caption.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
    add_para(doc, "Cara membaca Gambar 5: batang yang lebih pendek pada panel durasi berarti transfer lebih cepat, sedangkan batang yang lebih tinggi pada panel throughput berarti laju transfer lebih besar. Batch 4 paling cepat (92,8 ms) dan Batch 5 paling lambat (166,4 ms). Meskipun kecepatannya bervariasi, seluruh batch memiliki isi dan konfigurasi yang sama serta semuanya menerima ACK 180 record.")

    metric_caption = doc.add_paragraph(
        "Tabel 6. Rincian tujuh batch dari "
        "watch_log_2026-07-05_23-37-57.txt."
    )
    metric_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in metric_caption.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
    metric_table = doc.add_table(rows=1, cols=9)
    metric_table.style = "Table Grid"
    metric_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = (
        "Batch", "Waktu", "Record", "Payload (byte)", "Frame", "MTU",
        "Durasi (ms)", "Throughput (KiB/dtk)", "ACK",
    )
    for index, header in enumerate(headers):
        metric_table.rows[0].cells[index].text = header
    for index, row in enumerate(read_session5_metrics(), start=1):
        cells = metric_table.add_row().cells
        values = (
            str(index), row["time"], str(row["records"]),
            f'{row["payload"]:,}'.replace(",", "."), str(row["frames"]),
            str(row["mtu"]), f'{row["duration"]:.1f}'.replace(".", ","),
            f'{row["throughput_kib"]:.1f}'.replace(".", ","),
            f'{row["ack"]} record' if row["ack"] is not None else "Tidak ada",
        )
        for cell_index, value in enumerate(values):
            cells[cell_index].text = value
    add_para(doc, "Tabel 6 memperlihatkan bahwa ukuran setiap batch identik, sedangkan durasi dan throughput berubah. Variasi tersebut wajar karena waktu penjadwalan dan kondisi tautan BLE dapat berubah antarpengiriman. Semua ACK bernilai 180 record, sama dengan jumlah record masing-masing batch.")

    add_heading(doc, "6. Bahasa Klaim Ilmiah")
    add_para(doc, "Pola temporal menunjukkan bahwa kehilangan pada sesi Juni terkonsentrasi sebelum smartphone terhubung. Formulasi yang tepat adalah ‘berasosiasi dengan periode tanpa koneksi’, bukan ‘terbukti tidak disebabkan oleh kanal BLE’. Demikian pula, gunakan ‘meningkatkan keandalan’ alih-alih ‘menjamin tidak ada data hilang’, kecuali seluruh kondisi kegagalan telah diuji secara terkontrol.")

    add_heading(doc, "7. Status Manuskrip Setelah Koreksi")
    add_para(doc, "Kedua manuskrip telah diperbarui agar: menyebut lima sesi secara konsisten; memisahkan empat sesi versi awal dari Sesi 5 versi ACK-revisi; memperbaiki accuracy = 0 menjadi unreliable; mengubah unit throughput menjadi KiB/detik; membatasi klaim kausal; menjelaskan keterbatasan ACK; dan menempatkan pengujian disconnect–reconnect versi terbaru sebagai pekerjaan berikutnya.")

    doc.save(path)


if __name__ == "__main__":
    update_manuscript(ROOT / "Draft_Naskah_HR_BLE_SINTA2.docx", "id")
    update_manuscript(ROOT / "Draft_Manuscript_HR_BLE_EN.docx", "en")
    create_note(ROOT / "docs" / "26-07-07-Catatan-Faik-BLE-SmartWatch-sederhana.docx")
    print("Updated both manuscripts and created the 7 July correction note.")
