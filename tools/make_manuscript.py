# -*- coding: utf-8 -*-
"""Generate draft naskah (DOCX) level jurnal — Bahasa Indonesia + Abstract EN.

Memakai data nyata satu sesi pengujian dan gambar di figures/. Bagian yang
masih perlu dilengkapi penulis ditandai [[...]]. Tidak memalsukan sitasi:
referensi yang dapat diverifikasi diisi, literatur domain ditandai untuk diisi.

Pakai:  python3 tools/make_manuscript.py [OUT.docx]
Butuh:  pip install python-docx
"""
import os
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT = sys.argv[1] if len(sys.argv) > 1 else "Draft_Naskah_HR_BLE_SINTA2.docx"
FIG = "figures"
FONT = "Times New Roman"

doc = Document()

# ---- gaya dasar ----
st = doc.styles["Normal"]
st.font.name = FONT
st.font.size = Pt(11)
st.paragraph_format.space_after = Pt(6)
st.paragraph_format.line_spacing = 1.15


def set_font(run, size=11, bold=False, italic=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def para(text="", align="justify", size=11, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    a = {"justify": WD_ALIGN_PARAGRAPH.JUSTIFY, "center": WD_ALIGN_PARAGRAPH.CENTER,
         "left": WD_ALIGN_PARAGRAPH.LEFT}[align]
    p.alignment = a
    if text:
        r = p.add_run(text)
        set_font(r, size, bold, italic)
    return p


def heading(text, level=1):
    sizes = {1: 12, 2: 11}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, sizes.get(level, 11), bold=True)
    return p


def figure(path, caption, width=5.3):
    if not os.path.exists(path):
        para("[[Gambar tidak ditemukan: %s]]" % path, align="center", italic=True)
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Inches(width))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(c.add_run(caption), 10)
    c.paragraph_format.space_after = Pt(8)


def table(headers, rows, caption=None, widths=None):
    if caption:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(c.add_run(caption), 10)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        set_font(r, 10, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            set_font(cells[i].paragraphs[0].add_run(str(v)), 10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ===================== JUDUL & PENULIS =====================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(title.add_run(
    "Pengiriman Data Detak Jantung yang Andal dari Smartwatch ke "
    "Smartphone Berbasis Bluetooth Low Energy Menggunakan Pola "
    "Store-and-Forward dengan Konfirmasi Penerimaan"), 14, bold=True)

au = doc.add_paragraph(); au.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(au.add_run("Faik Irkham¹, [[Nama Pembimbing]]²"), 11, bold=True)
af = doc.add_paragraph(); af.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(af.add_run("¹²[[Program Studi/Jurusan, Fakultas, Universitas, Kota, Indonesia]]"), 10)
em = doc.add_paragraph(); em.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(em.add_run("e-mail: faikirkham@gmail.com"), 10)

# ===================== ABSTRAK (ID) =====================
heading("Abstrak", 1)
para(
    "Pemantauan detak jantung secara berkelanjutan menggunakan perangkat "
    "wearable merupakan komponen penting dalam Internet of Medical Things "
    "(IoMT). Tantangan utamanya adalah menjamin kelengkapan data saat "
    "dikirim dari smartwatch ke smartphone melalui Bluetooth Low Energy "
    "(BLE), mengingat keterbatasan ukuran paket (MTU), sifat notifikasi BLE "
    "yang tidak terjamin sampai, serta pembatasan proses latar belakang "
    "oleh sistem operasi. Penelitian ini merancang dan mengimplementasikan "
    "sistem dua aplikasi: aplikasi smartwatch (Wear OS) berperan sebagai "
    "peripheral/GATT server yang merekam detak jantung tiap detik ke basis "
    "data lokal, lalu mengirimkannya secara berkala dalam bentuk paket "
    "(batch) ke aplikasi smartphone (Android) yang berperan sebagai central. "
    "Untuk keandalan, paket dipecah menjadi rangkaian frame ber-opcode "
    "(START/DATA/END) dengan kendali aliran, dilengkapi pola "
    "store-and-forward, konfirmasi penerimaan tingkat aplikasi (ACK), "
    "penerima yang idempoten (anti-duplikat), dan foreground service agar "
    "tetap berjalan saat layar mati. Pada satu sesi pengujian "
    "berkelanjutan selama ±36 menit pada perangkat fisik, smartwatch "
    "merekam 2.167 pembacaan dan 2.164 di antaranya berhasil diterima "
    "smartphone (rasio keberhasilan 99,86%) tanpa duplikat; sisanya masih "
    "berstatus belum terkirim sesuai mekanisme store-and-forward. Sebanyak "
    "99,1% pembacaan tercatat pada akurasi sensor tertinggi. Satu paket "
    "berisi 228 pembacaan (10.717 byte, 24 frame) terkirim dalam ±0,32 "
    "detik pada MTU 512. Hasil menunjukkan mekanisme yang diusulkan andal "
    "pada kondisi normal; pengujian variasi kondisi menjadi pekerjaan "
    "lanjutan.", space_after=6)
kp = doc.add_paragraph()
set_font(kp.add_run("Kata kunci: "), 11, bold=True)
set_font(kp.add_run("IoMT; Bluetooth Low Energy; wearable; detak jantung; "
                    "store-and-forward; keandalan pengiriman; Wear OS"), 11, italic=True)

# ===================== ABSTRACT (EN) =====================
heading("Abstract", 1)
para(
    "Continuous heart-rate monitoring using wearable devices is a key "
    "component of the Internet of Medical Things (IoMT). A central challenge "
    "is guaranteeing data completeness when readings are transferred from a "
    "smartwatch to a smartphone over Bluetooth Low Energy (BLE), given the "
    "limited packet size (MTU), the best-effort nature of BLE notifications, "
    "and operating-system restrictions on background execution. This work "
    "designs and implements a two-application system: a smartwatch (Wear OS) "
    "application acting as a BLE peripheral/GATT server that records the "
    "heart rate every second into a local database and periodically sends it "
    "as a batch to a smartphone (Android) application acting as the central. "
    "For reliability, each batch is fragmented into opcode-tagged frames "
    "(START/DATA/END) with flow control, complemented by a store-and-forward "
    "scheme, application-level acknowledgement (ACK), an idempotent "
    "(duplicate-free) receiver, and a foreground service for background "
    "operation. In a single continuous ~36-minute test on physical devices, "
    "the smartwatch recorded 2,167 readings and 2,164 were successfully "
    "received by the smartphone (99.86% delivery ratio) with no duplicates; "
    "the remainder stayed pending as designed. 99.1% of readings were at the "
    "highest sensor-accuracy level. One batch of 228 readings (10,717 bytes, "
    "24 frames) was delivered in ~0.32 s at MTU 512. The results show the "
    "proposed mechanism is reliable under normal conditions; evaluation "
    "across varied conditions is future work.", space_after=6)
kp2 = doc.add_paragraph()
set_font(kp2.add_run("Keywords: "), 11, bold=True)
set_font(kp2.add_run("IoMT; Bluetooth Low Energy; wearable; heart rate; "
                     "store-and-forward; delivery reliability; Wear OS"), 11, italic=True)

# ===================== 1. PENDAHULUAN =====================
heading("1. Pendahuluan", 1)
para("Pemanfaatan perangkat wearable untuk pemantauan kesehatan secara "
     "berkelanjutan berkembang pesat seiring meluasnya konsep Internet of "
     "Medical Things (IoMT), yaitu jaringan perangkat medis/kesehatan yang "
     "mengumpulkan dan mempertukarkan data fisiologis [[lengkapi sitasi: "
     "survei IoMT]]. Smartwatch yang dilengkapi sensor detak jantung "
     "optik dapat berperan sebagai simpul sensor, sedangkan smartphone "
     "berfungsi sebagai gateway yang mengumpulkan, menyimpan, dan "
     "meneruskan data untuk analisis lebih lanjut.")
para("Bluetooth Low Energy (BLE) menjadi pilihan utama kanal komunikasi "
     "pada skenario ini karena konsumsi dayanya rendah [1]. Namun, "
     "pengiriman data kesehatan yang utuh melalui BLE menghadapi beberapa "
     "kendala: ukuran satu notifikasi dibatasi oleh Maximum Transmission "
     "Unit (MTU), notifikasi BLE bersifat best-effort (tanpa konfirmasi di "
     "lapisan atribut), dan sistem operasi—terutama pada perangkat Wear OS "
     "dan ponsel dengan manajemen daya agresif—membatasi proses yang "
     "berjalan di latar belakang [3]. Pada aplikasi kesehatan, kehilangan "
     "sebagian data dapat menurunkan kualitas analisis sehingga kelengkapan "
     "data menjadi kebutuhan penting.")
para("Banyak implementasi terdahulu menekankan penampilan nilai terbaru "
     "secara real-time (mis. mengirim satu snapshot pembacaan ketika "
     "diminta), namun pendekatan ini tidak menjamin seluruh pembacaan "
     "tersimpan di sisi penerima apabila koneksi terputus sesaat atau "
     "aplikasi penerima dihentikan sistem [[lengkapi sitasi: studi terkait "
     "wearable BLE/mHealth]]. Diperlukan mekanisme yang menjamin kelengkapan "
     "data, yaitu kombinasi penyimpanan-sementara (store-and-forward), "
     "konfirmasi penerimaan, pencegahan duplikasi, dan kemampuan berjalan di "
     "latar belakang.")
para("Kontribusi penelitian ini adalah: (1) rancangan dan implementasi "
     "sistem dua aplikasi Wear OS–Android untuk akuisisi dan "
     "pengiriman data detak jantung; (2) protokol pengiriman batch di atas "
     "notifikasi BLE dengan framing ber-opcode dan kendali aliran; (3) "
     "mekanisme keandalan berupa store-and-forward, konfirmasi penerimaan "
     "(ACK) tingkat aplikasi, dan penerima idempoten (anti-duplikat); (4) "
     "eksekusi latar belakang melalui foreground service; serta (5) evaluasi "
     "empiris awal terhadap rasio keberhasilan pengiriman, kualitas data, "
     "dan kinerja transfer, beserta perangkat bantu pengolahan data yang "
     "mendukung reproduksibilitas.")

# ===================== 2. METODE =====================
heading("2. Metode Penelitian", 1)
para("Penelitian ini menggunakan pendekatan rancang-bangun (design and "
     "implementation) yang dilanjutkan dengan pengujian pada perangkat "
     "fisik. Bagian ini menjelaskan arsitektur sistem, protokol komunikasi, "
     "mekanisme keandalan, arsitektur aplikasi, serta prosedur dan metrik "
     "evaluasi.")

heading("2.1 Arsitektur Sistem", 2)
para("Arsitektur sistem terdiri atas dua perangkat seperti pada Gambar 1. "
     "Smartwatch berperan sebagai peripheral/GATT server: sensor detak "
     "jantung dibaca oleh kode native, diteruskan ke lapisan aplikasi, "
     "disimpan ke basis data lokal SQLite, lalu diiklankan dan dikirim "
     "melalui BLE. Smartphone berperan sebagai central/GATT client: "
     "memindai, terhubung, berlangganan notifikasi, merangkai paket, "
     "menyimpan ke SQLite, dan menyediakan ekspor data. Aliran data "
     "mengikuti pola “U” bernomor 1–10, dari akuisisi di "
     "smartwatch hingga ekspor di smartphone.")
figure(os.path.join(FIG, "fig_architecture.png"),
       "Gambar 1. Arsitektur sistem dua perangkat (alur data 1–10).", width=4.2)

heading("2.2 Protokol Komunikasi dan Framing", 2)
para("Komunikasi memakai satu layanan GATT khusus dengan karakteristik "
     "“record” bertipe notify. Setelah terhubung, smartphone "
     "meminta MTU 512 byte dan mengaktifkan notifikasi. Karena satu paket "
     "(batch) berisi banyak pembacaan dan ukurannya melampaui kapasitas satu "
     "notifikasi, paket dipecah menjadi rangkaian frame yang masing-masing "
     "diawali opcode: START (0x01), DATA (0x02), dan END (0x03), seperti "
     "pada Gambar 2. Ukuran potongan data tiap frame adalah MTU dikurangi 4 "
     "byte (3 byte header ATT dan 1 byte opcode). Pengiriman antar-frame "
     "dikendalikan dengan flow control: frame berikutnya dikirim setelah "
     "callback pengiriman frame sebelumnya selesai. Penerima merangkai "
     "kembali frame menjadi payload JSON utuh. Urutan interaksi lengkap "
     "ditunjukkan pada Gambar 3.")
figure(os.path.join(FIG, "fig_framing.png"),
       "Gambar 2. Pemecahan satu batch JSON menjadi frame BLE (START/DATA/END).",
       width=6.2)
figure(os.path.join(FIG, "fig_sequence.png"),
       "Gambar 3. Diagram urutan komunikasi BLE: setup, transfer batch, dan ACK.",
       width=4.3)

heading("2.3 Mekanisme Store-and-Forward dan Konfirmasi (ACK)", 2)
para("Setiap pembacaan disimpan di smartwatch dengan penanda status "
     "terkirim (synced = 0). Secara berkala (interval dapat dipilih, "
     "misalnya 3 atau 5 menit), seluruh record yang belum terkirim diambil "
     "dan dikirim sebagai satu batch. Untuk menjamin data benar-benar "
     "diterima, smartphone menulis konfirmasi (ACK) berisi jumlah record "
     "yang berhasil disimpan ke sebuah karakteristik khusus; smartwatch baru "
     "menandai record sebagai terkirim (synced = 1) setelah ACK diterima. "
     "Apabila ACK tidak diterima dalam batas waktu (30 detik), record "
     "dibiarkan belum terkirim dan akan dikirim ulang pada interval "
     "berikutnya. Untuk mencegah duplikasi akibat kirim ulang, penerima "
     "dibuat idempoten melalui indeks unik pada atribut waktu (time) "
     "sehingga pengiriman berulang tidak menghasilkan baris ganda. Alur "
     "lengkap ditunjukkan pada Gambar 4.")
figure(os.path.join(FIG, "fig_storeforward.png"),
       "Gambar 4. Alur store-and-forward dengan konfirmasi penerimaan (ACK).",
       width=3.9)

heading("2.4 Arsitektur Aplikasi Smartwatch", 2)
para("Aplikasi smartwatch dibangun dengan Flutter, dengan pemisahan jelas "
     "antara lapisan Dart dan lapisan native (Kotlin) seperti pada Gambar 5. "
     "Lapisan Dart menggunakan pola BLoC: logika pemantauan (pencuplikan "
     "tiap detik, pengiriman per interval, store-and-forward, dan menunggu "
     "ACK) dipisahkan dari antarmuka. Lapisan native menampung pembacaan "
     "sensor, GATT server beserta advertiser, antrean frame dan kendali "
     "aliran, karakteristik ACK, serta foreground service yang menjaga "
     "proses tetap berjalan saat aplikasi di latar belakang atau layar mati. "
     "Kedua lapisan berkomunikasi melalui platform channel.")
figure(os.path.join(FIG, "fig_watch_arch.png"),
       "Gambar 5. Arsitektur internal aplikasi smartwatch (batas Flutter–Native).",
       width=5.6)

heading("2.5 Perangkat dan Skenario Pengujian", 2)
para("Pengujian dilakukan pada perangkat fisik karena komunikasi BLE tidak "
     "dapat diemulasikan. Spesifikasi perangkat uji ditunjukkan pada "
     "Tabel 1. Aplikasi dipasang dalam mode rilis (release) agar pengukuran "
     "kinerja tidak terbias oleh mode debug. Pada satu run, smartphone "
     "menghubungi smartwatch, lalu pemantauan dijalankan dan dibiarkan "
     "berlangsung selama beberapa interval; data kemudian diekspor ke "
     "format CSV dan basis data untuk dianalisis.")
table(["Komponen", "Spesifikasi"],
      [["Smartwatch (peripheral)", "Samsung Galaxy Watch (SM-R860), Wear OS"],
       ["Smartphone (central)", "Xiaomi Redmi Note 10 Pro (M2101K6G), Android 13"],
       ["Kerangka kerja", "Flutter 3.41.4; Dart SDK ^3.11.1"],
       ["Pustaka utama", "flutter_blue_plus, sqflite, flutter_bloc, permission_handler"],
       ["MTU diminta", "512 byte"],
       ["Interval pengiriman", "3 / 5 menit (dapat dipilih)"]],
      caption="Tabel 1. Spesifikasi perangkat dan lingkungan uji.")

heading("2.6 Metrik Evaluasi", 2)
para("Metrik yang diukur meliputi: (1) rasio keberhasilan pengiriman "
     "(delivery ratio), yaitu jumlah record yang cocok di smartphone dibagi "
     "jumlah record yang direkam smartwatch, dengan pencocokan berdasarkan "
     "atribut waktu (time) sehingga tidak bergantung pada sinkronisasi jam "
     "antar-perangkat; (2) jumlah duplikat di penerima; (3) distribusi "
     "akurasi sensor mengikuti konstanta status akurasi pada Android "
     "SensorManager (rentang -1 hingga 3) [2]; serta (4) latensi transfer "
     "dan throughput per batch yang dicatat otomatis oleh instrumentasi "
     "aplikasi. Throughput dihitung sebagai ukuran payload dibagi durasi "
     "transfer.")

# ===================== 3. HASIL =====================
heading("3. Hasil dan Pembahasan", 1)

heading("3.1 Verifikasi Fungsional End-to-End", 2)
para("Pengujian fungsional menunjukkan rangkaian koneksi berjalan sesuai "
     "rancangan: smartwatch mengiklankan layanan, smartphone memindai dan "
     "terhubung, negosiasi MTU 512 berhasil, notifikasi diaktifkan, dan satu "
     "batch dikirim utuh lalu dikonfirmasi dengan ACK. Hal ini membuktikan "
     "protokol framing dan mekanisme konfirmasi bekerja end-to-end pada "
     "perangkat nyata.")

heading("3.2 Keberhasilan Pengiriman (Delivery Ratio)", 2)
para("Pada satu sesi berkelanjutan selama ±36 menit, diperoleh hasil "
     "pada Tabel 2. Smartwatch merekam 2.167 pembacaan dan 2.164 di "
     "antaranya berhasil diterima smartphone, menghasilkan rasio "
     "keberhasilan 99,86% tanpa duplikat. Tiga pembacaan yang belum sampai "
     "bukan hilang, melainkan masih berstatus belum terkirim pada saat data "
     "diambil dan akan dikirim pada interval berikutnya—menegaskan "
     "mekanisme store-and-forward berfungsi sebagaimana mestinya. Seluruh "
     "data yang diterima identik dengan catatan smartwatch berdasarkan "
     "pencocokan waktu, sehingga integritas data terjaga.")
table(["Indikator", "Nilai"],
      [["Durasi sesi", "± 36 menit"],
       ["Direkam smartwatch", "2.167 pembacaan"],
       ["Ditandai terkirim (synced)", "2.164 pembacaan"],
       ["Diterima smartphone", "2.164 pembacaan"],
       ["Duplikat di smartphone", "0"],
       ["Rasio keberhasilan (delivery ratio)", "99,86%"]],
      caption="Tabel 2. Hasil pengiriman pada satu sesi pengujian.")

heading("3.3 Kualitas Data Sensor", 2)
para("Distribusi akurasi sensor ditunjukkan pada Tabel 3. Sebanyak 99,1% "
     "pembacaan berada pada akurasi tertinggi (nilai 3), dan hanya 0,9% "
     "(20 pembacaan) pada kondisi tanpa kontak (nilai -1), yang lazim "
     "terjadi saat sensor sesaat kehilangan kontak dengan kulit [2]. Tidak "
     "ditemukan pembacaan pada akurasi rendah (0–2). Statistik nilai "
     "detak jantung pada sesi tersebut adalah minimum 71 bpm, maksimum 96 "
     "bpm, rata-rata 80,4 bpm, dan simpangan baku 5,2 bpm. Distribusi "
     "akurasi yang sangat dominan pada nilai tertinggi mengindikasikan "
     "kualitas data yang baik.")
table(["Akurasi sensor", "Jumlah", "Persentase"],
      [["3 (tinggi)", "2.147", "99,1%"],
       ["-1 (tanpa kontak)", "20", "0,9%"],
       ["Total", "2.167", "100%"]],
      caption="Tabel 3. Distribusi nilai akurasi sensor.")

heading("3.4 Kinerja Transfer", 2)
para("Sebagai ilustrasi kinerja per batch, satu batch berisi 228 pembacaan "
     "(payload 10.717 byte) dipecah menjadi 24 frame dan terkirim dalam "
     "±0,32 detik (323,1 ms) pada MTU 512, setara throughput "
     "±33 KB/detik. Di sisi penerima, perakitan ulang frame memerlukan "
     "±250 ms dan penyimpanan batch ke basis data ±56 ms. Nilai "
     "ini menunjukkan transfer batch berlangsung cepat relatif terhadap "
     "interval pengiriman (menit). Pengukuran agregat lintas variasi "
     "kondisi (ukuran batch berbeda akibat interval 3 vs 5 menit, jarak "
     "antar-perangkat, dan skenario gangguan) merupakan bagian dari "
     "pekerjaan lanjutan. [[Lengkapi dengan tabel/grafik hasil replikasi.]]")

heading("3.5 Pembahasan dan Keterbatasan", 2)
para("Hasil menunjukkan kombinasi store-and-forward, konfirmasi ACK, dan "
     "penerima idempoten mampu menjaga kelengkapan dan integritas data pada "
     "kondisi normal, sementara framing ber-opcode dengan kendali aliran "
     "mengatasi keterbatasan ukuran notifikasi BLE. Pencocokan berbasis "
     "waktu membuat perhitungan keberhasilan tidak bergantung pada "
     "sinkronisasi jam.")
para("Penelitian ini memiliki beberapa keterbatasan. Pertama, evaluasi "
     "masih berbasis satu sesi sehingga diperlukan replikasi dan variasi "
     "kondisi untuk memperoleh rata-rata dan simpangan baku yang "
     "representatif. Kedua, eksekusi latar belakang mencakup kondisi "
     "aplikasi di latar belakang dan layar mati, tetapi belum menjamin "
     "operasi saat aplikasi ditutup paksa atau setelah perangkat reboot, "
     "yang juga dibatasi kebijakan hemat daya OS. Ketiga, komunikasi BLE "
     "belum menerapkan enkripsi/otentikasi, sehingga aspek keamanan data "
     "kesehatan menjadi agenda pengembangan. Keempat, sensor pada smartwatch "
     "konsumer belum tervalidasi secara klinis, sehingga fokus kontribusi "
     "adalah keandalan komunikasi, bukan akurasi medis nilai detak jantung.")

# ===================== 4. KESIMPULAN =====================
heading("4. Kesimpulan", 1)
para("Penelitian ini merancang dan mengimplementasikan sistem pengiriman "
     "data detak jantung dari smartwatch ke smartphone berbasis BLE dengan "
     "framing batch, store-and-forward, konfirmasi penerimaan (ACK), "
     "penerima idempoten, dan eksekusi latar belakang. Pengujian pada "
     "perangkat fisik menunjukkan rasio keberhasilan pengiriman 99,86% tanpa "
     "duplikat pada satu sesi ±36 menit, dengan 99,1% data berakurasi "
     "tertinggi, serta transfer batch yang cepat (±0,32 detik untuk 228 "
     "pembacaan pada MTU 512). Pekerjaan lanjutan meliputi pengukuran formal "
     "lintas variasi kondisi dengan replikasi, penguatan eksekusi latar "
     "belakang (termasuk auto-start setelah reboot), serta penambahan "
     "enkripsi untuk keamanan data.")

# ===================== UCAPAN TERIMA KASIH =====================
heading("Ucapan Terima Kasih", 1)
para("[[Opsional: ucapan terima kasih kepada pembimbing/institusi/pendanaan.]]")

# ===================== DAFTAR PUSTAKA =====================
heading("Daftar Pustaka", 1)
refs = [
    "Bluetooth SIG, “Bluetooth Core Specification,” Bluetooth "
    "Special Interest Group. [Daring]. Tersedia: https://www.bluetooth.com/specifications/specs/",
    "Android Developers, “SensorManager,” Android API Reference. "
    "[Daring]. Tersedia: https://developer.android.com/reference/android/hardware/SensorManager",
    "Android Developers, “Foreground services,” Android Developers "
    "Guide. [Daring]. Tersedia: https://developer.android.com/develop/background-work/services/foreground-services",
    "C. Boderick dkk., “flutter_blue_plus,” pub.dev. [Daring]. "
    "Tersedia: https://pub.dev/packages/flutter_blue_plus",
    "“sqflite,” pub.dev. [Daring]. Tersedia: https://pub.dev/packages/sqflite",
    "“flutter_bloc,” pub.dev. [Daring]. Tersedia: https://pub.dev/packages/flutter_bloc",
    "[[Lengkapi sitasi: survei/artikel Internet of Medical Things (IoMT).]]",
    "[[Lengkapi sitasi: studi pemantauan detak jantung berbasis wearable/BLE.]]",
    "[[Lengkapi sitasi: pola store-and-forward / keandalan pengiriman data pada IoT.]]",
    "[[Lengkapi sitasi: Wireless Body Area Network (WBAN) / mHealth.]]",
]
for i, r in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    p.paragraph_format.space_after = Pt(3)
    set_font(p.add_run("[%d] " % i), 10)
    set_font(p.add_run(r), 10)

# catatan untuk penulis
note = doc.add_paragraph()
note.paragraph_format.space_before = Pt(12)
set_font(note.add_run(
    "Catatan: dokumen ini adalah draf. Sesuaikan dengan template resmi "
    "jurnal SINTA 2 tujuan (format kolom, gaya sitasi, batas halaman). "
    "Bagian bertanda [[...]] perlu dilengkapi penulis."), 9, italic=True,
    color=(0x80, 0x80, 0x80))

doc.save(OUT)
print("[OK] tersimpan:", OUT)
