# -*- coding: utf-8 -*-
"""Generate the draft manuscript (DOCX) — full English version.

Uses real single-session data and the figures in figures/. Author-supplied
parts are marked [[...]]. Citations are not fabricated: verifiable references
are filled in; domain literature is marked for the author to complete.

Usage:  python3 tools/make_manuscript_en.py [OUT.docx]
Needs:  pip install python-docx
"""
import os
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT = sys.argv[1] if len(sys.argv) > 1 else "Draft_Manuscript_HR_BLE_EN.docx"
FIG = "figures"
FONT = "Times New Roman"

doc = Document()
st = doc.styles["Normal"]
st.font.name = FONT
st.font.size = Pt(11)
st.paragraph_format.space_after = Pt(6)
st.paragraph_format.line_spacing = 1.15


def set_font(run, size=11, bold=False, italic=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def para(text="", align="justify", size=11, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.alignment = {"justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
                   "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "left": WD_ALIGN_PARAGRAPH.LEFT}[align]
    if text:
        set_font(p.add_run(text), size, bold, italic)
    return p


def heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    set_font(p.add_run(text), {1: 12, 2: 11}.get(level, 11), bold=True)
    return p


def figure(path, caption, width=5.3):
    if not os.path.exists(path):
        para("[[Figure not found: %s]]" % path, align="center", italic=True)
    else:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Inches(width))
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(c.add_run(caption), 10)
    c.paragraph_format.space_after = Pt(8)


def table(headers, rows, caption=None):
    if caption:
        c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(c.add_run(caption), 10)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = ""
        set_font(t.rows[0].cells[i].paragraphs[0].add_run(h), 10, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            set_font(cells[i].paragraphs[0].add_run(str(v)), 10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ===================== TITLE & AUTHORS =====================
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(title.add_run(
    "Reliable Heart-Rate Data Delivery from a Smartwatch to a Smartphone "
    "over Bluetooth Low Energy Using a Store-and-Forward Scheme with "
    "Acknowledgement"), 14, bold=True)

au = doc.add_paragraph(); au.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(au.add_run("Faik Irkham¹, [[Advisor Name]]²"), 11, bold=True)
af = doc.add_paragraph(); af.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(af.add_run("¹²[[Department/Study Program, Faculty, University, City, Indonesia]]"), 10)
em = doc.add_paragraph(); em.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(em.add_run("e-mail: faikirkham@gmail.com"), 10)

# ===================== ABSTRACT =====================
heading("Abstract", 1)
para(
    "Continuous heart-rate monitoring using wearable devices is a key "
    "component of the Internet of Medical Things (IoMT). A central challenge "
    "is guaranteeing data completeness when readings are transferred from a "
    "smartwatch to a smartphone over Bluetooth Low Energy (BLE), given the "
    "limited packet size (MTU), the best-effort nature of BLE notifications, "
    "and operating-system restrictions on background execution. This work "
    "designs and implements a two-application system: a smartwatch (Wear OS) "
    "application acting as a BLE peripheral/GATT server that records the "
    "heart rate every second into a local database and periodically sends it "
    "as a batch to a smartphone (Android) application acting as the central. "
    "For reliability, each batch is fragmented into opcode-tagged frames "
    "(START/DATA/END) with flow control, complemented by a store-and-forward "
    "scheme, application-level acknowledgement (ACK), an idempotent "
    "(duplicate-free) receiver, and a foreground service for background "
    "operation. In a single continuous ~36-minute test on physical devices, "
    "the smartwatch recorded 2,167 readings and 2,164 were successfully "
    "received by the smartphone (99.86% delivery ratio) with no duplicates; "
    "the remainder stayed pending as designed by the store-and-forward "
    "scheme. 99.1% of readings were at the highest sensor-accuracy level. "
    "One batch of 228 readings (10,717 bytes, 24 frames) was delivered in "
    "~0.32 s at MTU 512. The results show the proposed mechanism is reliable "
    "under normal conditions; evaluation across varied conditions is future "
    "work.", space_after=6)
kp = doc.add_paragraph()
set_font(kp.add_run("Keywords: "), 11, bold=True)
set_font(kp.add_run("IoMT; Bluetooth Low Energy; wearable; heart rate; "
                    "store-and-forward; delivery reliability; Wear OS"), 11, italic=True)

# ===================== 1. INTRODUCTION =====================
heading("1. Introduction", 1)
para("The use of wearable devices for continuous health monitoring has grown "
     "rapidly along with the Internet of Medical Things (IoMT), a network of "
     "medical/health devices that collect and exchange physiological data "
     "[[add citation: IoMT survey]]. A smartwatch equipped with an optical "
     "heart-rate sensor can act as a sensor node, while a smartphone serves "
     "as a gateway that collects, stores, and forwards the data for further "
     "analysis.")
para("Bluetooth Low Energy (BLE) is a primary choice of communication "
     "channel in this scenario because of its low power consumption [1]. "
     "However, delivering complete health data over BLE faces several "
     "challenges: the size of a single notification is bounded by the "
     "Maximum Transmission Unit (MTU); BLE notifications are best-effort "
     "(no attribute-layer acknowledgement); and operating systems—"
     "especially Wear OS and phones with aggressive power management—"
     "restrict background processes [3]. In health applications, losing part "
     "of the data can degrade analysis quality, making data completeness an "
     "essential requirement.")
para("Many prior implementations emphasize displaying the latest value in "
     "real time (e.g., sending a single reading snapshot on request); this "
     "approach does not guarantee that all readings are stored at the "
     "receiver if the connection is briefly lost or the receiver application "
     "is stopped by the system [[add citation: related wearable BLE/mHealth "
     "study]]. A mechanism that guarantees data completeness is therefore "
     "needed, combining temporary storage (store-and-forward), delivery "
     "acknowledgement, duplicate prevention, and background operation.")
para("The contributions of this work are: (1) the design and implementation "
     "of a two-application Wear OS–Android system for acquiring and "
     "delivering heart-rate data; (2) a batch delivery protocol over BLE "
     "notifications using opcode-based framing with flow control; (3) "
     "reliability mechanisms consisting of store-and-forward, "
     "application-level acknowledgement (ACK), and an idempotent "
     "(duplicate-free) receiver; (4) background execution via a foreground "
     "service; and (5) an initial empirical evaluation of delivery ratio, "
     "data quality, and transfer performance, together with supporting data "
     "tools for reproducibility.")

# ===================== 2. METHOD =====================
heading("2. Materials and Methods", 1)
para("This study follows a design-and-implementation approach, followed by "
     "testing on physical devices. This section describes the system "
     "architecture, the communication protocol, the reliability mechanisms, "
     "the application architecture, and the evaluation procedure and "
     "metrics.")

heading("2.1 System Architecture", 2)
para("The system comprises two devices, as shown in Figure 1. The "
     "smartwatch acts as a peripheral/GATT server: the heart-rate sensor is "
     "read by native code, forwarded to the application layer, stored in a "
     "local SQLite database, and then advertised and sent over BLE. The "
     "smartphone acts as a central/GATT client: it scans, connects, "
     "subscribes to notifications, reassembles the batch, stores it in "
     "SQLite, and provides data export. The data follows a numbered “U”-"
     "shaped flow (1–10), from acquisition on the smartwatch to export on "
     "the smartphone.")
figure(os.path.join(FIG, "fig_architecture.png"),
       "Figure 1. Two-device system architecture (data flow 1–10).", width=4.2)

heading("2.2 Communication Protocol and Framing", 2)
para("Communication uses a dedicated GATT service with a “record” "
     "characteristic of the notify type. After connecting, the smartphone "
     "requests a 512-byte MTU and enables notifications. Because one batch "
     "contains many readings whose size exceeds the capacity of a single "
     "notification, the batch is fragmented into a sequence of frames, each "
     "prefixed with an opcode: START (0x01), DATA (0x02), and END (0x03), as "
     "shown in Figure 2. The data chunk per frame is MTU minus 4 bytes "
     "(3 bytes of ATT header and 1 byte of opcode). Inter-frame transmission "
     "is flow-controlled: the next frame is sent only after the "
     "send-completion callback of the previous frame. The receiver "
     "reassembles the frames into a complete JSON payload. The full "
     "interaction sequence is shown in Figure 3.")
figure(os.path.join(FIG, "fig_framing.png"),
       "Figure 2. Fragmentation of one JSON batch into BLE frames "
       "(START/DATA/END).", width=6.2)
figure(os.path.join(FIG, "fig_sequence.png"),
       "Figure 3. BLE communication sequence: setup, batch transfer, and ACK.",
       width=4.3)

heading("2.3 Store-and-Forward and Acknowledgement (ACK)", 2)
para("Each reading is stored on the smartwatch with a sent-status flag "
     "(synced = 0). Periodically (a selectable interval, e.g., 3 or 5 "
     "minutes), all not-yet-sent records are retrieved and transmitted as a "
     "single batch. To guarantee that data is actually received, the "
     "smartphone writes an acknowledgement (ACK) containing the number of "
     "successfully stored records to a dedicated characteristic; the "
     "smartwatch marks records as sent (synced = 1) only after the ACK is "
     "received. If no ACK arrives within a timeout (30 s), the records "
     "remain not-yet-sent and are retransmitted in the next interval. To "
     "prevent duplicates caused by retransmission, the receiver is made "
     "idempotent through a unique index on the time attribute, so repeated "
     "delivery does not create duplicate rows. The complete flow is shown in "
     "Figure 4.")
figure(os.path.join(FIG, "fig_storeforward.png"),
       "Figure 4. Store-and-forward flow with delivery acknowledgement (ACK).",
       width=3.9)

heading("2.4 Smartwatch Application Architecture", 2)
para("The smartwatch application is built with Flutter, with a clear "
     "separation between the Dart layer and the native (Kotlin) layer, as "
     "shown in Figure 5. The Dart layer uses the BLoC pattern: the "
     "monitoring logic (per-second sampling, per-interval sending, "
     "store-and-forward, and waiting for ACK) is separated from the user "
     "interface. The native layer hosts the sensor listener, the GATT "
     "server and advertiser, the frame queue and flow control, the ACK "
     "characteristic, and a foreground service that keeps the process alive "
     "when the application is in the background or the screen is off. The two "
     "layers communicate via platform channels.")
figure(os.path.join(FIG, "fig_watch_arch.png"),
       "Figure 5. Internal architecture of the smartwatch application "
       "(Flutter–Native boundary).", width=5.6)

heading("2.5 Devices and Test Scenario", 2)
para("Testing was conducted on physical devices because BLE communication "
     "cannot be emulated. The device specifications are listed in Table 1. "
     "The applications were installed in release mode so that performance "
     "measurements are not biased by debug mode. In one run, the smartphone "
     "connects to the smartwatch, monitoring is started and left running for "
     "several intervals, and the data is then exported to CSV and database "
     "formats for analysis.")
table(["Component", "Specification"],
      [["Smartwatch (peripheral)", "Samsung Galaxy Watch (SM-R860), Wear OS"],
       ["Smartphone (central)", "Xiaomi Redmi Note 10 Pro (M2101K6G), Android 13"],
       ["Framework", "Flutter 3.41.4; Dart SDK ^3.11.1"],
       ["Main libraries", "flutter_blue_plus, sqflite, flutter_bloc, permission_handler"],
       ["Requested MTU", "512 bytes"],
       ["Sending interval", "3 / 5 minutes (selectable)"]],
      caption="Table 1. Device and test-environment specifications.")

heading("2.6 Evaluation Metrics", 2)
para("The measured metrics are: (1) delivery ratio, i.e., the number of "
     "records matched at the smartphone divided by the number of records "
     "recorded by the smartwatch, with matching based on the time attribute "
     "so it does not depend on clock synchronization between devices; (2) "
     "the number of duplicates at the receiver; (3) the sensor-accuracy "
     "distribution following the accuracy-status constants of the Android "
     "SensorManager (range -1 to 3) [2]; and (4) per-batch transfer latency "
     "and throughput recorded automatically by the application "
     "instrumentation. Throughput is computed as the payload size divided by "
     "the transfer duration.")

# ===================== 3. RESULTS =====================
heading("3. Results and Discussion", 1)

heading("3.1 End-to-End Functional Verification", 2)
para("Functional testing showed that the connection sequence works as "
     "designed: the smartwatch advertises the service, the smartphone scans "
     "and connects, the 512-byte MTU negotiation succeeds, notifications are "
     "enabled, and one batch is delivered intact and then confirmed with an "
     "ACK. This demonstrates that the framing protocol and the "
     "acknowledgement mechanism work end-to-end on real devices.")

heading("3.2 Delivery Ratio", 2)
para("In a single continuous ~36-minute session, the results in Table 2 "
     "were obtained. The smartwatch recorded 2,167 readings and 2,164 of "
     "them were successfully received by the smartphone, yielding a 99.86% "
     "delivery ratio with no duplicates. The three readings that did not "
     "arrive were not lost but were still pending at the time the data was "
     "captured and would be sent in the next interval—confirming that the "
     "store-and-forward mechanism works as intended. All received data was "
     "identical to the smartwatch records based on time matching, so data "
     "integrity was preserved.")
table(["Indicator", "Value"],
      [["Session duration", "~36 minutes"],
       ["Recorded on smartwatch", "2,167 readings"],
       ["Marked as sent (synced)", "2,164 readings"],
       ["Received on smartphone", "2,164 readings"],
       ["Duplicates on smartphone", "0"],
       ["Delivery ratio", "99.86%"]],
      caption="Table 2. Delivery results for one test session.")

heading("3.3 Sensor Data Quality", 2)
para("The sensor-accuracy distribution is shown in Table 3. 99.1% of "
     "readings were at the highest accuracy (value 3), and only 0.9% (20 "
     "readings) were at the no-contact condition (value -1), which commonly "
     "occurs when the sensor momentarily loses skin contact [2]. No readings "
     "were at low accuracy (0–2). The heart-rate statistics for the session "
     "were: minimum 71 bpm, maximum 96 bpm, mean 80.4 bpm, and standard "
     "deviation 5.2 bpm. The distribution being strongly dominated by the "
     "highest accuracy indicates good data quality.")
table(["Sensor accuracy", "Count", "Percentage"],
      [["3 (high)", "2,147", "99.1%"],
       ["-1 (no contact)", "20", "0.9%"],
       ["Total", "2,167", "100%"]],
      caption="Table 3. Distribution of sensor-accuracy values.")

heading("3.4 Transfer Performance", 2)
para("As an illustration of per-batch performance, one batch of 228 "
     "readings (10,717-byte payload) was fragmented into 24 frames and "
     "delivered in ~0.32 s (323.1 ms) at MTU 512, equivalent to a "
     "throughput of ~33 KB/s. On the receiver side, frame reassembly took "
     "~250 ms and storing the batch to the database took ~56 ms. These "
     "values show that batch transfer is fast relative to the sending "
     "interval (minutes). Aggregate measurements across varied conditions "
     "(different batch sizes due to 3- vs 5-minute intervals, device "
     "distance, and disruption scenarios) are part of future work. "
     "[[Complete with replication results table/charts.]]")

heading("3.5 Discussion and Limitations", 2)
para("The results show that the combination of store-and-forward, ACK, and "
     "an idempotent receiver can preserve data completeness and integrity "
     "under normal conditions, while opcode-based framing with flow control "
     "overcomes the BLE notification size limit. Time-based matching makes "
     "the delivery computation independent of clock synchronization.")
para("This study has several limitations. First, the evaluation is still "
     "based on a single session, so replication and varied conditions are "
     "needed to obtain representative means and standard deviations. Second, "
     "background execution covers the application being in the background and "
     "the screen off, but does not yet guarantee operation when the "
     "application is force-closed or after a device reboot, which is also "
     "constrained by OS power-saving policies. Third, the BLE communication "
     "does not yet apply encryption/authentication, so the security of "
     "health data is a development agenda. Fourth, the sensor on a consumer "
     "smartwatch is not clinically validated, so the contribution focuses on "
     "communication reliability rather than the medical accuracy of the "
     "heart-rate values.")

# ===================== 4. CONCLUSION =====================
heading("4. Conclusion", 1)
para("This study designed and implemented a BLE-based heart-rate data "
     "delivery system from a smartwatch to a smartphone using batch framing, "
     "store-and-forward, delivery acknowledgement (ACK), an idempotent "
     "receiver, and background execution. Testing on physical devices showed "
     "a 99.86% delivery ratio with no duplicates in a single ~36-minute "
     "session, with 99.1% of data at the highest accuracy, and fast batch "
     "transfer (~0.32 s for 228 readings at MTU 512). Future work includes "
     "formal measurement across varied conditions with replication, "
     "strengthening background execution (including auto-start after "
     "reboot), and adding encryption for data security.")

# ===================== ACKNOWLEDGEMENT =====================
heading("Acknowledgement", 1)
para("[[Optional: acknowledgement to advisor/institution/funding.]]")

# ===================== REFERENCES =====================
heading("References", 1)
refs = [
    "Bluetooth SIG, “Bluetooth Core Specification,” Bluetooth Special "
    "Interest Group. [Online]. Available: https://www.bluetooth.com/specifications/specs/",
    "Android Developers, “SensorManager,” Android API Reference. [Online]. "
    "Available: https://developer.android.com/reference/android/hardware/SensorManager",
    "Android Developers, “Foreground services,” Android Developers Guide. "
    "[Online]. Available: https://developer.android.com/develop/background-work/services/foreground-services",
    "“flutter_blue_plus,” pub.dev. [Online]. Available: "
    "https://pub.dev/packages/flutter_blue_plus",
    "“sqflite,” pub.dev. [Online]. Available: https://pub.dev/packages/sqflite",
    "“flutter_bloc,” pub.dev. [Online]. Available: https://pub.dev/packages/flutter_bloc",
    "[[Add citation: Internet of Medical Things (IoMT) survey/article.]]",
    "[[Add citation: wearable/BLE-based heart-rate monitoring study.]]",
    "[[Add citation: store-and-forward / data delivery reliability in IoT.]]",
    "[[Add citation: Wireless Body Area Network (WBAN) / mHealth.]]",
]
for i, r in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    p.paragraph_format.space_after = Pt(3)
    set_font(p.add_run("[%d] " % i), 10)
    set_font(p.add_run(r), 10)

note = doc.add_paragraph()
note.paragraph_format.space_before = Pt(12)
set_font(note.add_run(
    "Note: this is a draft. Adapt it to the target SINTA-2 journal template "
    "(column format, citation style, page limit). Items marked [[...]] must "
    "be completed by the author."), 9, italic=True, color=(0x80, 0x80, 0x80))

doc.save(OUT)
print("[OK] saved:", OUT)
