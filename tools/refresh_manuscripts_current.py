#!/usr/bin/env python3
"""Synchronize the current Indonesian and English manuscripts.

The DOCX files contain the authoritative five-session revision.  This script
updates stale summary text/captions and replaces embedded figures with the
current project assets without rebuilding the document from an older draft.
"""
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile
import os
import shutil
import tempfile

from docx import Document


ROOT = Path(__file__).resolve().parents[1]

ID_DOC = ROOT / "Draft_Naskah_HR_BLE_SINTA2.docx"
EN_DOC = ROOT / "Draft_Manuscript_HR_BLE_EN.docx"

FIGURE_MEDIA = {
    "image1.png": "fig_architecture.png",
    "image2.png": "fig_framing.png",
    "image3.png": "fig_sequence.png",
    "image4.png": "fig_storeforward.png",
    "image5.png": "fig_watch_arch.png",
    "image6.png": "fig_hr_completeness.png",
    "image7.png": "fig_hr_timeline.png",
    "image8.png": "fig_hr_bpm_dist.png",
    "image9.png": "fig_hr_contact.png",
}


ID_ABSTRACT = (
    "Pemantauan detak jantung secara berkelanjutan menggunakan perangkat "
    "wearable merupakan komponen penting dalam Internet of Medical Things "
    "(IoMT). Tantangan utamanya adalah menjaga kelengkapan data saat pembacaan "
    "dikirim dari smartwatch ke smartphone melalui Bluetooth Low Energy "
    "(BLE), dengan mempertimbangkan batas Maximum Transmission Unit (MTU), "
    "sifat best-effort notifikasi BLE, dan pembatasan proses latar belakang. "
    "Penelitian ini merancang sistem dua aplikasi: smartwatch Wear OS sebagai "
    "peripheral/GATT server merekam pembacaan ke SQLite dan mengirim batch "
    "ber-frame START/DATA/END, sedangkan smartphone Android sebagai "
    "central/GATT client merangkai ulang dan menyimpan data. Keandalan "
    "didukung oleh store-and-forward, ACK tingkat aplikasi, penerima "
    "idempoten, kendali aliran, dan foreground service. Evaluasi historis "
    "lima sesi pada satu pasangan perangkat menghasilkan 22.429 dari 25.191 "
    "record yang cocok di smartphone (delivery agregat deskriptif 89,04%) "
    "dengan fidelitas nilai 100% dan tanpa duplikat basis data. Empat sesi "
    "versi awal memperlihatkan kegagalan backfill akibat penandaan synced "
    "prematur. Pada Sesi 5, versi revisi yang menunggu ACK mengirim 1.974 "
    "dari 1.976 record (99,90%); dua record terakhir tetap pending. Tujuh "
    "batch berisi 180 record pada MTU 512 ditransfer rata-rata dalam "
    "121,5 ms dengan throughput 70,3 KiB/detik. Hasil mendukung kelayakan "
    "protokol pada kondisi terhubung yang diuji, tetapi ketahanan versi revisi "
    "terhadap putus–sambung masih memerlukan eksperimen terkontrol."
)

EN_ABSTRACT = (
    "Continuous heart-rate monitoring with wearable devices is an important "
    "Internet of Medical Things (IoMT) use case. A central challenge is "
    "maintaining data completeness when readings are transferred from a "
    "smartwatch to a smartphone over Bluetooth Low Energy (BLE), given the "
    "Maximum Transmission Unit (MTU), best-effort notifications, and "
    "background-execution restrictions. This study implements two "
    "applications: a Wear OS peripheral/GATT server that stores readings in "
    "SQLite and transmits START/DATA/END-framed batches, and an Android "
    "central/GATT client that reassembles and stores them. Reliability is "
    "supported by store-and-forward, application-level acknowledgements, an "
    "idempotent receiver, flow control, and a foreground service. Across five "
    "historical sessions on one device pair, 22,429 of 25,191 watch records "
    "were matched on the smartphone (a descriptive aggregate delivery ratio "
    "of 89.04%), with 100% value fidelity and no database duplicates. Four "
    "initial-version sessions exposed a backfill failure caused by premature "
    "synced marking. In Session 5, the ACK-waiting revision delivered 1,974 "
    "of 1,976 records (99.90%); the final two remained pending. Seven "
    "180-record batches at MTU 512 transferred in 121.5 ms on average at "
    "70.3 KiB/s. These results support protocol feasibility under the tested "
    "connected condition, while controlled disconnect–reconnect evaluation "
    "of the revised implementation remains necessary."
)


def replace_paragraph(paragraph, text):
    """Replace text while preserving the paragraph-level style."""
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def replace_starts(doc, replacements):
    found = set()
    for paragraph in doc.paragraphs:
        original = paragraph.text.strip()
        for prefix, replacement in replacements.items():
            if original.startswith(prefix):
                replace_paragraph(paragraph, replacement)
                found.add(prefix)
                break
    missing = set(replacements) - found
    if missing:
        raise RuntimeError(f"Paragraph prefixes not found: {sorted(missing)}")


def refresh_text(path, language):
    doc = Document(path)
    if language == "id":
        replacements = {
            "Pengujian dilakukan pada perangkat fisik karena komunikasi BLE tidak dapat diemulasikan.": (
                "Pengujian dilakukan pada perangkat fisik agar perilaku radio, "
                "sensor, dan eksekusi latar belakang yang tidak sepenuhnya "
                "tercakup oleh emulator dapat diamati. Spesifikasi perangkat "
                "uji ditunjukkan pada Tabel 1. Aplikasi dipasang dalam mode "
                "rilis agar pengukuran kinerja tidak terbias oleh mode debug. "
                "Pengujian mencakup empat sesi pada 23–28 Juni 2026 yang "
                "menggunakan implementasi awal dan satu sesi pada 5 Juli 2026 "
                "yang menggunakan revisi penantian ACK sebelum markSynced. "
                "Agregat lima sesi bersifat deskriptif historis karena "
                "mencakup dua versi perangkat lunak."
            ),
            "Atribut waktu (time) pada metrik di atas": (
                "Atribut waktu (time) adalah timestamp pencatatan aplikasi "
                "smartwatch, bukan waktu fisiologis terjadinya detak. Sekitar "
                "sekali per detik, aplikasi mencatat nilai bpm terbaru bersama "
                "waktu perangkat sebagai epoch milidetik. Pada implementasi "
                "ini, timestamp tersebut digunakan sebagai kunci unik record "
                "serta dasar pencocokan watch–smartphone dan deteksi duplikat."
            ),
            "Gambar 6. Kelengkapan data per sesi pengujian": (
                "Gambar 6. Kelengkapan data pada empat sesi evaluasi Juni "
                "(diterima vs hilang)."
            ),
            "Gambar 8. Distribusi kualitas kontak sensor pada periode pengukuran.": (
                "Gambar 8. Distribusi kualitas kontak sensor pada empat sesi "
                "evaluasi Juni."
            ),
            "Hasil menunjukkan kombinasi store-and-forward": (
                "Hasil menunjukkan bahwa framing ber-opcode dan kendali aliran "
                "mempertahankan fidelitas nilai 100% pada record yang diterima. "
                "Empat sesi versi awal mengungkap kegagalan backfill akibat "
                "penandaan synced prematur ketika koneksi tidak tersedia. "
                "Revisi ACK pada Sesi 5 menghilangkan false-sent dalam kondisi "
                "koneksi kontinu yang diuji, tetapi belum dievaluasi melalui "
                "eksperimen putus–sambung terkontrol. Karena itu, klaim "
                "ketahanan store-and-forward versi terbaru dibatasi pada "
                "kondisi yang telah diuji."
            ),
        }
    else:
        replacements = {
            "The time attribute above is the smartwatch application's recording timestamp": (
                "The time attribute is the smartwatch application's recording "
                "timestamp rather than the physiological time of a heartbeat. "
                "Approximately once per second, the application records the "
                "latest bpm value with the device time as epoch milliseconds. "
                "In this implementation, that timestamp is used as the unique "
                "record key and as the basis for watch–smartphone matching and "
                "duplicate detection."
            ),
            "Figure 6. Data completeness per test session": (
                "Figure 6. Data completeness in the four June evaluation "
                "sessions (received vs lost)."
            ),
            "Figure 8. Sensor contact-quality distribution during the measurement period.": (
                "Figure 8. Sensor contact-quality distribution in the four "
                "June evaluation sessions."
            ),
        }
    replace_starts(doc, replacements)
    doc.save(path)


def refresh_media(path):
    with tempfile.TemporaryDirectory() as tmp_name:
        tmp = Path(tmp_name)
        with ZipFile(path, "r") as zin:
            zin.extractall(tmp)
        for media_name, figure_name in FIGURE_MEDIA.items():
            target = tmp / "word" / "media" / media_name
            source = ROOT / "figures" / figure_name
            if not target.exists() or not source.exists():
                raise FileNotFoundError(target if not target.exists() else source)
            shutil.copyfile(source, target)
        replacement = path.with_suffix(".new.docx")
        with ZipFile(replacement, "w", ZIP_DEFLATED) as zout:
            for directory, _, files in os.walk(tmp):
                for filename in files:
                    item = Path(directory) / filename
                    zout.write(item, item.relative_to(tmp))
        replacement.replace(path)


def main():
    refresh_text(ID_DOC, "id")
    refresh_text(EN_DOC, "en")
    refresh_media(ID_DOC)
    refresh_media(EN_DOC)
    print(f"[OK] refreshed: {ID_DOC.name}")
    print(f"[OK] refreshed: {EN_DOC.name}")


if __name__ == "__main__":
    main()
