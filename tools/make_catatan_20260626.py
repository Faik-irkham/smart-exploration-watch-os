# -*- coding: utf-8 -*-
"""Membuat catatan harian (DOCX) untuk pembimbing — 26 Juni 2026.

Melengkapi catatan 23 Juni (gambaran umum sistem) dengan rincian protokol
komunikasi BLE pada lapisan GATT: struktur layanan, penyiapan koneksi,
berlangganan (CCCD), transfer batch berbingkai dengan flow-control, konfirmasi
penerimaan (ACK) end-to-end, dan pencegahan duplikasi. Menyisipkan gambar
figures/fig_gatt.png dan figures/fig_framing.png.

Pakai:  python3 tools/make_catatan_20260626.py
Output: docs/26-06-26-Catatan-Faik-BLE-SmartWatch.docx

PERINGATAN: berkas .docx keluaran SUDAH DISUNTING MANUAL di Word (mis. penjelasan
gambar ditulis ulang, ditambah paragraf nuansa MTU & notify-vs-indicate). Skrip
ini berisi teks LAMA; menjalankannya ulang akan MENIMPA suntingan tersebut. Karena
itu skrip menolak menimpa berkas yang sudah ada — pakai --force hanya jika benar
benar ingin membangun ulang dari nol (suntingan manual akan hilang).
"""
import os
import sys
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(ROOT, "docs", "26-06-26-Catatan-Faik-BLE-SmartWatch.docx")

doc = Document()


def title(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)


def heading(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)


def body(text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def mono(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)


def figure(rel_path, caption, width_in):
    path = os.path.join(ROOT, rel_path)
    if not os.path.exists(path):
        body(f"[Gambar tidak ditemukan: {rel_path}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width_in))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)


# ---------------------------------------------------------------------------
title("DETAIL JALUR KOMUNIKASI BLE/GATT SMARTWATCH KE SMARTPHONE "
      "(26 Juni 2026)")

body(
    "Catatan ini melengkapi catatan sebelumnya dengan rincian protokol "
    "komunikasi BLE pada lapisan GATT, yaitu bagaimana smartwatch (peripheral / "
    "GATT server) dan smartphone (central / GATT client) saling bertukar data "
    "sejak awal koneksi hingga konfirmasi penerimaan. Tujuannya agar mekanisme "
    "keandalan pengiriman data dapat dijelaskan secara rinci dan dapat "
    "direproduksi dalam paper."
)

heading("Struktur Layanan GATT")
body(
    "Smartwatch menyediakan satu layanan kustom (UUID 0000a100) yang memuat dua "
    "karakteristik. Karakteristik pertama (0000a101) bertipe notify dan "
    "dilengkapi deskriptor CCCD (0x2902); lewat karakteristik inilah paket data "
    "detak jantung dikirim ke smartphone. Karakteristik kedua (0000a102) bertipe "
    "write dan dipakai smartphone untuk menuliskan konfirmasi penerimaan "
    "(acknowledgement, ACK) atas data yang diterima. Ketiga UUID merupakan UUID "
    "128-bit yang dibangun dari alias 16-bit A100/A101/A102 pada base UUID "
    "Bluetooth."
)

heading("Penyiapan Koneksi")
body(
    "Smartwatch mengiklankan (advertising) layanan 0000a100 agar mudah "
    "ditemukan, sementara nama perangkat diletakkan pada scan response agar "
    "paket iklan tidak melampaui batas 31 byte. Smartphone memindai dengan "
    "filter UUID layanan tersebut lalu terhubung (connect). Setelah terhubung, "
    "smartphone meminta MTU besar (hingga 512 byte) supaya tiap notifikasi bisa "
    "memuat potongan data yang lebih besar, kemudian melakukan service discovery "
    "untuk menemukan karakteristik record, karakteristik ACK, dan deskriptor "
    "CCCD."
)
body(
    "Perlu dicatat bahwa 512 byte hanyalah nilai yang diminta; MTU final "
    "ditentukan melalui negosiasi dan sama dengan kemampuan terkecil di antara "
    "kedua perangkat, sehingga bisa lebih kecil dari 512. Sebagian perangkat "
    "bahkan menolak permintaan ini dan tetap memakai MTU default (23 byte). "
    "Sistem menangani kedua kemungkinan tersebut: ukuran potongan (chunk) tiap "
    "frame dihitung menyesuaikan MTU yang disepakati, yaitu maksimum MTU-4 byte. "
    "Bila MTU besar, jumlah frame sedikit dan transfer lebih cepat; bila MTU "
    "kecil, jumlah frame lebih banyak dan transfer lebih lambat, tetapi data "
    "tetap utuh dan tidak ada yang hilang. Dengan kata lain, MTU memengaruhi "
    "kecepatan/efisiensi pengiriman, bukan kebenaran datanya."
)

heading("Berlangganan (Subscribe)")
body(
    "Smartphone mengaktifkan notifikasi dengan menuliskan nilai 0x01 0x00 ke "
    "deskriptor CCCD pada karakteristik record. Di sisi smartwatch, permintaan "
    "ini ditangani lewat callback onDescriptorWriteRequest, dan perangkat "
    "smartphone dicatat sebagai pelanggan (subscriber) yang berhak menerima "
    "notifikasi. Sejak titik ini, smartwatch dapat mengirim paket data."
)

heading("Pengiriman Batch Ber-frame dengan Flow Control")
body(
    "Karena satu paket (JSON array berisi banyak pembacaan) berukuran lebih "
    "besar daripada kapasitas satu notifikasi BLE, paket dipecah menjadi "
    "beberapa frame. Setiap frame diawali satu byte opcode: START (0x01) "
    "menandai awal paket, DATA (0x02) membawa potongan JSON sebesar maksimum "
    "MTU-4 byte, dan END (0x03) menandai akhir paket. Frame dikirim satu per "
    "satu dengan flow control: frame berikutnya baru dikirim setelah sistem "
    "mengonfirmasi frame sebelumnya tuntas terkirim (callback onNotificationSent). "
    "Smartphone merangkai kembali potongan-potongan tersebut sampai menerima "
    "END, lalu men-decode JSON menjadi daftar pembacaan dan menyimpannya ke "
    "SQLite dalam satu transaksi."
)
body(
    "Notifikasi BLE bersifat fire-and-forget: pada lapisan ATT tidak ada "
    "konfirmasi penerimaan, sehingga smartwatch tidak otomatis tahu apakah paket "
    "benar-benar sampai. BLE menyediakan mekanisme alternatif bernama indication "
    "yang berkonfirmasi, tetapi konfirmasinya terjadi per pesan sehingga lebih "
    "lambat. Sistem ini sengaja memilih notify (cepat) lalu menambahkan "
    "konfirmasi sendiri di lapisan aplikasi (ACK) sekali per paket - lebih "
    "efisien daripada indication yang berkonfirmasi tiap frame. Mekanisme "
    "konfirmasi ini dijelaskan pada bagian berikut."
)

heading("Konfirmasi Penerimaan (ACK) End-to-End")
body(
    "Setelah seluruh paket tersimpan di smartphone, smartphone menuliskan jumlah "
    "record yang berhasil disimpan ke karakteristik ACK (0000a102). Smartwatch "
    "baru menandai data tersebut sebagai sudah terkirim (synced = 1) setelah "
    "menerima ACK ini. Bila ACK tidak diterima dalam tenggang waktu (30 detik) - "
    "misalnya aplikasi penerima berhenti atau koneksi terputus - data tetap "
    "berstatus belum terkirim dan akan dikirim ulang pada giliran berikutnya. "
    "Mekanisme ini menutup celah “data dianggap terkirim padahal penerima "
    "tidak benar-benar menerimanya”."
)

heading("Pencegahan Duplikasi (Idempoten)")
body(
    "Karena pengiriman ulang dapat membuat satu pembacaan terkirim lebih dari "
    "sekali, basis data smartphone diberi indeks unik pada kolom waktu (time) "
    "dan penyimpanan dilakukan dengan mode “abaikan bila sudah ada” "
    "(INSERT OR IGNORE). Dengan begitu, pengiriman ulang tidak menghasilkan "
    "baris ganda dan jumlah duplikat pada hasil ekspor tetap nol. Pencocokan "
    "antara data smartwatch dan smartphone tetap memakai kolom time (Unix epoch "
    "yang dibuat di smartwatch), sehingga perhitungan tidak bergantung pada "
    "sinkronisasi jam kedua perangkat."
)

body("Contoh isi satu paket (JSON array) yang dikirim melalui BLE:")
mono('[{"bpm": 78, "accuracy": 3, "time": 1750636800123},')
mono(' {"bpm": 79, "accuracy": 3, "time": 1750636801124}, ... ]')

figure("figures/fig_gatt.png",
       "Gambar 1. Jalur komunikasi GATT lengkap: tabel atribut, penyiapan "
       "koneksi, berlangganan (CCCD), transfer berbingkai dengan flow control, "
       "dan penulisan ACK (konfirmasi penerimaan) kembali ke smartwatch.", 4.2)

body(
    "Gambar 1 merangkum keseluruhan jalur komunikasi GATT dalam satu pandangan. "
    "Bagian atas adalah tabel atribut yang di-host smartwatch, yaitu satu "
    "layanan yang memuat karakteristik record (beserta deskriptor CCCD) dan "
    "karakteristik ACK. Bagian bawah memperlihatkan urutan pesan antara "
    "smartwatch (server) dan smartphone (client) yang dikelompokkan menjadi lima "
    "fase: (1) penyiapan layanan dan advertising di smartwatch; (2) koneksi, "
    "negosiasi MTU, dan service discovery oleh smartphone; (3) berlangganan "
    "dengan menuliskan CCCD; (4) transfer batch yang dikirim frame demi frame "
    "dengan flow control; dan (5) penulisan ACK dari smartphone yang menjadi "
    "dasar smartwatch menandai data sebagai sudah terkirim. Garis penuh "
    "menandakan aliran notifikasi/data, sedangkan garis putus-putus menandakan "
    "ACK dan balasan (response)."
)

figure("figures/fig_framing.png",
       "Gambar 2. Struktur pemecahan paket menjadi frame "
       "START / DATA / END (satu frame = satu notifikasi BLE).", 6.0)

body(
    "Gambar 2 memperlihatkan cara satu paket JSON dipecah menjadi rangkaian "
    "frame sebelum dikirim: diawali frame START, diikuti beberapa frame DATA "
    "yang masing-masing membawa potongan JSON, lalu ditutup frame END. Sisipan "
    "di bawahnya menunjukkan struktur satu frame DATA, yakni satu byte opcode "
    "(penanda jenis frame) diikuti potongan data berukuran maksimum MTU-4 byte. "
    "Setiap frame dikirim sebagai satu notifikasi BLE, dan smartphone "
    "merangkainya kembali menjadi paket utuh ketika frame END diterima."
)

heading("Catatan & Langkah Berikutnya")
body(
    "Mekanisme ACK, anti-duplikat, dan flow control di atas sudah terpasang pada "
    "kode dan lolos analisis statis. Verifikasi end-to-end pada perangkat fisik "
    "untuk skenario gangguan koneksi (perangkat menjauh atau Bluetooth sempat "
    "dimatikan) menjadi langkah pengujian berikutnya sebelum pengumpulan data "
    "eksperimen formal."
)

if os.path.exists(OUT) and "--force" not in sys.argv:
    raise SystemExit(
        f"[BATAL] {OUT}\n"
        f"        sudah ada dan mungkin berisi suntingan manual di Word.\n"
        f"        Menjalankan ulang akan menimpanya. Bila yakin ingin membangun\n"
        f"        ulang dari nol: python3 {os.path.relpath(sys.argv[0], ROOT)} --force"
    )
os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print("[OK] tersimpan:", OUT)
