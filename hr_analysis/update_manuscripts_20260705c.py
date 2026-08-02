#!/usr/bin/env python3
"""Perbaikan kesiapan SINTA 2 (yang dapat dikerjakan tanpa input eksternal):
1. Tukar Gambar 8 <-> 9 (ID) / Figure 8 <-> 9 (EN): figur kualitas kontak
   dipindah SEBELUM figur distribusi bpm dan dinomori 8, karena dirujuk lebih
   dulu di §3.3 — urutan rujukan pertama menjadi monoton naik.
2. Sitasi pustaka yang menganggur [4] flutter_blue_plus, [5] sqflite,
   [6] flutter_bloc ke dalam teks (§2.1 dan §2.4).
3. Bahasa baku ID: "fidelity nilai" -> "fidelitas nilai" (glos "value
   fidelity" dipertahankan pada definisi metrik §2.6)."""
from docx import Document
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]

def set_text(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        p.add_run(text)

def find_caption(doc, prefix):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    raise SystemExit(f"CAPTION TIDAK DITEMUKAN: {prefix!r}")

def swap_fig_blocks(doc, cap_a_prefix, cap_b_prefix):
    """Pindahkan blok figur B (gambar+caption+penjelasan) ke sebelum blok A,
    lalu tukar nomor pada kedua caption."""
    cap_a = find_caption(doc, cap_a_prefix)   # figur yang semula lebih dulu
    cap_b = find_caption(doc, cap_b_prefix)   # figur yang akan dipindah ke depan
    img_a = cap_a._p.getprevious()
    img_b = cap_b._p.getprevious()
    expl_b = cap_b._p.getnext()
    # pindahkan [img_b, cap_b, expl_b] ke sebelum img_a (urutan terjaga)
    img_a.addprevious(img_b)
    img_a.addprevious(cap_b._p)
    img_a.addprevious(expl_b)
    # tukar nomor caption
    a_text, b_text = cap_a.text, cap_b.text
    set_text(cap_a, a_text.replace(cap_a_prefix, cap_b_prefix, 1))
    set_text(cap_b, b_text.replace(cap_b_prefix, cap_a_prefix, 1))

def replace_in_paragraphs(doc, replacements):
    done = {old: False for old, _ in replacements}
    for p in doc.paragraphs:
        t = p.text
        hit = False
        for old, new in replacements:
            if old in t:
                t = t.replace(old, new)
                done[old] = True
                hit = True
        if hit:
            set_text(p, t)
    missing = [k for k, v in done.items() if not v]
    if missing:
        raise SystemExit(f"TEKS TIDAK DITEMUKAN: {missing}")

# ============================ NASKAH ID ============================
doc = Document(str(ROOT / "Draft_Naskah_HR_BLE_SINTA2.docx"))
swap_fig_blocks(doc, "Gambar 8.", "Gambar 9.")
replace_in_paragraphs(doc, [
    # rujukan §3.3 mengikuti penomoran baru
    ("ditunjukkan pada Tabel 3 dan Gambar 9", "ditunjukkan pada Tabel 3 dan Gambar 8"),
    ("(n=21.087, Gambar 8)", "(n=21.087, Gambar 9)"),
    ("Sebaran nilai detak jantung pada pembacaan terbaik serta komposisi kualitas kontak tersebut "
     "divisualkan pada Gambar 8 dan Gambar 9.",
     "Komposisi kualitas kontak tersebut serta sebaran nilai detak jantung pada pembacaan terbaik "
     "divisualkan pada Gambar 8 dan Gambar 9."),
    # sitasi pustaka [4][5][6] ke dalam teks
    ("Smartphone berperan sebagai central/GATT client:",
     "Smartphone berperan sebagai central/GATT client (pustaka flutter_blue_plus) [4]:"),
    ("disimpan ke basis data lokal SQLite",
     "disimpan ke basis data lokal SQLite (pustaka sqflite) [5]"),
    ("Lapisan Dart menggunakan pola BLoC:",
     "Lapisan Dart menggunakan pola BLoC [6]:"),
    # bahasa baku: fidelitas
    ("fidelity nilai", "fidelitas nilai"),
])
# glos EN pada definisi metrik §2.6
replace_in_paragraphs(doc, [
    ("(2) fidelitas nilai, yaitu proporsi",
     "(2) fidelitas nilai (value fidelity), yaitu proporsi"),
])
doc.save(str(ROOT / "Draft_Naskah_HR_BLE_SINTA2.docx"))
print("OK -> Draft_Naskah_HR_BLE_SINTA2.docx")

# ============================ MANUSCRIPT EN ============================
doc = Document(str(ROOT / "Draft_Manuscript_HR_BLE_EN.docx"))
swap_fig_blocks(doc, "Figure 8.", "Figure 9.")
replace_in_paragraphs(doc, [
    ("shown in Table 3 and Figure 9", "shown in Table 3 and Figure 8"),
    ("(n=21,087, Figure 8)", "(n=21,087, Figure 9)"),
    ("The distribution of heart-rate values for the best readings and the composition of contact "
     "quality are visualised in Figure 8 and Figure 9.",
     "The composition of contact quality and the distribution of heart-rate values for the best "
     "readings are visualised in Figure 8 and Figure 9."),
    ("The smartphone acts as a central/GATT client:",
     "The smartphone acts as a central/GATT client (flutter_blue_plus library) [4]:"),
    ("stored in a local SQLite database",
     "stored in a local SQLite database (sqflite library) [5]"),
    ("The Dart layer uses the BLoC pattern:",
     "The Dart layer uses the BLoC pattern [6]:"),
])
doc.save(str(ROOT / "Draft_Manuscript_HR_BLE_EN.docx"))
print("OK -> Draft_Manuscript_HR_BLE_EN.docx")
