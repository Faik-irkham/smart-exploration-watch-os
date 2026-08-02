#!/usr/bin/env python3
"""Catatan 3 Juli 2026 — GABUNGAN, gaya ilmiah namun tetap terbaca orang awam.
Menggabungkan seluruh isi catatan 28 Juni (hasil uji + temuan backfill) ke dalam
satu dokumen, ditambah penjelasan asal timestamp + diagram, cara menghitung
data diterima/hilang, dan klarifikasi dua angka (2.760 vs 2.879).

Register bahasa: baku/ilmiah (smartwatch, ponsel, valid, duplikat), dengan
penjelasan singkat dalam kurung untuk istilah teknis agar tetap ramah awam.
Pola tiap visual: pengantar -> tabel/gambar -> paragraf penjelasan.

Urutan: pengambilan data -> pembuangan data off-wrist -> asal timestamp
(+diagram) -> cara hitung diterima/hilang -> hasil pengiriman (tabel+figur)
-> kualitas sensor -> temuan backfill -> klarifikasi angka -> ringkasan ->
tindak lanjut naskah.

Nomor: Tabel 1 = hasil per sesi, Tabel 2 = akurasi, Tabel 3 = dua angka;
Gambar 1 = timestamp, Gambar 2 = kelengkapan, Gambar 3 = timeline.
Memakai catatan 28 Juni sebagai template agar gaya/font konsisten."""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH as ALIGN
from docx.oxml.ns import qn
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "docs" / "26-06-28-Catatan-Faik-BLE-SmartWatch.docx"
OUT = ROOT / "docs" / "26-07-03-Catatan-Faik-BLE-SmartWatch-sederhana.docx"
FIG = ROOT / "figures"

d = Document(str(SRC))
body = d.element.body
for child in list(body):                       # kosongkan isi, simpan sectPr
    if child.tag == qn('w:sectPr'):
        continue
    body.remove(child)

def head(text, size=12):
    p = d.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(size)
    return p

def body_p(text):
    return d.add_paragraph(text)

def caption(text):
    p = d.add_paragraph(text); p.alignment = ALIGN.CENTER
    p.runs[0].font.size = Pt(10); p.runs[0].italic = True
    return p

def figure(img, cap, width):
    p = d.add_paragraph(); p.alignment = ALIGN.CENTER
    p.add_run().add_picture(str(FIG / img), width=Inches(width))
    caption(cap)

def table(rows, header_bold=True, total_bold=True):
    t = d.add_table(rows=len(rows), cols=len(rows[0])); t.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            c = t.rows[i].cells[j]; c.text = v
            if (header_bold and i == 0) or (total_bold and row[0] == "Total"):
                for rn in c.paragraphs[0].runs: rn.bold = True
    return t

# ============================== ISI ==============================
head("UJI PENGIRIMAN DATA DETAK JANTUNG", 13)

# ---------------- pengambilan data + kebersihan ----------------
head("Pengambilan Data")
body_p("Data diambil langsung dari kedua perangkat melalui koneksi adb (jalur penghubung antara komputer "
       "dan perangkat Android), dengan aplikasi versi rilis (versi final sebagaimana dipakai pengguna, "
       "bukan versi pengembangan). Berkas yang dianalisis adalah hasil ekspor kedua aplikasi dalam format "
       "CSV (tabel data berbentuk teks) yang disalin dari folder Download: watch-2026-06-28.csv (45.446 "
       "baris) dari smartwatch Samsung Galaxy Watch (SM-R860), dan phone-2026-06-28.csv (42.562 baris) "
       "dari ponsel Xiaomi Redmi Note 10 Pro. Satu baris pada berkas tersebut mewakili satu data "
       "pembacaan detak jantung. Rekaman mencakup 23–28 Juni 2026 dan, berdasarkan jeda waktu "
       "antar-pembacaan, terbagi menjadi empat sesi (periode perekaman yang terpisah) dengan kondisi "
       "berbeda: ada sesi singkat, sesi panjang (Sesi 3, sekitar 5,5 jam), dan satu sesi yang ponselnya "
       "gagal terhubung (Sesi 2).")
body_p("Sebagai gambaran bentuk datanya, Tabel 1 menampilkan tiga baris pertama dari berkas smartwatch.")
table([
    ("id", "bpm", "accuracy", "time_ms", "time_iso", "synced"),
    ("1", "76.0", "3", "1782171536608", "2026-06-23T06:38:56.608", "1"),
    ("2", "76.0", "-1", "1782171537606", "2026-06-23T06:38:57.606", "1"),
    ("3", "76.0", "-1", "1782171538606", "2026-06-23T06:38:58.606", "1"),
], total_bold=False)
caption("Tabel 1. Contoh isi berkas data smartwatch (tiga baris pertama).")
body_p("Kolom id adalah nomor urut data; bpm adalah nilai detak jantung; accuracy adalah kualitas kontak "
       "sensor (di contoh ini baris pertama berkontak baik, dua baris berikutnya sempat kehilangan "
       "kontak). Kolom time_ms dan time_iso berisi timestamp yang sama dalam dua bentuk: time_ms berupa "
       "bilangan milidetik (dibahas pada bagian berikutnya), sedangkan time_iso adalah bentuk yang mudah "
       "dibaca manusia — terlihat ketiga baris berjarak sekitar satu detik (06:38:56, 06:38:57, "
       "06:38:58). Kolom synced adalah penanda status kirim (1 berarti sudah ditandai terkirim), yang "
       "akan berperan pada bagian Temuan.")

head("Membuang Data Saat Smartwatch Tidak Dipakai")
body_p("Hampir separuh rekaman (22.231 baris) ternyata tidak valid: nilai detaknya konstan (beku) di "
       "angka 106 bpm dengan kualitas kontak yang buruk. Kualitas kontak ini tercatat pada nilai "
       "accuracy, yaitu penanda seberapa baik sensor menempel ke kulit: 3 berarti kontak baik, 0 berarti "
       "sedang, dan -1 berarti tidak ada kontak. Kondisi nilai beku tersebut terjadi saat smartwatch "
       "dilepas dari pergelangan, tetapi aplikasinya masih merekam; karena bukan detak jantung yang "
       "sebenarnya, data itu dikeluarkan dari perhitungan. Sisanya, yaitu data pengukuran yang valid, "
       "berjumlah 23.215 baris. Kriteria pemotongannya: setiap sesi dipotong sampai data dengan kontak "
       "baik (accuracy = 3) yang terakhir.")

# ---------------- asal timestamp ----------------
head("Dari Mana Timestamp Berasal")
body_p("Data yang dihasilkan sensor detak jantung hanya berisi dua nilai, yaitu bpm (beats per minute, "
       "jumlah detak jantung per menit) dan tingkat akurasi. Timestamp (kolom time pada database, berisi "
       "penanda waktu tiap data) tidak berasal dari sensor, melainkan dibuat oleh aplikasi pada saat data "
       "disimpan. Alur lengkapnya, dari sensor sampai tersimpan di database (tempat penyimpanan data di "
       "dalam perangkat), ditunjukkan pada Gambar 1.")
figure("fig_timestamp.png",
       "Gambar 1. Alur dari sensor sampai tersimpan di database — timestamp dibuat di sisi aplikasi "
       "(langkah 5), bukan dari sensor.", 4.3)
body_p("Sensor detak jantung ditangani oleh kode native/Kotlin, yaitu bagian program yang berhubungan "
       "langsung dengan perangkat keras; bagian ini hanya menghasilkan dua angka — nilai bpm dan tingkat "
       "akurasi — tanpa informasi waktu yang menyertainya. Kedua angka itu kemudian diteruskan ke lapisan "
       "aplikasi (Flutter/Dart, bagian program yang mengatur tampilan dan logika aplikasi) melalui jalur "
       "penghubung antar-lapisan yang disebut platform channel; sampai tahap ini pun belum ada informasi "
       "waktu, dan aplikasi hanya menyimpan nilai terbaru untuk sementara. Aplikasi lalu menjalankan "
       "pengatur waktu (timer) yang aktif setiap 1 detik: pada tiap detiknya, jika tersedia nilai bpm "
       "yang sah, aplikasi membentuk satu baris data — pada saat itulah timestamp ditambahkan dengan "
       "membaca jam perangkat (melalui perintah DateTime.now()). Timestamp tersebut disimpan sebagai "
       "bilangan milidetik terhitung sejak 1 Januari 1970 (format epoch, cara komputer menuliskan waktu "
       "agar mudah diurutkan), misalnya 1750636800123.")
body_p("Dengan demikian, timestamp dibuat di sisi aplikasi pada saat data disimpan, bukan berasal dari "
       "sensor. Maknanya adalah “kapan data dicatat oleh aplikasi”, bukan waktu persis satu detak jantung "
       "terjadi. Karena timer berjalan setiap 1 detik, timestamp antar-data berjarak sekitar satu detik "
       "dan tidak pernah sama (unik).")

head("Cara Menghitung Data yang Diterima dan yang Hilang")
body_p("Timestamp inilah yang menjadi kunci untuk menentukan sebuah data “diterima” atau “hilang”. "
       "Setiap data yang dicatat smartwatch telah memiliki timestamp. Data kemudian dikirim ke ponsel "
       "melalui Bluetooth secara berkala (setiap beberapa menit) dalam satu kumpulan (batch), dan "
       "timestamp ikut terkirim bersama nilai bpm dan akurasinya — sehingga satu data yang sama memiliki "
       "timestamp identik di kedua perangkat. Daftar data di smartwatch dan di ponsel kemudian "
       "dicocokkan berdasarkan timestamp — bukan dengan membandingkan jam kedua perangkat — sehingga "
       "perbedaan setelan jam tidak memengaruhi hasil. Data yang timestamp-nya ada di smartwatch tetapi "
       "tidak ditemukan di ponsel dihitung “hilang”, sedangkan yang ada di keduanya dihitung “diterima”. "
       "Inilah dasar kolom “Diterima” dan “Hilang” pada Tabel 3.")

# ---------------- ukuran data & MTU ----------------
head("Ukuran Data yang Dikirim dan Peran MTU")
body_p("Sebelum dikirim, setiap record dikemas dalam format JSON (format teks untuk bertukar data antar "
       "program) berbentuk {\"bpm\":76.0,\"accuracy\":3,\"time\":1782171536608}, berukuran sekitar 46–47 "
       "byte. Dengan pencatatan satu data per detik, satu batch berisi sekitar 180 record (±8 KB) untuk "
       "interval pengiriman 3 menit, atau sekitar 300 record (±14 KB) untuk interval 5 menit; bila ada "
       "data tertunda yang menumpuk, batch bisa lebih besar. Sebagai angka nyata dari pengujian, satu "
       "batch berisi 228 record berukuran 10.717 byte. Angka-angka ini berasal dari log metrik aplikasi: "
       "setiap selesai mengirim satu batch, aplikasi smartwatch menulis satu baris berpenanda HR-METRIC "
       "ke log sistem Android (Logcat, dibaca melalui adb) yang memuat jumlah record, ukuran byte, jumlah "
       "frame, MTU, durasi, dan kecepatan — contohnya: tx_batch,228,10717,24,512,323.1,33168.")
body_p("Karena satu paket Bluetooth ukurannya terbatas, batch tersebut tidak dikirim sekaligus, melainkan "
       "dipotong menjadi rangkaian frame (paket kecil yang dikirim berurutan). Batas ukuran paket ini "
       "disebut MTU (Maximum Transmission Unit); tiap frame memuat potongan data sebesar MTU dikurangi 4 "
       "byte. Ponsel meminta MTU 512, tetapi setiap smartwatch dapat menyetujui nilai yang berbeda "
       "sesuai kemampuannya — yang dipakai adalah nilai terkecil dari keduanya, dan bila negosiasi gagal "
       "dipakai nilai bawaan 23. Pada smartwatch yang diuji (Galaxy Watch SM-R860), nilai 512 disetujui, "
       "terbukti dari log metrik di atas (kolom mtu = 512, dan 10.717 byte terbagi menjadi 24 frame). "
       "Tabel 2 memperlihatkan dampak MTU terhadap jumlah frame untuk batch nyata yang sama.")
table([
    ("MTU yang disetujui", "Ukuran potongan (MTU − 4)", "Total frame (batch 10.717 byte)"),
    ("512 (smartwatch yang diuji)", "508 byte", "24"),
    ("247", "243 byte", "47"),
    ("185", "181 byte", "62"),
    ("23 (bawaan, bila negosiasi gagal)", "19 byte", "567"),
], total_bold=False)
caption("Tabel 2. Pengaruh MTU terhadap jumlah frame untuk batch yang sama (10.717 byte, 228 record).")
body_p("Terlihat bahwa MTU hanya memengaruhi banyaknya frame dan kecepatan transfer: MTU kecil berarti "
       "potongan kecil dan frame lebih banyak, sehingga pengiriman lebih lama. Kelengkapan data tidak "
       "terpengaruh — berapa pun MTU yang disetujui, seluruh 228 record tetap sampai utuh, karena "
       "kelengkapan dijaga oleh mekanisme pengiriman ulang dan pencocokan timestamp, bukan oleh ukuran "
       "paket.")
body_p("Seluruh angka pada bagian ini dapat diturunkan dari tiga fakta dasar. Pertama, ukuran rata-rata "
       "satu record adalah 47 byte, diperoleh dari log metrik: 10.717 byte ÷ 228 record = 47,0 byte. "
       "Kedua, aplikasi mencatat satu record per detik, sehingga interval 3 menit menghasilkan 3 × 60 = "
       "180 record (180 × 47 = 8.460 byte ≈ 8,3 KB) dan interval 5 menit menghasilkan 300 record (300 × "
       "47 = 14.100 byte ≈ 13,8 KB). Ketiga, tiap frame memuat potongan sebesar MTU − 4 byte, sehingga "
       "batch 10.717 byte pada MTU 512 membutuhkan 10.717 ÷ 508 = 21,1 — dibulatkan ke atas menjadi 22 "
       "frame DATA — ditambah frame START dan END menjadi 24 frame. Hasil hitungan ini persis sama "
       "dengan yang tercatat di log (24 frame), yang sekaligus membuktikan rumusnya akurat. Dengan cara "
       "yang sama, beban terburuk dapat diperkirakan: bila ponsel terputus selama satu jam, tumpukan "
       "berisi 3.600 record × 47 byte ≈ 165 KB (sekitar 335 frame), yang pada kecepatan terukur 33.168 "
       "byte per detik (kolom terakhir log metrik yang sama) membutuhkan sekitar 5 detik untuk terkirim "
       "ulang.")
body_p("Rangkaian ukuran tersebut — dari satu record, terkumpul menjadi satu batch, hingga dipotong "
       "menjadi frame sesuai MTU — divisualkan pada Gambar 2.")
figure("fig_batch_mtu.png",
       "Gambar 2. Ukuran data yang dikirim dan pengaruh MTU: satu record (±47 byte) terkumpul menjadi "
       "satu batch (228 record = 10.717 byte), lalu dipotong menjadi frame sesuai MTU yang disetujui.",
       6.3)
body_p("Pada gambar terlihat batch yang sama dipotong dengan dua cara: pada MTU 512 potongannya besar "
       "sehingga cukup 24 frame (terukur ±0,32 detik), sedangkan pada MTU 23 potongannya kecil sehingga "
       "diperlukan 567 frame. Meski jumlah frame jauh berbeda, isi yang sampai tetap sama-sama lengkap "
       "— 228 record utuh pada kedua kasus.")

# ---------------- hasil pengiriman ----------------
head("Hasil Pengiriman Data")
body_p("Hasil pengiriman pada setiap sesi dirangkum pada Tabel 3.")
table([
    ("Sesi", "Mulai (WIB)", "Durasi", "Direkam", "Diterima", "Hilang", "Delivery"),
    ("Sesi 1", "23/06 06:38", "36 mnt", "2.177", "2.164", "13", "99,40%"),
    ("Sesi 2", "23/06 08:41", "1 mnt", "53", "0", "53", "0,00%"),
    ("Sesi 3", "23/06 14:15", "345 mnt", "20.702", "18.013", "2.689", "87,01%"),
    ("Sesi 4", "28/06 21:29", "5 mnt", "283", "278", "5", "98,23%"),
    ("Total", "—", "387 mnt", "23.215", "20.455", "2.760", "88,11%"),
])
caption("Tabel 3. Hasil pengiriman per sesi (periode pengukuran).")
body_p("Secara keseluruhan, smartwatch merekam 23.215 data dan 20.455 di antaranya sampai ke ponsel, "
       "sehingga tingkat keberhasilan pengiriman (delivery) mencapai 88,11% tanpa data ganda (duplikat). "
       "Tingkat keberhasilan antar-sesi berbeda cukup jauh: sesi dengan ponsel terhubung penuh mencapai "
       "98–99%, sedangkan satu sesi singkat yang ponselnya tidak pernah terhubung (Sesi 2) kehilangan "
       "seluruh datanya.")
body_p("Perbandingan jumlah data yang diterima dan yang hilang pada tiap sesi ditampilkan pada Gambar 3.")
figure("fig_hr_completeness.png", "Gambar 3. Kelengkapan data tiap sesi (diterima vs hilang).", 5.6)
body_p("Dari gambar terlihat bahwa kehilangan terbesar terjadi pada sesi terpanjang (Sesi 3) dan sesi "
       "yang gagal terhubung (Sesi 2), sementara sesi lainnya hampir lengkap.")
body_p("Pola kehilangan data terhadap waktu pada sesi terpanjang ditunjukkan pada Gambar 4.")
figure("fig_hr_timeline.png", "Gambar 4. Grafik detak jantung dan titik data yang hilang pada sesi "
       "terpanjang; data hanya hilang sebelum ponsel terhubung.", 6.3)
body_p("Grafik memperjelas bahwa kehilangan hanya terjadi di awal sesi, yaitu sebelum ponsel terhubung "
       "sekitar pukul 15.00; setelah terhubung, hampir tidak ada data yang hilang. Selain itu, seluruh "
       "data yang diterima isinya sama persis dengan catatan smartwatch (nilai bpm dan akurasi), sehingga "
       "tidak ada data yang rusak selama pengiriman.")

head("Kualitas Data Sensor")
body_p("Sebaran kualitas kontak sensor pada data pengukuran dirangkum pada Tabel 4.")
table([
    ("Akurasi sensor", "Jumlah", "Persentase"),
    ("3 (tinggi)", "21.087", "90,8%"),
    ("0 (sedang)", "15", "0,1%"),
    ("-1 (tanpa kontak)", "2.113", "9,1%"),
    ("Total", "23.215", "100%"),
])
caption("Tabel 4. Distribusi nilai akurasi sensor (periode pengukuran).")
body_p("Sebagian besar data (90,8%) memiliki kontak baik (accuracy 3); sisanya 9,1% tanpa kontak (-1) "
       "dan 0,1% kontak sedang (0). Pada 21.087 data dengan kontak baik, detak jantung terendah 60 bpm, "
       "tertinggi 123 bpm, dan rata-rata 83,4 bpm, dengan standard deviation 9,9 bpm (ukuran seberapa "
       "jauh nilai menyebar dari rata-rata; nilai sebesar ini masih wajar untuk aktivitas santai).")

# ---------------- temuan backfill ----------------
head("TEMUAN PENTING — Data Tertunda Tidak Dikirim Ulang Saat Koneksi Putus-Sambung")
body_p("Temuan ini melengkapi catatan 26 Juni, khususnya bagian ACK (balasan tanda terima dari ponsel "
       "yang menyatakan data sudah diterima) dan store-and-forward (pola “simpan dulu di smartwatch, "
       "kirim kemudian”). Analisis silang antara penanda synced di smartwatch dan keberadaan data di "
       "ponsel — dihitung pada periode pengukuran yang sama dengan Tabel 3 — menunjukkan: dari 2.760 "
       "data yang hilang, 2.755 di antaranya (99,8%) ternyata sudah ditandai “terkirim” (synced = 1) "
       "padahal tidak ada di ponsel; hanya 5 data terakhir yang memang masih menunggu giliran kirim "
       "(wajar, menunggu jadwal pengiriman berikutnya). Dengan demikian angka hilang pada Tabel 3 "
       "terurai persis: 2.755 + 5 = 2.760. Seluruh 2.755 data itu hilang pada periode ketika ponsel "
       "sedang terputus — termasuk 53 data pada Sesi 2 yang ponselnya tidak pernah terhubung sama "
       "sekali, bukti paling jelas bahwa penandaan dilakukan terlalu dini.")
body_p("Artinya, pada implementasi saat ini smartwatch menandai data “terkirim” begitu data dikirimkan, "
       "belum benar-benar menunggu balasan ACK dari ponsel ketika koneksi terputus. Akibatnya, data yang "
       "menumpuk selama koneksi terputus tidak otomatis dikirim ulang saat tersambung kembali. Dengan "
       "kata lain, mekanisme ini bekerja baik selama koneksi tersambung, tetapi belum tahan terhadap "
       "kondisi putus-sambung.")
body_p("Saran perbaikan: (1) data tidak ditandai “terkirim” sebelum balasan ACK benar-benar diterima; "
       "(2) menambahkan proses kirim ulang otomatis yang memeriksa data belum terkonfirmasi setiap kali "
       "koneksi tersambung kembali; (3) menggunakan tiga status, yaitu belum dikirim, sudah dikirim "
       "tetapi belum dibalas, dan sudah dipastikan diterima. Intinya, konfirmasi pada tingkat aplikasi "
       "(ACK) lebih dapat diandalkan daripada hanya mengandalkan notifikasi BLE (Bluetooth Low Energy, "
       "jenis Bluetooth hemat daya yang dipakai untuk pengiriman ini), karena notifikasi BLE tidak "
       "memberi kepastian bahwa data benar-benar sampai.")

# ---------------- klarifikasi angka ----------------
head("Hubungan Angka 2.755, 2.760, dan 2.879")
body_p("Dalam analisis ini muncul tiga angka yang saling berkaitan. Agar tidak membingungkan, hubungan "
       "ketiganya dirangkum pada Tabel 5.")
table([
    ("Angka", "Arti", "Cakupan data"),
    ("2.760", "Total data hilang (kolom “Hilang” pada Tabel 3)",
     "Periode pengukuran valid (23.215 baris)"),
    ("2.755", "Bagian dari 2.760 yang keliru ditandai terkirim; 5 sisanya pending wajar",
     "Periode pengukuran valid (23.215 baris)"),
    ("2.879", "Data bertanda terkirim yang tidak ada di ponsel bila seluruh rekaman ikut dihitung",
     "Seluruh rekaman, termasuk saat jam dilepas (45.446 baris)"),
], total_bold=False)
caption("Tabel 5. Hubungan tiga angka pada analisis kehilangan data.")
body_p("Angka 2.879 diperoleh dari selisih seluruh isi rekaman: 45.446 baris di smartwatch dikurangi "
       "42.562 baris di ponsel menghasilkan 2.884, lalu dikurangi 5 data pending menjadi 2.879. Adapun "
       "2.755 adalah bagian dari 2.879 yang berada pada periode pengukuran valid — dan bersama 5 data "
       "pending menyusun persis total hilang 2.760 pada Tabel 3. Ketiganya konsisten; perbedaannya "
       "semata pada cakupan data yang dihitung. Pada naskah, angka 2.755 (cakupan seragam dengan tabel "
       "hasil) dipakai sebagai angka utama temuan.")

head("Ringkasan")
body_p("Timestamp tidak berasal dari sensor, melainkan dibuat aplikasi pada saat data disimpan (sekitar "
       "setiap 1 detik) dan berfungsi sebagai penanda unik tiap data sekaligus dasar pencocokan data "
       "smartwatch–ponsel. Hasil uji empat sesi menunjukkan dari 23.215 data pengukuran, 20.455 diterima "
       "ponsel (88,11%) tanpa duplikat, dan isi data yang diterima 100% sama dengan catatan smartwatch. "
       "Kehilangan data tidak merata: hampir nol saat koneksi tersambung dan menumpuk saat terputus, "
       "dengan temuan penting bahwa 2.755 dari 2.760 data yang hilang (99,8%) ternyata keliru ditandai "
       "terkirim sehingga tidak dikirim ulang saat koneksi pulih (perlu penguatan ACK dan proses kirim "
       "ulang/backfill). Adapun perbedaan angka 2.755, 2.760, dan 2.879 semata soal cakupan data yang "
       "dihitung, bukan kesalahan.")

d.save(str(OUT))
print("OK ->", OUT.relative_to(ROOT))
print("paragraphs:", len(d.paragraphs), "tables:", len(d.tables),
      "images:", len(d.element.body.findall('.//'+qn('pic:pic'))))
