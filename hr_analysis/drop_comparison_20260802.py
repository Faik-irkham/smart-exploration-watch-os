#!/usr/bin/env python3
"""Hapus subbab perbandingan dengan penelitian terdahulu dari Pembahasan.

Yang dihapus: judul subbab, paragraf pengantar, caption Tabel 4, tabel
perbandingan itu sendiri, dan dua paragraf narasinya. Sub-subbab sesudahnya
dinomori ulang sehingga Pembahasan menjadi:

    F. Pembahasan
       1) Interpretasi Hasil
       2) Implikasi Penelitian
       3) Keterbatasan

Pemosisian terhadap studi terdahulu tetap ada di Bagian 1, jadi tidak ada
substansi yang hilang — hanya pengulangannya di Pembahasan yang dibuang.
"""
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]


def set_text(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        p.add_run(text)


def drop(path, prefixes, table_caption, renumber):
    doc = Document(str(path))

    # --- tabel perbandingan -------------------------------------------------
    hit = [t for t in doc.tables
           if t.rows[0].cells[0].text.strip() in ("Studi", "Study")]
    if len(hit) != 1:
        raise SystemExit(f"{path.name}: tabel perbandingan tidak unik ({len(hit)})")
    hit[0]._tbl.getparent().remove(hit[0]._tbl)

    # --- paragraf terkait ---------------------------------------------------
    removed = 0
    for p in list(doc.paragraphs):
        t = p.text.strip()
        if any(t.startswith(pre) for pre in prefixes) or t.startswith(table_caption):
            p._p.getparent().remove(p._p)
            removed += 1
    if removed != len(prefixes) + 1:
        raise SystemExit(f"{path.name}: terhapus {removed}, diharapkan {len(prefixes)+1}")

    # --- penomoran ulang sub-subbab ----------------------------------------
    for p in doc.paragraphs:
        t = p.text.strip()
        for old, new in renumber:
            if t == old:
                set_text(p, new)

    doc.save(str(path))
    print(f"OK — {path.name}: tabel + {removed} paragraf dihapus")


drop(ROOT / "Draft_Naskah_HR_BLE_SINTA2.docx",
     prefixes=[
         "2) Perbandingan dengan Penelitian Terdahulu",
         "Posisi penelitian ini terhadap studi sejenis dirangkum pada Tabel 4",
         "Tabel tersebut memperlihatkan pembagian kerja yang cukup jelas",
         "Perbedaan yang paling menentukan terletak pada kolom terakhir Tabel 4",
     ],
     table_caption="Tabel 4. Posisi penelitian ini terhadap studi sejenis.",
     renumber=[("3) Implikasi Penelitian", "2) Implikasi Penelitian"),
               ("4) Keterbatasan", "3) Keterbatasan")])

drop(ROOT / "Draft_Manuscript_HR_BLE_EN_v2.docx",
     prefixes=[
         "2) Comparison with Related Studies",
         "The position of this study relative to comparable work is summarised in Table 4",
         "The table reveals a fairly clear division of labour among prior studies",
         "The decisive difference lies in the last column of Table 4",
     ],
     table_caption="Table 4. Position of this study relative to comparable work.",
     renumber=[("3) Implications", "2) Implications"),
               ("4) Limitations", "3) Limitations")])
