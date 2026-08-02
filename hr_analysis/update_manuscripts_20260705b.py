#!/usr/bin/env python3
"""Tambahkan paragraf penjelasan SETELAH setiap tabel di kedua manuskrip,
sehingga tidak ada tabel yang langsung ditempel gambar/judul subbab.
Tabel diidentifikasi berdasarkan urutan (Tabel/Table 1..3)."""
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]

def add_after_tables(path, texts):
    doc = Document(str(path))
    tables = doc.tables
    if len(tables) != len(texts):
        raise SystemExit(f"{path.name}: ditemukan {len(tables)} tabel, "
                         f"diharapkan {len(texts)}")
    parent = doc.paragraphs[0]._parent
    for tbl, text in zip(tables, texts):
        new_p = OxmlElement("w:p")
        tbl._tbl.addnext(new_p)
        Paragraph(new_p, parent).add_run(text)
    doc.save(str(path))
    print(f"OK -> {path.name} (+{len(texts)} penjelasan tabel)")

# ============================ NASKAH ID ============================
add_after_tables(ROOT / "Draft_Naskah_HR_BLE_SINTA2.docx", [
    # Tabel 1 — spesifikasi perangkat & lingkungan uji (§2.5)
    "Tabel tersebut merangkum lingkungan pengujian yang dijaga tetap: sepasang perangkat konsumer "
    "(Samsung Galaxy Watch SM-R860 sebagai peripheral dan Xiaomi Redmi Note 10 Pro sebagai central), "
    "kerangka kerja Flutter dengan aplikasi terpasang dalam mode rilis, MTU yang diminta 512 byte, "
    "serta interval pengiriman yang dapat dipilih 3/5 menit. Dataset akhir mencakup empat sesi pada "
    "23–28 Juni 2026 dengan total 23.215 pembacaan pengukuran.",
    # Tabel 2 — hasil per sesi (§3.2)
    "Dari tabel terlihat rentang delivery antar-sesi sangat lebar (0,00–99,40%): dua sesi dengan "
    "koneksi penuh hampir lengkap, Sesi 3 (terpanjang) turun ke 87,01% akibat periode awal tanpa "
    "koneksi, dan Sesi 2 kehilangan seluruh 53 pembacaannya karena smartphone tidak pernah terhubung. "
    "Profil temporal kehilangan tersebut divisualkan pada Gambar 6 dan Gambar 7.",
    # Tabel 3 — distribusi akurasi (§3.3)
    "Dari tabel terlihat kualitas kontak didominasi tingkat tertinggi (90,8%), dengan porsi tanpa "
    "kontak 9,1% dan tingkat sedang hanya 0,1%. Sebaran nilai detak jantung pada pembacaan terbaik "
    "serta komposisi kualitas kontak tersebut divisualkan pada Gambar 8 dan Gambar 9.",
])

# ============================ MANUSCRIPT EN ============================
add_after_tables(ROOT / "Draft_Manuscript_HR_BLE_EN.docx", [
    "The table summarises the controlled test environment: a consumer device pair (Samsung Galaxy "
    "Watch SM-R860 as the peripheral and Xiaomi Redmi Note 10 Pro as the central), the Flutter "
    "framework with the applications installed in release mode, a requested MTU of 512 bytes, and a "
    "selectable 3/5-minute sending interval. The final dataset covers four sessions during 23–28 June "
    "2026 with 23,215 measurement readings in total.",
    "The table shows a wide range of per-session delivery (0.00–99.40%): the two fully connected "
    "sessions are nearly complete, Session 3 (the longest) drops to 87.01% due to its initial "
    "disconnected period, and Session 2 lost all of its 53 readings because the smartphone never "
    "connected. The temporal profile of these losses is visualised in Figure 6 and Figure 7.",
    "The table shows that contact quality is dominated by the highest level (90.8%), with 9.1% at "
    "no-contact and only 0.1% at the medium level. The distribution of heart-rate values for the best "
    "readings and the composition of contact quality are visualised in Figure 8 and Figure 9.",
])
