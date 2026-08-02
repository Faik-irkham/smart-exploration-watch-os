#!/usr/bin/env python3
"""Buat versi BAHASA SEDERHANA dari catatan 26 & 28 Juni.
File asli tidak diubah; gambar & tabel ikut terbawa karena hanya teks
paragraf yang diganti."""
from docx import Document
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1] / "docs"

def set_text(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        p.add_run(text)

def rewrite(src, dst, repl):
    d = Document(str(ROOT / src))
    for i, txt in repl.items():
        set_text(d.paragraphs[i], txt)
    d.save(str(ROOT / dst))
    print("OK ->", dst)

# ====================================================================
# CATATAN 26 JUNI — versi sederhana
# ====================================================================
r26 = {
0: "CARA SMARTWATCH KIRIM DATA KE HP LEWAT BLE (PENJELASAN SEDERHANA)",
1: "Jalur Data di Smartwatch (Layanan GATT)",
2: "Smartwatch membuka satu “layanan” (ibarat pintu data, kodenya 0000a100) yang punya dua "
   "jalur. Jalur pertama (0000a101) untuk mengirim data detak jantung ke HP; jalur ini bisa "
   "memberi tahu HP otomatis tiap ada data baru (disebut notify) dan punya saklar langganan "
   "bernama CCCD. Jalur kedua (0000a102) untuk HP membalas “sudah saya terima” (ACK). Kode-kode "
   "itu hanya nomor pengenal unik untuk tiap jalur.",
3: "Cara Menyambung",
4: "Smartwatch terus menyiarkan keberadaannya supaya gampang ditemukan HP. Nama perangkat dikirim "
   "terpisah agar siaran tidak kepenuhan (batasnya 31 byte). HP mencari smartwatch yang punya "
   "layanan tadi lalu menyambung. Setelah tersambung, HP minta jatah paket sebesar mungkin (sampai "
   "512 byte) supaya sekali kirim bisa banyak data, lalu mengecek jalur-jalur yang tersedia.",
5: "Angka 512 byte itu permintaan dari HP. HP minta 512, smartwatch menjawab sebesar yang ia mampu. "
   "Yang dipakai adalah angka terkecil di antara keduanya — jadi bisa kurang dari 512. Tawar-menawar "
   "ini diurus sistem Bluetooth, bukan aplikasi. Kalau gagal (misalnya perangkat tak mendukung), "
   "dipakai ukuran default yang kecil: 23 byte.",
6: "Ukuran tiap potongan data mengikuti jatah paket tadi (jatah dikurangi 4 byte). Jadi jatah paket "
   "hanya menentukan seberapa BANYAK potongannya, bukan lengkap-tidaknya data. Tiap kiriman selalu "
   "diawali penanda START dan diakhiri END; yang berubah cuma jumlah potongan DATA di tengah. Jatah "
   "besar berarti potongan besar, jumlahnya sedikit, transfer cepat. Jatah kecil berarti potongan "
   "kecil, jumlahnya banyak, transfer lambat. Tabel di bawah membandingkan satu kiriman (sekitar 228 "
   "data, ~10.717 byte) pada dua ukuran:",
7: "Pada dua kasus itu, data yang sampai sama persis (228 data). Jadi ukuran paket hanya memengaruhi "
   "kecepatan, bukan kelengkapan. Ringkasan tawar-menawar ini ada di Gambar 1.",
9: "Gambar 1. Tawar-menawar ukuran paket (MTU): HP minta ukuran, smartwatch menjawab sebesar yang ia "
   "bisa, lalu dipakai yang terkecil. Kalau gagal, dipakai ukuran default 23 byte.",
10: "Berlangganan Data (Subscribe)",
11: "Supaya bisa menerima data otomatis, HP berlangganan dulu ke jalur data (menulis kode 0x01 0x00 "
    "ke saklar CCCD). Smartwatch lalu mencatat HP itu sebagai pelanggan. Mulai dari sini, smartwatch "
    "boleh mengirim data.",
12: "Kirim Banyak Data Sekaligus (Dipotong-potong)",
13: "Satu kiriman data (berisi banyak pembacaan dalam bentuk JSON) biasanya lebih besar dari kapasitas "
    "sekali kirim, jadi dipotong jadi beberapa bagian (frame). Byte pertama tiap bagian menandai "
    "jenisnya: START (0x01) = mulai, DATA (0x02) = isi potongan, END (0x03) = selesai. Potongan "
    "dikirim satu per satu — yang berikutnya menunggu yang sebelumnya benar-benar terkirim. HP "
    "menyusun ulang semua potongan sampai ketemu END, lalu mengubahnya jadi daftar pembacaan dan "
    "menyimpannya ke database sekaligus.",
14: "Notifikasi BLE sifatnya “kirim lalu lupakan” — smartwatch tidak otomatis tahu apakah data benar-"
    "benar sampai. BLE punya cara lain (indication) yang ada konfirmasinya, tapi konfirmasinya terjadi "
    "tiap pesan sehingga lambat. Sistem ini sengaja pakai cara cepat (notify), lalu menambah "
    "konfirmasi sendiri satu kali per kiriman (ACK) — lebih hemat daripada konfirmasi tiap potongan. "
    "Penjelasan ACK ada di bagian berikut.",
15: "Balasan “Sudah Diterima” (ACK)",
16: "Setelah semua data tersimpan di HP, HP menuliskan berapa banyak data yang berhasil disimpan ke "
    "jalur ACK. Smartwatch baru menandai data itu “sudah terkirim” (synced = 1) SETELAH menerima "
    "balasan ini. Kalau balasan tak datang dalam 30 detik (misalnya aplikasi HP berhenti atau koneksi "
    "putus), data tetap dianggap belum terkirim dan akan dikirim ulang nanti. Ini menutup celah "
    "“dikira terkirim padahal belum sampai”.",
17: "Mencegah Data Dobel",
18: "Karena ada kirim ulang, satu pembacaan bisa terkirim dua kali. Supaya tidak dobel, database HP "
    "memberi kunci unik pada kolom waktu dan mengabaikan data yang sudah ada (mode INSERT OR IGNORE). "
    "Jadi kirim ulang tidak membuat baris ganda, dan jumlah duplikat tetap nol. Pencocokan data watch "
    "vs HP memakai kolom waktu (dibuat di watch), sehingga tidak bergantung pada kecocokan jam kedua "
    "perangkat.",
23: "Gambar 2. Jalur lengkap komunikasi: daftar jalur di smartwatch, penyambungan, berlangganan, "
    "pengiriman data dipotong-potong, dan balasan ACK kembali ke smartwatch.",
24: "Gambar di atas menunjukkan keseluruhan cara smartwatch (penyedia data) dan HP (peminta data) "
    "berkomunikasi lewat Bluetooth. Bagian atas: daftar jalur yang dibuka smartwatch — satu layanan "
    "utama (0000a100) yang menaungi jalur data (0000a101, tipe notify, ada saklar CCCD) dan jalur "
    "balasan (0000a102, tipe write). Bagian bawah: urutan pesan dari waktu ke waktu, dibagi lima tahap.",
26: "Tahap 1 (Setup): smartwatch menyiapkan jalur-jalurnya lalu menyiarkan keberadaannya. Tahap 2 "
    "(Connect & discover): HP mencari, menyambung, menawar ukuran paket (sampai 512 byte), dan "
    "mengecek jalur yang dibutuhkan. Tahap 3 (Subscribe): HP berlangganan ke jalur data; smartwatch "
    "mencatatnya sebagai pelanggan. Tahap 4 (Notify): smartwatch mengirim data dipotong-potong, satu "
    "per satu (potongan berikut menunggu yang sebelumnya selesai), sampai penanda END membuat HP "
    "menyusun dan menyimpan data. Tahap 5 (ACK): HP menulis jumlah data yang tersimpan ke jalur "
    "balasan, lalu smartwatch menandai data “sudah terkirim”. Garis penuh = aliran data dari "
    "smartwatch ke HP; garis putus-putus = balasan ACK.",
28: "Gambar 3. Cara satu kiriman dipotong jadi frame START / DATA / END (satu frame = satu notifikasi BLE).",
29: "Gambar 3 menjelaskan cara memotong satu kiriman yang ukurannya lebih besar dari kapasitas sekali "
    "kirim. Satu kiriman (daftar JSON berisi bpm, accuracy, time) dipotong jadi rangkaian frame: "
    "diawali START, lalu beberapa DATA (masing-masing membawa satu potongan), ditutup END. Kotak kecil "
    "di bawah memperlihatkan isi satu frame DATA: byte pertama penanda jenis (0x02), sisanya potongan "
    "data sepanjang maksimum MTU−4 byte. Tiap frame dikirim sebagai satu notifikasi dan menunggu yang "
    "sebelumnya selesai, jadi tidak ada data menumpuk atau hilang. Di HP, potongan disusun ulang sampai "
    "END, lalu diubah jadi satu daftar JSON yang utuh.",
}
rewrite("26-06-26-Catatan-Faik-BLE-SmartWatch.docx",
        "26-06-26-Catatan-Faik-BLE-SmartWatch-sederhana.docx", r26)

# ====================================================================
# CATATAN 28 JUNI — versi sederhana
# ====================================================================
r28 = {
0: "CATATAN HASIL UJI — KIRIM DATA DETAK JANTUNG & TEMUAN PENTING",
2: "Cara Mengambil Datanya",
3: "Data diambil langsung dari kedua alat lewat adb (aplikasi versi rilis): file CSV hasil ekspor "
   "disalin dari folder Download. Smartwatch-nya Samsung Galaxy Watch (SM-R860), HP-nya Xiaomi Redmi "
   "Note 10 Pro. File mentahnya: watch-2026-06-28.csv (45.446 baris) dan phone-2026-06-28.csv (42.562 "
   "baris), dari 23 sampai 28 Juni 2026. Dari jeda waktunya, data terbagi jadi empat sesi.",
4: "Membuang Data Saat Jam Tidak Dipakai",
5: "Hampir separuh data (22.231 baris) ternyata tidak berguna: nilainya beku di 106 dengan kontak "
   "jelek (accuracy ≤0). Ini terjadi saat jam dilepas tapi aplikasinya masih merekam. Karena bukan "
   "detak jantung asli, data ini dibuang dari hitungan. Sisanya, yaitu data pengukuran yang benar, "
   "ada 23.215 baris. Patokan membuang: tiap sesi dipotong sampai data berkontak bagus (accuracy=3) "
   "yang terakhir.",
6: "Hasil Pengiriman Data",
7: "Hasil tiap sesi ada di Tabel 1 dan Gambar 1–2. Totalnya, smartwatch merekam 23.215 data dan "
   "20.455 di antaranya sampai ke HP (berhasil 88,11%) tanpa data dobel. Yang penting: data yang "
   "hilang TIDAK merata. Saat Bluetooth tersambung, hampir tidak ada yang hilang (misalnya setelah HP "
   "menyambung di sesi besar sekitar jam 15:00, hampir 100% sampai). Yang hilang justru menumpuk saat "
   "HP sedang TIDAK tersambung — termasuk satu sesi pendek yang HP-nya tak pernah menyambung, jadi "
   "100% datanya tidak tersimpan. Selain itu, semua data yang sampai isinya sama persis dengan catatan "
   "jam (nilai bpm dan akurasi), jadi tidak ada data yang rusak di perjalanan.",
10: "Gambar 1. Kelengkapan data tiap sesi (yang sampai vs yang hilang).",
12: "Gambar 2. Grafik detak jantung dan titik data hilang di sesi besar; data hanya hilang sebelum HP "
    "tersambung.",
13: "Kualitas Data Sensor",
14: "Pada data pengukuran yang benar, 90,8% punya kontak bagus (accuracy 3); sisanya 9,1% tanpa kontak "
    "(-1) dan 0,1% kontak sedang (0) — lihat Tabel 2. Untuk data berkontak bagus (21.087 baris): detak "
    "terendah 60 bpm, tertinggi 123 bpm, dan rata-rata 83,4 bpm.",
16: "TEMUAN PENTING — Data Lama Tidak Dikirim Ulang Saat Putus-Sambung",
17: "Temuan ini melengkapi catatan 26 Juni (bagian ACK dan store-and-forward). Saat dibandingkan, ada "
    "2.879 data yang DITANDAI “sudah terkirim” (synced = 1) oleh jam, tapi sebenarnya TIDAK ADA di HP. "
    "Hanya 5 data terakhir yang memang masih antre (itu wajar). Semua 2.879 data tadi hilang pada saat "
    "HP sedang putus sambungan.",
18: "Artinya: saat ini jam menandai “terkirim” begitu data dikirim, belum benar-benar menunggu balasan "
    "(ACK) dari HP ketika koneksi putus. Akibatnya, data yang menumpuk selama putus tidak otomatis "
    "dikirim ulang saat tersambung lagi. Jadi mekanismenya bagus SELAMA tersambung, tapi belum tahan "
    "saat putus-sambung.",
19: "Saran perbaikan: (1) jangan tandai “terkirim” sebelum balasan ACK benar-benar diterima; (2) "
    "tambahkan proses kirim-ulang otomatis yang mengecek data belum-terkonfirmasi setiap kali "
    "tersambung kembali; (3) bisa pakai tiga status: belum dikirim / sudah dikirim tapi belum dibalas / "
    "sudah dipastikan. Intinya, balasan di tingkat aplikasi (ACK) lebih bisa diandalkan daripada hanya "
    "mengandalkan notifikasi BLE.",
20: "Yang Sudah Diubah di Naskah",
21: "Kedua draf naskah (Indonesia dan Inggris) sudah diperbarui dengan data ini: Tabel 2 menjadi hasil "
    "per sesi, Tabel 3 (akurasi), tambahan ukuran kesamaan nilai (fidelity), subbab baru 3.5 soal "
    "temuan kirim-ulang ini, plus Gambar 6–9. Angka utama berubah: dari delivery 99,86% (satu sesi "
    "lama) menjadi “isi data 100% sama dan hampir tidak ada yang hilang saat tersambung”, dengan angka "
    "gabungan 88,1% yang justru memperlihatkan kelemahan kirim-ulang tadi. Skrip analisisnya: "
    "hr_analysis/make_figs.py.",
}
rewrite("26-06-28-Catatan-Faik-BLE-SmartWatch.docx",
        "26-06-28-Catatan-Faik-BLE-SmartWatch-sederhana.docx", r28)

print("done")
