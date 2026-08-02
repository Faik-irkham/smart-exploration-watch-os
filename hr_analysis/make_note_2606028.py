#!/usr/bin/env python3
"""Buat catatan baru 26-06-28 (lanjutan 26-06-26) berisi hasil evaluasi
empiris terbaru + temuan keterbatasan backfill store-and-forward.
Memakai dokumen 26-06-26 sebagai template agar gaya/font konsisten."""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH as ALIGN
from docx.oxml.ns import qn
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "docs" / "26-06-26-Catatan-Faik-BLE-SmartWatch.docx"
OUT = ROOT / "docs" / "26-06-28-Catatan-Faik-BLE-SmartWatch.docx"
FIG = ROOT / "figures"

d = Document(str(SRC))
body = d.element.body
for child in list(body):                       # kosongkan isi, simpan sectPr
    if child.tag == qn('w:sectPr'):
        continue
    body.remove(child)

def head(text, size=12):
    p = d.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(size)
    return p

def body_p(text):
    return d.add_paragraph(text)

def caption(text):
    p = d.add_paragraph(text); p.alignment = ALIGN.CENTER
    p.runs[0].font.size = Pt(10); p.runs[0].italic = True
    return p

def figure(img, cap, width):
    p = d.add_paragraph(); p.alignment = ALIGN.CENTER
    p.add_run().add_picture(str(FIG / img), width=Inches(width))
    caption(cap)

def table(rows, header_bold=True, total_bold=True):
    t = d.add_table(rows=len(rows), cols=len(rows[0])); t.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            c = t.rows[i].cells[j]; c.text = v
            if (header_bold and i == 0) or (total_bold and row[0] in ("Total", "Total*")):
                for rn in c.paragraphs[0].runs: rn.bold = True
    return t

# ============================== ISI ==============================
head("CATATAN EVALUASI EMPIRIS — PENGIRIMAN DATA HR & TEMUAN BACKFILL", 13)
sub = d.add_paragraph("28 Juni 2026 · lanjutan dari catatan 26 Juni 2026 (jalur komunikasi BLE/GATT)")
sub.runs[0].italic = True; sub.runs[0].font.size = Pt(10)

head("Pengambilan Data Uji")
body_p("Data diambil langsung dari kedua perangkat melalui adb (mode rilis): file CSV hasil ekspor "
       "ditarik dari folder Download. Smartwatch yang digunakan adalah Samsung Galaxy Watch (SM-R860) "
       "dan smartphone Xiaomi Redmi Note 10 Pro (M2101K6G). Berkas mentah: watch-2026-06-28.csv "
       "(45.446 baris) dan phone-2026-06-28.csv (42.562 baris), mencakup rentang 23–28 Juni 2026. "
       "Berdasarkan jeda waktu antar-pembacaan, data terbagi menjadi empat sesi pengujian.")

head("Pemisahan Data Off-Wrist (kebersihan data)")
body_p("Ditemukan bahwa hampir separuh rekaman (22.231 sampel) berada di ekor sesi dengan nilai BPM "
       "beku (konstan 106) dan akurasi ≤0 — pola khas saat smartwatch tidak terpasang di pergelangan "
       "tetapi aplikasi masih merekam. Sampel ini bukan detak jantung nyata, sehingga dikecualikan dari "
       "seluruh metrik. Periode pengukuran valid (sensor terpasang) menyisakan 23.215 sampel. Aturan "
       "pemisahan: sebuah sesi dipotong sampai sampel ber-accuracy=3 (kontak valid) terakhir.")

head("Hasil Pengiriman (Delivery)")
body_p("Hasil per sesi pada periode pengukuran ditunjukkan pada Tabel 1 dan Gambar 1–2. Secara agregat, "
       "smartwatch merekam 23.215 pembacaan dan 20.455 di antaranya tercatat di smartphone (rasio "
       "keberhasilan 88,11%) tanpa duplikat. Yang penting, kehilangan TIDAK merata: saat tautan BLE aktif "
       "kehilangan mendekati nol (mis. setelah ponsel terhubung pada sesi utama sekitar pukul 15:00, "
       "delivery praktis 100%), sedangkan seluruh kehilangan terkonsentrasi pada periode ketika ponsel "
       "sedang TIDAK terhubung — termasuk satu sesi singkat yang ponselnya tidak pernah terhubung sehingga "
       "100% pembacaannya tidak tersimpan. Selain itu, seluruh pembacaan yang berhasil diterima identik "
       "dengan catatan watch (nilai bpm dan akurasi), sehingga fidelity nilai = 100% — tidak ada korupsi "
       "data pada kanal BLE.")
table([
    ("Sesi", "Mulai (WIB)", "Durasi", "Direkam", "Diterima", "Hilang", "Delivery"),
    ("Sesi 1", "23/06 06:38", "36 mnt", "2.177", "2.164", "13", "99,40%"),
    ("Sesi 2", "23/06 08:41", "1 mnt", "53", "0", "53", "0,00%"),
    ("Sesi 3", "23/06 14:15", "345 mnt", "20.702", "18.013", "2.689", "87,01%"),
    ("Sesi 4", "28/06 21:29", "5 mnt", "283", "278", "5", "98,23%"),
    ("Total", "—", "387 mnt", "23.215", "20.455", "2.760", "88,11%"),
])
caption("Tabel 1. Hasil pengiriman per sesi (periode pengukuran).")
figure("fig_hr_completeness.png", "Gambar 1. Kelengkapan data per sesi (diterima vs hilang).", 5.6)
figure("fig_hr_timeline.png", "Gambar 2. Sinyal detak jantung dan kehilangan paket BLE pada sesi utama; "
       "kehilangan hanya terjadi sebelum ponsel terhubung.", 6.3)

head("Kualitas Data Sensor")
body_p("Pada periode pengukuran, 90,8% pembacaan berada pada akurasi tertinggi (nilai 3); selebihnya "
       "9,1% tanpa kontak (-1) dan 0,1% pada nilai 0 (Tabel 2). Statistik detak jantung pada pembacaan "
       "akurasi tertinggi (n=21.087): minimum 60 bpm, maksimum 123 bpm, rata-rata 83,4 bpm, simpangan "
       "baku 9,9 bpm.")
table([
    ("Akurasi sensor", "Jumlah", "Persentase"),
    ("3 (tinggi)", "21.087", "90,8%"),
    ("0 (sedang)", "15", "0,1%"),
    ("-1 (tanpa kontak)", "2.113", "9,1%"),
    ("Total", "23.215", "100%"),
])
caption("Tabel 2. Distribusi nilai akurasi sensor (periode pengukuran).")

head("TEMUAN PENTING — Keterbatasan Backfill Store-and-Forward")
body_p("Temuan ini melengkapi (dan sedikit mengoreksi) bagian “Konfirmasi Penerimaan (ACK)” dan "
       "“store-and-forward” pada catatan 26 Juni 2026. Analisis silang antara penanda status terkirim "
       "(synced) di watch dengan keberadaan record di ponsel menunjukkan: sebanyak 2.879 pembacaan "
       "ditandai TELAH terkirim (synced = 1) oleh watch, tetapi TIDAK ada di basis data ponsel; hanya 5 "
       "pembacaan terakhir yang memang masih berstatus belum terkirim (wajar, menunggu interval). Seluruh "
       "selisih 2.879 record itu terjadi pada periode saat ponsel sedang terputus.")
body_p("Interpretasi: pada implementasi sekarang, penandaan synced tampaknya dilakukan pada tingkat "
       "PENGIRIMAN notifikasi, belum sepenuhnya menunggu ACK tingkat aplikasi ketika tautan tidak aktif. "
       "Akibatnya, backlog yang menumpuk selama putus koneksi tidak selalu dikirim ulang (di-backfill) "
       "setelah tautan pulih, padahal idealnya store-and-forward harus memindai record belum-terkonfirmasi "
       "dan mengirim ulang. Jadi: mekanisme menjaga kelengkapan dengan baik SELAMA tautan aktif, tetapi "
       "belum tahan terhadap skenario putus-sambung.")
body_p("Rekomendasi perbaikan: (1) tunda penandaan synced=1 hingga ACK tingkat aplikasi benar-benar "
       "diterima; (2) tambahkan pemicu backfill yang memindai record berstatus belum-terkonfirmasi setiap "
       "kali koneksi terbentuk kembali; (3) pertimbangkan penanda status tiga-keadaan (belum-dikirim / "
       "terkirim-menunggu-ACK / terkonfirmasi). Temuan ini memperkuat pentingnya ACK tingkat aplikasi "
       "dibanding mengandalkan notifikasi BLE semata.")

head("Tindak Lanjut pada Naskah")
body_p("Kedua draf naskah (Indonesia: Draft_Naskah_HR_BLE_SINTA2.docx; Inggris: "
       "Draft_Manuscript_HR_BLE_EN.docx) telah diperbarui dengan data ini: Tabel 2 menjadi hasil per "
       "sesi, Tabel 3 (akurasi periode pengukuran), penambahan metrik fidelity nilai, subbab baru 3.5 "
       "“Temuan: Keterbatasan Backfill Store-and-Forward”, serta Gambar 6–9 (kelengkapan, timeline, "
       "distribusi BPM, kualitas kontak). Angka headline berubah dari delivery 99,86% (satu sesi lama) "
       "menjadi: fidelity 100% & kehilangan ≈0 saat link aktif, dengan delivery agregat 88,1% yang "
       "mengungkap keterbatasan backfill. Skrip analisis reproducible: hr_analysis/make_figs.py.")

d.save(str(OUT))
print("OK ->", OUT.relative_to(ROOT))
print("paragraphs:", len(d.paragraphs), "tables:", len(d.tables),
      "images:", len(d.element.body.findall('.//'+qn('pic:pic'))))
