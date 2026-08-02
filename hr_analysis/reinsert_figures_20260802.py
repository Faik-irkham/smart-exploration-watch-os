#!/usr/bin/env python3
"""Sisipkan ulang Gambar 1–9 ke kedua naskah.

Berkas media pada kedua .docx sudah menjadi yatim (ada di word/media tetapi
tidak lagi dirujuk dari document.xml), sehingga naksah hanya memuat caption
tanpa gambar. Skrip ini mengisi paragraf kosong tepat di atas tiap caption
dengan berkas figur terkini dari folder figures/.

Idempoten: paragraf yang sudah memuat gambar dilewati.
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH as ALIGN
from docx.shared import Inches

ROOT = Path(__file__).resolve().parents[1]
FIG = ROOT / "figures"

# caption prefix -> (berkas figur, lebar inci)
FIGURES = [
    ("1.", "fig_architecture.png", 6.0),
    ("2.", "fig_framing.png", 5.8),
    ("3.", "fig_sequence.png", 5.6),
    ("4.", "fig_storeforward.png", 5.2),
    ("5.", "fig_watch_arch.png", 6.0),
    ("6.", "fig_hr_completeness.png", 5.9),
    ("7.", "fig_hr_timeline.png", 6.3),
    ("8.", "fig_hr_contact.png", 5.0),
    ("9.", "fig_hr_bpm_dist.png", 4.7),
]


def has_image(p):
    return bool(p._p.findall(
        ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"))


def fill(path, word):
    doc = Document(str(path))
    paras = doc.paragraphs
    done = 0
    for num, img, width in FIGURES:
        target = None
        for i, p in enumerate(paras):
            if p.text.strip().startswith(f"{word} {num}") and i > 0:
                target = paras[i - 1]
                break
        if target is None:
            raise SystemExit(f"{path.name}: caption {word} {num} tidak ditemukan")
        if has_image(target):
            continue
        if target.text.strip():
            raise SystemExit(f"{path.name}: paragraf sebelum {word} {num} tidak kosong")
        target.alignment = ALIGN.CENTER
        target.add_run().add_picture(str(FIG / img), width=Inches(width))
        done += 1
    doc.save(str(path))
    print(f"OK — {path.name}: {done} gambar disisipkan")


fill(ROOT / "Draft_Naskah_HR_BLE_SINTA2.docx", "Gambar")
fill(ROOT / "Draft_Manuscript_HR_BLE_EN_v2.docx", "Figure")
