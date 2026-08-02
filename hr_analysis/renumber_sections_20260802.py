#!/usr/bin/env python3
"""Ganti penomoran subbab desimal (2.1, 3.5, 3.6.1) menjadi huruf (A, B, ...).

Konvensi yang dipakai — mengikuti gaya IEEE yang lazim pada jurnal SINTA 2:

    2. Metode Penelitian          <- bab utama tetap bernomor
       A. Arsitektur Sistem       <- subbab memakai huruf, ulang tiap bab
       B. Protokol Komunikasi ...
    3. Hasil dan Pembahasan
       A. Verifikasi Fungsional ...
       F. Pembahasan
          1) Interpretasi Hasil   <- sub-subbab memakai angka berkurung
          2) Perbandingan ...

Rujukan silang di badan teks ikut disesuaikan (mis. "Bagian 2.5" -> "Bagian 2.E",
"Bagian 3.1 sampai 3.5" -> "Bagian 3.A sampai 3.E").
"""
import re
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]

LETTER = {1: "A", 2: "B", 3: "C", 4: "D", 5: "E", 6: "F",
          7: "G", 8: "H", 9: "I", 10: "J"}


def set_text(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        p.add_run(text)


def convert(path, xrefs):
    doc = Document(str(path))

    n_head = 0
    for p in doc.paragraphs:
        t = p.text.strip()

        # sub-subbab: 3.6.1 Judul -> 1) Judul
        m = re.match(r'^\d+\.\d+\.(\d+)\s+(.*)$', t)
        if m:
            set_text(p, f"{int(m.group(1))}) {m.group(2)}")
            n_head += 1
            continue

        # subbab: 2.1 Judul -> A. Judul
        m = re.match(r'^\d+\.(\d+)\s+(\S.*)$', t)
        if m and len(t) < 80:
            set_text(p, f"{LETTER[int(m.group(1))]}. {m.group(2)}")
            n_head += 1

    n_ref = 0
    for p in doc.paragraphs:
        new = p.text
        for old, rep in xrefs:
            new = new.replace(old, rep)
        if new != p.text:
            set_text(p, new)
            n_ref += 1

    doc.save(str(path))
    print(f"OK — {path.name}: {n_head} judul, {n_ref} paragraf rujukan silang")


convert(ROOT / "Draft_Naskah_HR_BLE_SINTA2.docx", [
    ("Bagian 2.5", "Bagian 2.E"),
    ("Bagian 3.1 sampai 3.5", "Bagian 3.A sampai 3.E"),
])

convert(ROOT / "Draft_Manuscript_HR_BLE_EN_v2.docx", [
    ("Section 2.5", "Section 2.E"),
    ("Sections 3.1 to 3.5", "Sections 3.A to 3.E"),
])
