#!/usr/bin/env python3
"""Catatan 4 Juli 2026 — analisis ulang temuan backfill dengan cakupan seragam.
Melengkapi catatan 3 Juli: analisis silang synced vs keberadaan data di ponsel
dihitung ulang HANYA pada periode pengukuran valid (23.215 data) — cakupan yang
sama dengan tabel hasil pengiriman — sehingga angkanya langsung sebanding:
2.760 hilang = 2.755 keliru ditandai terkirim + 5 pending wajar.

Angka direproduksi dari watch/phone-2026-06-28.csv dengan logika pemotongan
sesi yang sama dengan hr_analysis/make_figs.py. Gaya bahasa: ilmiah namun
terbaca awam; pola tabel: pengantar -> tabel -> penjelasan; semua paragraf.
Memakai catatan 28 Juni sebagai template agar gaya/font konsisten."""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH as ALIGN
from docx.oxml.ns import qn
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "docs" / "26-06-28-Catatan-Faik-BLE-SmartWatch.docx"
OUT = ROOT / "docs" / "26-07-04-Catatan-Faik-BLE-SmartWatch-sederhana.docx"

d = Document(str(SRC))
body = d.element.body
for child in list(body):                       # kosongkan isi, simpan sectPr
    if child.tag == qn('w:sectPr'):
        continue
    body.remove(child)

def head(text, size=12):
    p = d.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(size)
    return p

def body_p(text):
    return d.add_paragraph(text)

def caption(text):
    p = d.add_paragraph(text); p.alignment = ALIGN.CENTER
    p.runs[0].font.size = Pt(10); p.runs[0].italic = True
    return p

def table(rows, header_bold=True, total_bold=True):
    t = d.add_table(rows=len(rows), cols=len(rows[0])); t.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            c = t.rows[i].cells[j]; c.text = v
            if (header_bold and i == 0) or (total_bold and row[0] == "Total"):
                for rn in c.paragraphs[0].runs: rn.bold = True
    return t

# ============================== ISI ==============================
head("ANALISIS ULANG TEMUAN BACKFILL DENGAN CAKUPAN SERAGAM", 13)

head("Kenapa Dihitung Ulang")
body_p("Pada catatan 3 Juli, temuan backfill (data keliru ditandai terkirim) dihitung atas seluruh "
       "rekaman smartwatch (45.446 baris) dan menghasilkan angka 2.879, sedangkan tabel hasil pengiriman "
       "dihitung atas periode pengukuran valid (23.215 baris) dan menghasilkan 2.760 data hilang. "
       "Keduanya benar, tetapi perbedaan cakupan itu membutuhkan penjelasan tersendiri dan berpotensi "
       "membingungkan pembaca maupun penelaah jurnal. Untuk standar publikasi, analisis lebih kuat bila "
       "seluruh metrik dihitung pada cakupan yang sama. Karena itu, analisis silang antara penanda "
       "synced di smartwatch dan keberadaan data di ponsel dihitung ulang hanya pada periode pengukuran "
       "valid — cakupan yang sama dengan tabel hasil pengiriman. Perhitungan dilakukan dari berkas yang "
       "sama (watch-2026-06-28.csv dan phone-2026-06-28.csv) dengan aturan pemotongan sesi yang sama "
       "seperti analisis sebelumnya.")

head("Hasil dengan Cakupan Seragam")
body_p("Hasil hitung ulang untuk setiap sesi dirangkum pada Tabel 1.")
table([
    ("Sesi", "Hilang", "Keliru ditandai terkirim (synced = 1)", "Pending wajar (synced = 0)"),
    ("Sesi 1", "13", "13", "0"),
    ("Sesi 2", "53", "53", "0"),
    ("Sesi 3", "2.689", "2.689", "0"),
    ("Sesi 4", "5", "0", "5"),
    ("Total", "2.760", "2.755", "5"),
])
caption("Tabel 1. Komposisi data hilang per sesi pada periode pengukuran valid.")
body_p("Dari 2.760 data yang hilang, 2.755 di antaranya (99,8%) berstatus synced = 1 — artinya "
       "smartwatch menganggapnya sudah terkirim padahal data tersebut tidak pernah ada di ponsel. Hanya "
       "5 data yang berstatus synced = 0, yaitu lima data terakhir Sesi 4 yang memang masih menunggu "
       "jadwal pengiriman berikutnya (pending wajar). Dengan demikian angka hilang pada tabel hasil "
       "pengiriman terurai persis: 2.755 + 5 = 2.760, tanpa selisih. Sesi 2 menjadi bukti paling jelas: "
       "seluruh 53 datanya berstatus terkirim padahal ponsel pada sesi itu tidak pernah terhubung sama "
       "sekali — penandaan jelas dilakukan sebelum ada kepastian data diterima.")

head("Makna Temuan")
body_p("Dengan cakupan seragam, temuan menjadi lebih tajam: praktis seluruh kehilangan data pada "
       "periode pengukuran — kecuali 5 data pending yang wajar — disebabkan oleh satu mekanisme yang "
       "sama, yaitu penandaan terkirim yang terlalu dini ketika koneksi sedang terputus, sehingga data "
       "tersebut tidak pernah dikirim ulang setelah koneksi pulih. Kesimpulan pada catatan sebelumnya "
       "tetap berlaku (mekanisme bekerja baik selama tersambung, belum tahan putus-sambung), dan saran "
       "perbaikannya tidak berubah: tunda penandaan terkirim sampai balasan ACK benar-benar diterima, "
       "tambahkan proses kirim ulang otomatis saat koneksi tersambung kembali, dan gunakan tiga status "
       "(belum dikirim / sudah dikirim tetapi belum dibalas / sudah dipastikan diterima).")

head("Ringkasan")
body_p("Analisis silang backfill dihitung ulang pada cakupan yang sama dengan tabel hasil pengiriman "
       "(periode pengukuran valid, 23.215 data). Hasilnya: 2.760 data hilang terurai persis menjadi "
       "2.755 data yang keliru ditandai terkirim (99,8%) dan 5 data pending wajar, dengan Sesi 2 (53 "
       "data, ponsel tidak pernah terhubung) sebagai bukti penandaan prematur yang paling jelas. Angka "
       "2.755 kini dipakai sebagai angka utama temuan pada kedua naskah (Indonesia dan Inggris), "
       "menggantikan 2.879 yang berpindah menjadi catatan sekunder (cakupan seluruh rekaman); catatan "
       "3 Juli juga telah diselaraskan. Temuan, interpretasi, dan saran perbaikan tidak berubah — "
       "hanya menjadi lebih mudah diperiksa karena semua angka kini berada pada cakupan yang sama.")

d.save(str(OUT))
print("OK ->", OUT.relative_to(ROOT))
print("paragraphs:", len(d.paragraphs), "tables:", len(d.tables))
