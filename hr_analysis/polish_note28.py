#!/usr/bin/env python3
"""Rapikan bahasa & tanda baca catatan 28-sederhana untuk laporan ke pembimbing:
sederhana tapi baku, istilah konsisten, tanpa kapital-teriak."""
from docx import Document
from pathlib import Path

F = Path(__file__).resolve().parents[1] / "docs" / "26-06-28-Catatan-Faik-BLE-SmartWatch-sederhana.docx"
d = Document(str(F))

def set_text(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        p.add_run(text)

repl = {
0: "CATATAN HASIL UJI — PENGIRIMAN DATA DETAK JANTUNG DAN TEMUAN PENTING",
2: "Pengambilan Data",
3: "Data diambil langsung dari kedua perangkat melalui koneksi adb (aplikasi versi rilis): berkas "
   "CSV hasil ekspor disalin dari folder Download. Smartwatch yang digunakan adalah Samsung Galaxy "
   "Watch (SM-R860), dan ponselnya Xiaomi Redmi Note 10 Pro. Berkas mentahnya bernama "
   "watch-2026-06-28.csv (45.446 baris) dan phone-2026-06-28.csv (42.562 baris), mencakup tanggal 23 "
   "sampai 28 Juni 2026. Berdasarkan jeda waktunya, data terbagi menjadi empat sesi.",
4: "Membuang Data Saat Smartwatch Tidak Dipakai",
5: "Hampir separuh data (22.231 baris) ternyata tidak valid: nilainya beku di angka 106 dengan kontak "
   "buruk (accuracy ≤ 0). Kondisi ini terjadi saat smartwatch dilepas, tetapi aplikasinya masih "
   "merekam. Karena bukan detak jantung yang sebenarnya, data ini dikeluarkan dari perhitungan. "
   "Sisanya, yaitu data pengukuran yang sah, berjumlah 23.215 baris. Patokan pemotongannya: setiap "
   "sesi dipotong sampai data dengan kontak baik (accuracy = 3) yang terakhir.",
6: "Hasil Pengiriman Data",
7: "Hasil tiap sesi disajikan pada Tabel 1 serta Gambar 1 dan 2. Secara keseluruhan, smartwatch "
   "merekam 23.215 data dan 20.455 di antaranya sampai ke ponsel (tingkat keberhasilan 88,11%) tanpa "
   "data ganda. Hal penting yang perlu dicatat adalah data yang hilang tidak merata. Ketika Bluetooth "
   "tersambung, hampir tidak ada data yang hilang — misalnya setelah ponsel tersambung pada sesi "
   "besar sekitar pukul 15.00, hampir 100% data sampai. Sebaliknya, kehilangan justru menumpuk saat "
   "ponsel sedang tidak tersambung, termasuk satu sesi pendek yang ponselnya tidak pernah tersambung "
   "sehingga seluruh datanya tidak tersimpan. Selain itu, semua data yang sampai memiliki isi yang "
   "sama persis dengan catatan smartwatch (nilai bpm dan akurasi), sehingga tidak ada data yang rusak "
   "selama pengiriman.",
10: "Gambar 1. Kelengkapan data tiap sesi (data yang sampai dibanding yang hilang).",
12: "Gambar 2. Grafik detak jantung dan titik data yang hilang pada sesi besar; data hanya hilang "
    "sebelum ponsel tersambung.",
13: "Kualitas Data Sensor",
14: "Pada data pengukuran yang sah, 90,8% memiliki kontak baik (accuracy 3); sisanya 9,1% tanpa kontak "
    "(-1) dan 0,1% kontak sedang (0) — lihat Tabel 2. Untuk data dengan kontak baik (21.087 baris): "
    "detak terendah 60 bpm, tertinggi 123 bpm, dan rata-rata 83,4 bpm.",
16: "TEMUAN PENTING — Data Tidak Dikirim Ulang Saat Koneksi Putus-Sambung",
17: "Temuan ini melengkapi catatan 26 Juni (bagian ACK dan store-and-forward). Setelah dibandingkan, "
    "terdapat 2.879 data yang ditandai “sudah terkirim” (synced = 1) oleh smartwatch, tetapi "
    "sebenarnya tidak ada di ponsel. Hanya 5 data terakhir yang memang masih dalam antrean, dan itu "
    "wajar. Seluruh 2.879 data tersebut hilang ketika ponsel sedang terputus.",
18: "Artinya, untuk saat ini smartwatch menandai data “terkirim” begitu data dikirim, belum benar-benar "
    "menunggu balasan (ACK) dari ponsel ketika koneksi terputus. Akibatnya, data yang menumpuk selama "
    "koneksi terputus tidak otomatis dikirim ulang saat tersambung kembali. Dengan kata lain, "
    "mekanisme ini bekerja baik selama koneksi tersambung, tetapi belum tahan terhadap kondisi "
    "putus-sambung.",
19: "Saran perbaikan: (1) jangan menandai data “terkirim” sebelum balasan ACK benar-benar diterima; "
    "(2) tambahkan proses kirim ulang otomatis yang memeriksa data yang belum terkonfirmasi setiap "
    "kali koneksi tersambung kembali; (3) gunakan tiga status, yaitu belum dikirim, sudah dikirim "
    "tetapi belum dibalas, dan sudah dipastikan. Intinya, konfirmasi pada tingkat aplikasi (ACK) lebih "
    "dapat diandalkan daripada hanya mengandalkan notifikasi BLE.",
20: "Perubahan pada Naskah",
21: "Kedua draf naskah (versi Indonesia dan Inggris) telah diperbarui dengan data ini: Tabel 2 menjadi "
    "hasil per sesi, Tabel 3 berisi distribusi akurasi, ditambahkan ukuran kesamaan nilai (fidelity), "
    "serta subbab baru 3.5 mengenai temuan kirim ulang ini, dilengkapi Gambar 6–9. Angka utama "
    "berubah: dari delivery 99,86% (satu sesi lama) menjadi isi data 100% sama dan hampir tidak ada "
    "yang hilang saat koneksi tersambung, dengan angka gabungan 88,1% yang justru memperlihatkan "
    "kelemahan kirim ulang tersebut. Skrip analisisnya: hr_analysis/make_figs.py.",
}
for i, t in repl.items():
    set_text(d.paragraphs[i], t)
d.save(str(F))
print("OK — bahasa & tanda baca dirapikan:", F.name)
