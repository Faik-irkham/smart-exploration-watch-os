#!/usr/bin/env python3
"""Susun ulang §3.6 menjadi Pembahasan berstruktur baku.

Tidak ada hasil baru yang diperkenalkan; seluruh angka yang dirujuk sudah
dilaporkan pada §3.1–§3.5. Struktur baru:

    3.6   Pembahasan
    3.6.1 Interpretasi Hasil          — pemaknaan mendalam atas temuan
    3.6.2 Perbandingan dengan Penelitian Terdahulu — Tabel 4 + narasi
    3.6.3 Implikasi Penelitian        — praktis, metodologis, dan perancangan
    3.6.4 Keterbatasan                — batas keberlakuan hasil

Tabel 4 disusun dari studi yang sudah disitasi pada Bagian 1; kolom
"verifikasi kelengkapan per record" adalah pembeda utama penelitian ini.
"""
import copy
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[1]


def set_text(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        p.add_run(text)


def find(doc, prefix):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    raise SystemExit(f"PARAGRAF TIDAK DITEMUKAN: {prefix!r}")


def insert_before(ref, *texts):
    """Sisipkan paragraf baru sebelum `ref`, mewarisi format `ref`."""
    out = []
    for t in texts:
        el = copy.deepcopy(ref._p)
        ref._p.addprevious(el)
        new = Paragraph(el, ref._parent)
        set_text(new, t)
        out.append(new)
    return out


def insert_after(ref, *texts):
    last = ref
    for t in texts:
        el = copy.deepcopy(ref._p)
        last._p.addnext(el)
        new = Paragraph(el, ref._parent)
        set_text(new, t)
        last = new
    return last


def add_table_after(doc, ref, caption, header, rows):
    """Sisipkan caption lalu tabel tepat setelah `ref`; kembalikan paragraf tabel."""
    p_cap = insert_after(ref, caption)
    tbl = doc.add_table(rows=1, cols=len(header))
    tbl.style = doc.tables[0].style
    for i, h in enumerate(header):
        tbl.rows[0].cells[i].text = h
        for p in tbl.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = tbl.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v
    p_cap._p.addnext(tbl._tbl)

    # paragraf kosong sesudah tabel, dipakai sebagai jangkar narasi
    anchor_el = copy.deepcopy(ref._p)
    tbl._tbl.addnext(anchor_el)
    anchor = Paragraph(anchor_el, ref._parent)
    set_text(anchor, "")
    return anchor


# ==========================================================================
# NASKAH INDONESIA
# ==========================================================================
doc = Document(str(ROOT / "Draft_Naskah_HR_BLE_SINTA2.docx"))

set_text(find(doc, "3.6 Pembahasan"), "3.6 Pembahasan")

# ---- 3.6.1 Interpretasi Hasil -------------------------------------------
p_interp = find(doc, "Secara keseluruhan, hasil pengujian menunjukkan")
insert_before(p_interp, "3.6.1 Interpretasi Hasil")

set_text(p_interp,
    "Hasil pada Bagian 3.1 sampai 3.5 dapat dibaca sebagai tiga lapis temuan yang saling "
    "menopang. Lapis pertama menyangkut kanal: fidelitas nilai 100% pada seluruh record yang "
    "diterima, ketiadaan duplikat pada kedua versi, dan negosiasi MTU yang konsisten menghasilkan "
    "512 byte menunjukkan bahwa tautan BLE beserta skema framing ber-opcode dan kendali aliran "
    "tidak merusak maupun menggandakan data. Dengan kata lain, sepanjang sebuah record sampai di "
    "penerima, ia sampai dalam keadaan utuh. Lapis kedua menyangkut lapisan aplikasi: meskipun "
    "kanal bersih, versi awal tetap kehilangan 2.707 record, dan 99,8% di antaranya ternyata sudah "
    "bertanda terkirim. Kehilangan karena itu tidak dapat dijelaskan oleh kualitas radio, "
    "melainkan oleh satu keputusan perangkat lunak mengenai kapan sebuah record boleh dinyatakan "
    "selesai. Lapis ketiga adalah konsekuensinya: pada sistem pemantauan berkelanjutan yang "
    "memakai penyangga lokal, keandalan ujung-ke-ujung lebih ditentukan oleh semantik penanda "
    "status kirim daripada oleh mutu tautan nirkabel.")

insert_after(p_interp,
    "Pembacaan tersebut menjelaskan mengapa keberadaan penyangga store-and-forward saja tidak "
    "memadai. Penyangga hanya berfungsi bila record yang belum terkonfirmasi tetap berada di dalam "
    "antrean; begitu penanda dimutakhirkan oleh peristiwa lokal di sisi pengirim, penyangga "
    "kehilangan daftar pekerjaannya dan mekanisme pengiriman ulang tidak pernah aktif. Inilah "
    "sebabnya kegagalan versi awal justru paling parah pada sesi terpanjang, yaitu sesi yang "
    "peluang terputusnya paling besar, sedangkan sesi pendek yang tersambung sejak awal nyaris "
    "tidak terpengaruh. Pola tersebut konsisten dengan mekanisme yang diusulkan, meskipun perlu "
    "ditegaskan bahwa konsistensi pola bukanlah pembuktian kausal selama eksperimen ablasi "
    "terkendali belum dilakukan.",
    "Angka kinerja transfer memberi konteks yang melengkapi pembacaan di atas. Karena satu batch "
    "hanya menempati sekitar 0,07% dari interval pengirimannya, biaya kanal untuk menambahkan "
    "mekanisme keandalan pada rancangan ini nyaris tidak berarti. Ruang yang lapang itu bermakna "
    "praktis: penambahan identitas batch, validasi jumlah record yang dikonfirmasi, atau bahkan "
    "pengiriman ulang selektif dapat dilakukan tanpa menekan anggaran daya yang justru menjadi "
    "alasan pemilihan BLE sejak awal. Dengan demikian, pertukaran antara keandalan dan efisiensi "
    "energi pada pemantauan berbasis tren tidak setajam pada aliran waktu-nyata, tempat setiap "
    "paket menuntut balasan.")

# ---- 3.6.2 Perbandingan --------------------------------------------------
p_cmp = find(doc, "Dibandingkan dengan pekerjaan terdahulu")
insert_before(p_cmp, "3.6.2 Perbandingan dengan Penelitian Terdahulu")

set_text(p_cmp,
    "Posisi penelitian ini terhadap studi sejenis dirangkum pada Tabel 4. Perbandingan disusun "
    "menurut empat atribut yang menentukan apakah sebuah rancangan dapat menjamin kelengkapan "
    "data: tataran kajian, ada atau tidaknya penyanggaan di sisi sumber, ada atau tidaknya "
    "konfirmasi tingkat aplikasi, dan ada atau tidaknya verifikasi kelengkapan yang diperiksa per "
    "record.")

anchor = add_table_after(doc, p_cmp,
    "Tabel 4. Posisi penelitian ini terhadap studi sejenis.",
    ["Studi", "Fokus dan tataran kajian", "Penyanggaan dan konfirmasi di sisi sumber",
     "Verifikasi kelengkapan per record", "Platform evaluasi"],
    [
        ["Balas dkk. [8]", "Perilaku protokol BLE pada aplikasi pemantauan detak jantung",
         "Tidak dibahas", "Tidak — metrik tataran tautan", "Perangkat BLE, pengukuran protokol"],
        ["Xu dkk. [11]", "Pemodelan kinerja BLE dan pengaruh parameter lapisan tautan",
         "Tidak dibahas", "Tidak — metrik throughput", "Model analitis"],
        ["Xiong dan Jiang [9]", "Perutean delay-tolerant berbasis store-carry-forward",
         "Ya, penyanggaan dan konfirmasi antar-simpul", "Tidak — metrik perutean",
         "Simulasi jaringan"],
        ["Majumdar dkk. [10]", "Survei keandalan, ketangguhan, dan efisiensi energi pada WBAN",
         "Ditinjau secara umum", "Tidak — kajian pustaka", "Studi pustaka"],
        ["Schweizer dan Gilgen-Ammann [12]", "Validasi akurasi sensor detak jantung pergelangan tangan",
         "Tidak relevan", "Tidak — fokus akurasi sensor", "Uji validasi lintas kondisi aktivitas"],
        ["Penelitian ini", "Pengiriman data detak jantung smartwatch–smartphone pada tataran aplikasi",
         "Ya: SQLite di smartwatch, ACK tingkat aplikasi, penerima idempoten",
         "Ya: pencocokan timestamp atas 25.191 pembacaan",
         "Sepasang perangkat konsumer (Wear OS dan Android), lima sesi"],
    ])

set_text(anchor,
    "Tabel tersebut memperlihatkan pembagian kerja yang cukup jelas di antara studi terdahulu. "
    "Kajian pada tataran protokol menjelaskan bagaimana parameter koneksi dan lapisan tautan "
    "memengaruhi kualitas serta throughput pengiriman BLE [8], [11], tetapi berhenti pada perilaku "
    "tautan dan tidak menelusuri apakah setiap pembacaan yang direkam benar-benar tersimpan di "
    "penerima. Kajian pada tataran jaringan menyediakan kerangka keandalan yang lebih luas: survei "
    "WBAN merangkum pendekatan keandalan dan efisiensi energi pada arsitektur jaringan [10], "
    "sedangkan protokol perutean delay-tolerant menerapkan gagasan penyanggaan yang secara "
    "konseptual paling dekat dengan penelitian ini [9]. Namun gagasan tersebut bekerja pada "
    "penyanggaan antar-simpul dalam simulasi jaringan, bukan pada satu tautan wearable–gateway "
    "nyata yang dibatasi MTU dan pembatasan eksekusi latar belakang sistem operasi seluler [3]. "
    "Pada tataran perangkat, studi validasi sensor menegaskan pentingnya melaporkan kualitas "
    "kontak bersama data [12] — sebuah aspek yang penelitian ini tangani dengan mengirim dan "
    "menyimpan status akurasi apa adanya, tetapi bukan aspek keandalan pengiriman.")

insert_after(anchor,
    "Perbedaan yang paling menentukan terletak pada kolom terakhir Tabel 4. Tidak satu pun studi "
    "pembanding memverifikasi kelengkapan pada tingkat record individual dengan membandingkan basis "
    "data pengirim dan penerima. Verifikasi semacam itulah yang memungkinkan penelitian ini "
    "menemukan kegagalan backfill pada versi awal — kegagalan yang tidak akan terdeteksi oleh "
    "metrik tataran tautan, tidak muncul pada simulasi perutean, dan tidak tampak pada pengujian "
    "fungsional biasa. Karena itu kontribusi penelitian ini bersifat melengkapi, bukan menggantikan "
    "studi-studi tersebut: ia menurunkan gagasan penyanggaan yang sudah mapan ke tataran aplikasi "
    "pada perangkat konsumer nyata, sekaligus menyediakan bukti kuantitatif atas satu mode "
    "kegagalan yang selama ini luput dari kedua tataran kajian di atasnya.")

# ---- 3.6.3 Implikasi -----------------------------------------------------
p_lim = find(doc, "Penelitian ini memiliki beberapa keterbatasan")
insert_before(p_lim,
    "3.6.3 Implikasi Penelitian",
    "Bagi perancang sistem pemantauan berbasis wearable, implikasi paling langsung adalah aturan "
    "perancangan yang sederhana namun mengikat: penanda status kirim hanya boleh dimutakhirkan "
    "oleh peristiwa yang berasal dari penerima. Aturan ini tidak khas BLE. Setiap kanal yang "
    "bersifat best-effort — dan setiap rancangan yang menaruh penyangga di sisi sumber — "
    "menghadapi risiko yang sama, sehingga pemeriksaan serupa layak dilakukan pada sistem yang "
    "memakai kanal lain dengan pola store-and-forward. Implikasi turunannya, konfirmasi tingkat "
    "aplikasi sebaiknya diperlakukan sebagai bagian dari kontrak penyimpanan data, bukan sekadar "
    "pelengkap protokol.",
    "Bagi metodologi pengujian, hasil penelitian ini menunjukkan bahwa pengujian fungsional dan "
    "indikator pada antarmuka tidak memadai untuk menilai kelengkapan data. Pada versi awal, "
    "aplikasi berjalan normal dan melaporkan seluruh data terkirim, sementara hampir tiga ribu "
    "record tidak pernah sampai. Verifikasi kelengkapan berbasis pencocokan timestamp antara basis "
    "data kedua sisi karena itu perlu menjadi prosedur baku, bukan analisis tambahan. Biaya "
    "penerapannya rendah — cukup ekspor kedua basis data lalu cocokkan berdasarkan kunci waktu — "
    "dan karena timestamp dibuat di sisi sumber lalu dikirim apa adanya, prosedur ini tidak "
    "bergantung pada sinkronisasi jam antar-perangkat.",
    "Bagi perancangan protokol pada aplikasi IoMT, temuan mengenai penggunaan kanal memberi arah "
    "yang cukup praktis. Anggaran kanal yang terpakai sangat kecil menunjukkan bahwa penguatan "
    "keandalan pada skema batch periodik hampir tidak berbiaya, sehingga pilihan rancangan tidak "
    "perlu diposisikan sebagai pertukaran antara kelengkapan data dan hemat daya. Perlu ditegaskan "
    "bahwa implikasi ini berlaku untuk kelas sistem yang diuji, yaitu pemantauan berbasis tren "
    "pada satu tautan wearable–gateway dengan toleransi tunda beberapa menit; sistem yang menuntut "
    "peringatan waktu-nyata memiliki kendala berbeda dan tidak tercakup oleh hasil ini.")

# ---- 3.6.4 Keterbatasan --------------------------------------------------
insert_before(p_lim, "3.6.4 Keterbatasan")

doc.save(str(ROOT / "Draft_Naskah_HR_BLE_SINTA2.docx"))
print("OK — Draft_Naskah_HR_BLE_SINTA2.docx")


# ==========================================================================
# NASKAH INGGRIS
# ==========================================================================
doc = Document(str(ROOT / "Draft_Manuscript_HR_BLE_EN_v2.docx"))

set_text(find(doc, "3.6 Discussion"), "3.6 Discussion")

p_interp = find(doc, "Taken together, the results show that the protocol design")
insert_before(p_interp, "3.6.1 Interpretation of the Results")

set_text(p_interp,
    "The results in Sections 3.1 to 3.5 can be read as three mutually supporting layers of "
    "finding. The first layer concerns the channel: 100% value fidelity across every received "
    "record, the absence of duplicates under either version, and MTU negotiation that consistently "
    "settled at 512 bytes together show that the BLE link, the opcode framing, and the flow "
    "control neither corrupt nor duplicate data. Put differently, whenever a record reached the "
    "receiver, it arrived intact. The second layer concerns the application: despite that clean "
    "channel, the initial version still lost 2,707 records, 99.8% of which had already been marked "
    "as sent. The loss therefore cannot be explained by radio quality, but by a single software "
    "decision about when a record may be declared complete. The third layer is the consequence: in "
    "continuous monitoring systems that rely on a local buffer, end-to-end reliability is governed "
    "more by the semantics of the send-status flag than by the quality of the wireless link.")

insert_after(p_interp,
    "This reading explains why the presence of a store-and-forward buffer is not sufficient on its "
    "own. A buffer only works while unconfirmed records remain in the queue; once the flag is "
    "advanced by a local event at the sender, the buffer loses its worklist and the retransmission "
    "path never activates. That is why the failure of the initial version was most severe in the "
    "longest session — the session with the greatest opportunity to lose the link — while short "
    "sessions that were connected from the outset were barely affected. The pattern is consistent "
    "with the proposed mechanism, though it bears repeating that a consistent pattern is not causal "
    "proof so long as a controlled ablation experiment has not been performed.",
    "The transfer-performance figures supply context that completes this picture. Because a single "
    "batch occupies only about 0.07% of its sending interval, the channel cost of adding "
    "reliability machinery to this design is close to negligible. That headroom has practical "
    "meaning: batch identity, validation of the acknowledged record count, or even selective "
    "retransmission can be added without eroding the power budget that motivated the choice of BLE "
    "in the first place. The trade-off between reliability and energy efficiency is therefore far "
    "less acute in trend-based monitoring than in real-time streaming, where every packet demands "
    "a reply.")

p_cmp = find(doc, "Relative to prior work, this study occupies a distinct position")
insert_before(p_cmp, "3.6.2 Comparison with Related Studies")

set_text(p_cmp,
    "The position of this study relative to comparable work is summarised in Table 4. The "
    "comparison is organised around the attributes that determine whether a design can guarantee "
    "data completeness: the level at which the study operates, whether buffering exists at the "
    "source, whether an application-level acknowledgement is used, and whether completeness is "
    "verified at the level of individual records.")

anchor = add_table_after(doc, p_cmp,
    "Table 4. Position of this study relative to comparable work.",
    ["Study", "Focus and level of analysis", "Buffering and confirmation at the source",
     "Per-record completeness verification", "Evaluation platform"],
    [
        ["Wan et al. [14]", "Wearable IoT health monitoring streamed directly to the cloud",
         "Not implemented", "No", "Wearable prototype"],
        ["Manjunath et al. [15]", "Battery-free skin-mountable multi-sensor patch, continuous BLE streaming",
         "Not implemented", "No", "Hardware prototype"],
        ["Balas et al. [20]", "BLE protocol behaviour in a heart-rate monitoring application",
         "Not addressed", "No — link-level metrics", "BLE devices, protocol measurement"],
        ["Xu et al. [8]", "Analytical modelling of BLE performance",
         "Not addressed", "No — throughput metrics", "Analytical model"],
        ["Xiong and Jiang [21]", "Delay-tolerant routing based on store-carry-forward",
         "Yes, buffering and confirmation between nodes", "No — routing metrics",
         "Network simulation"],
        ["Domingues et al. [24]", "MQTT retransmission in the wake of network disconnection",
         "Yes, retransmission after reconnection", "No — payload-sizing metrics",
         "MQTT network testbed"],
        ["Franco et al. [25]", "BLE connection reliability for mobile health applications",
         "Not addressed", "No — connection-level metrics", "Mobile devices"],
        ["Majumdar et al. [22]", "Survey of reliability, robustness, and energy efficiency in WBANs",
         "Reviewed generically", "No — literature study", "Literature study"],
        ["This work", "Application-level heart-rate delivery from smartwatch to smartphone over BLE",
         "Yes: SQLite on the smartwatch, application-level ACK, idempotent receiver",
         "Yes: timestamp matching over 25,191 readings",
         "Consumer device pair (Wear OS and Android), five sessions"],
    ])

set_text(anchor,
    "The table reveals a fairly clear division of labour among prior studies. Direct-transmission "
    "wearable architectures optimise for immediacy: readings are streamed to a cloud service or to "
    "a receiver as they are produced [14], [15], which makes them well suited to live display but "
    "leaves no provision for recovering a backlog once the link or the power source is interrupted. "
    "Protocol-level studies explain how connection parameters and link-layer behaviour shape BLE "
    "delivery quality and throughput [8], [20], and connection-reliability evaluations characterise "
    "how often and how long mobile-health links fail [25]; all of them, however, stop at link "
    "behaviour and do not trace whether each recorded reading was ultimately stored at the "
    "receiver. Network-level work supplies the broader reliability framing: the WBAN survey "
    "summarises reliability and energy-efficiency approaches at the architectural level [22], while "
    "delay-tolerant routing applies the buffering idea that is conceptually closest to this study "
    "[21], and MQTT retransmission work addresses recovery after disconnection in a different "
    "protocol stack [24]. Those mechanisms nonetheless operate between nodes, in simulation or on "
    "a broker-mediated testbed, rather than on a single wearable–gateway link constrained by MTU "
    "and by mobile-OS background-execution limits [10], [26].")

insert_after(anchor,
    "The decisive difference lies in the last column of Table 4. None of the comparison studies "
    "verifies completeness at the level of individual records by comparing the sender and receiver "
    "databases. It was precisely that verification which exposed the backfill failure of the "
    "initial version — a failure that link-level metrics would not register, that routing "
    "simulation would not reproduce, and that ordinary functional testing would not reveal. The "
    "contribution of this work is therefore complementary rather than competing: it brings a "
    "well-established buffering idea down to the application level on real consumer devices, and "
    "supplies quantitative evidence of a failure mode that has fallen between the two levels of "
    "analysis above it.")

p_lim = find(doc, "This study has several limitations")
insert_before(p_lim,
    "3.6.3 Implications",
    "For designers of wearable monitoring systems, the most immediate implication is a simple but "
    "binding design rule: a send-status flag may be advanced only by an event originating at the "
    "receiver. The rule is not specific to BLE. Any best-effort channel — and any design that "
    "places a buffer at the source — carries the same risk, so a comparable audit is warranted for "
    "systems that combine store-and-forward with other transports. A corollary is that "
    "application-level acknowledgement should be treated as part of the data-persistence contract "
    "rather than as an optional protocol embellishment.",
    "For testing methodology, these results show that functional testing and user-interface "
    "indicators are inadequate for assessing data completeness. Under the initial version the "
    "application ran normally and reported all data as sent, while nearly three thousand records "
    "never arrived. Completeness verification by timestamp matching between the two databases "
    "should therefore become standard procedure rather than a supplementary analysis. The cost of "
    "adopting it is low — export both databases and match on the time key — and because the "
    "timestamp is created at the source and transmitted unchanged, the procedure does not depend "
    "on clock synchronisation between devices.",
    "For protocol design in IoMT applications, the channel-occupancy finding offers a practical "
    "direction. The very small share of channel time consumed indicates that strengthening "
    "reliability in a periodic batch scheme is almost free, so design choices need not be framed as "
    "a trade-off between data completeness and power thrift. It should be stressed that this "
    "implication holds for the class of system tested here — trend-based monitoring over a single "
    "wearable–gateway link that tolerates a delay of several minutes; systems requiring real-time "
    "alerting face different constraints and are not covered by these results.")

insert_before(p_lim, "3.6.4 Limitations")

doc.save(str(ROOT / "Draft_Manuscript_HR_BLE_EN_v2.docx"))
print("OK — Draft_Manuscript_HR_BLE_EN_v2.docx")
