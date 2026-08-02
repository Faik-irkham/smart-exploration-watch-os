#!/usr/bin/env python3
"""Lengkapi perangkat kepustakaan kedua manuskrip (syarat SINTA 2):
1. Isi placeholder daftar pustaka [7]-[10] dengan referensi internasional
   >=2023 yang TERVERIFIKASI via Crossref (judul/penulis/venue/DOI), dan
   tambahkan [11]-[12].
2. Ganti dua placeholder sitasi di Pendahuluan dengan sitasi nyata.
3. Sisipkan dua paragraf Penelitian Terkait (related work + celah riset)
   sebelum paragraf kontribusi.
4. Sitasi pendukung: [11] pada kalimat MTU (Pendahuluan) dan [12] pada
   keterbatasan validasi klinis sensor (Pembahasan).

Referensi (semua diverifikasi 5 Jul 2026 melalui api.crossref.org):
[7]  He dkk., Digital Communications and Networks 12(5):717-742, 2026,
     doi:10.1016/j.dcan.2024.11.013
[8]  Balas dkk., Procedia Computer Science 225:63-69, 2023,
     doi:10.1016/j.procs.2023.09.092
[9]  Xiong & Jiang, Electronics 12(21):4528, 2023,
     doi:10.3390/electronics12214528
[10] Majumdar dkk., Sensors 24(20):6531, 2024, doi:10.3390/s24206531
[11] Xu dkk., IEEE Communications Letters 28(3):732-736, 2024,
     doi:10.1109/LCOMM.2024.3352545
[12] Schweizer & Gilgen-Ammann, JMIR Cardio 9:e67110, 2025, doi:10.2196/67110
"""
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]

def set_text(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        p.add_run(text)

def find(doc, prefix):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    raise SystemExit(f"PARAGRAF TIDAK DITEMUKAN: {prefix!r}")

def insert_before(anchor, text):
    new_p = OxmlElement("w:p")
    anchor._p.addprevious(new_p)
    np = Paragraph(new_p, anchor._parent)
    np.add_run(text)
    return np

def insert_after(anchor, text):
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    np = Paragraph(new_p, anchor._parent)
    np.add_run(text)
    return np

def replace_in_paragraphs(doc, replacements):
    done = {old: False for old, _ in replacements}
    for p in doc.paragraphs:
        t = p.text
        hit = False
        for old, new in replacements:
            if old in t:
                t = t.replace(old, new); done[old] = True; hit = True
        if hit:
            set_text(p, t)
    missing = [k for k, v in done.items() if not v]
    if missing:
        raise SystemExit(f"TEKS TIDAK DITEMUKAN: {missing}")

# ============================ NASKAH ID ============================
doc = Document(str(ROOT / "Draft_Naskah_HR_BLE_SINTA2.docx"))

replace_in_paragraphs(doc, [
    ("data fisiologis [[lengkapi sitasi: survei IoMT]]", "data fisiologis [7]"),
    ("aplikasi penerima dihentikan sistem [[lengkapi sitasi: studi terkait wearable BLE/mHealth]]",
     "aplikasi penerima dihentikan sistem [8], [10]"),
    ("dibatasi oleh Maximum Transmission Unit (MTU)",
     "dibatasi oleh Maximum Transmission Unit (MTU) [11]"),
    ("Kelima, sensor pada smartwatch konsumer belum tervalidasi secara klinis",
     "Kelima, sensor pada smartwatch konsumer belum tervalidasi secara klinis dan akurasinya "
     "diketahui bervariasi antar-kondisi aktivitas [12]"),
])

kontrib = find(doc, "Kontribusi penelitian ini adalah")
insert_before(kontrib,
    "Sejumlah studi terkini telah mengkaji aspek-aspek yang bersinggungan dengan penelitian ini. He "
    "dkk. [7] memetakan arsitektur, aplikasi, dan tantangan IoMT, termasuk kebutuhan keandalan "
    "pengiriman data pada pemantauan pasien jarak jauh. Pada tataran protokol, Balas dkk. [8] mengamati "
    "perilaku BLE pada aplikasi pemantauan detak jantung dan menunjukkan bahwa parameter koneksi "
    "memengaruhi kualitas pengiriman data, sementara Xu dkk. [11] memodelkan kinerja BLE dan "
    "memperlihatkan pengaruh parameter lapisan tautan terhadap throughput. Pada tataran jaringan, "
    "survei Majumdar dkk. [10] merangkum pendekatan keandalan, ketangguhan, dan efisiensi energi pada "
    "wireless body area network (WBAN), dan Xiong dan Jiang [9] mengembangkan protokol perutean "
    "delay-tolerant network (DTN) berbasis paradigma store-carry-forward untuk jaringan dengan "
    "konektivitas terputus-putus. Pada tataran perangkat, studi validasi Schweizer dan Gilgen-Ammann "
    "[12] menegaskan bahwa akurasi sensor detak jantung pergelangan tangan bervariasi antar-kondisi "
    "aktivitas, sehingga kualitas kontak sensor perlu dilaporkan bersama data.")
insert_before(kontrib,
    "Meskipun demikian, studi-studi tersebut umumnya menitikberatkan pada perilaku lapisan tautan, "
    "protokol perutean, atau akurasi sensor, dan belum membahas jaminan kelengkapan data pada tingkat "
    "aplikasi untuk pasangan perangkat konsumer (smartwatch Wear OS dan smartphone Android) dalam "
    "kondisi koneksi putus-sambung — termasuk verifikasi empiris bahwa seluruh record yang direkam "
    "benar-benar tiba di penerima tanpa duplikasi. Celah itulah yang diisi oleh penelitian ini melalui "
    "kombinasi store-and-forward, konfirmasi tingkat aplikasi (ACK), penerima idempoten, dan evaluasi "
    "kelengkapan berbasis pencocokan timestamp.")

set_text(find(doc, "[7]"),
    "[7] P. He, D. Huang, D. Wu, H. He, Y. Wei, Y. Cui, R. Wang, dan L. Peng, “A survey of Internet of "
    "medical things: technology, application and future directions,” Digital Communications and "
    "Networks, vol. 12, no. 5, hlm. 717–742, 2026, doi: 10.1016/j.dcan.2024.11.013.")
set_text(find(doc, "[8]"),
    "[8] Z. Balas, K. Tokarz, B. Zieliński, dan T. Guźniczak, “Research on the behaviour of Bluetooth "
    "Low Energy protocol in the heart rate monitoring application,” Procedia Computer Science, "
    "vol. 225, hlm. 63–69, 2023, doi: 10.1016/j.procs.2023.09.092.")
set_text(find(doc, "[9]"),
    "[9] Y. Xiong dan S. Jiang, “Multi-decision dynamic intelligent routing protocol for "
    "delay-tolerant networks,” Electronics, vol. 12, no. 21, art. 4528, 2023, "
    "doi: 10.3390/electronics12214528.")
p10 = find(doc, "[10]")
set_text(p10,
    "[10] P. Majumdar, S. Roy, S. Sikdar, P. Ghosh, dan N. Ghosh, “A survey on data-driven approaches "
    "for reliability, robustness, and energy efficiency in wireless body area networks,” Sensors, "
    "vol. 24, no. 20, art. 6531, 2024, doi: 10.3390/s24206531.")
p11 = insert_after(p10,
    "[11] H. Xu, Z. Yan, B. Li, dan M. Yang, “Modeling and analysis of the performance for Bluetooth "
    "Low Energy,” IEEE Communications Letters, vol. 28, no. 3, hlm. 732–736, 2024, "
    "doi: 10.1109/LCOMM.2024.3352545.")
insert_after(p11,
    "[12] T. Schweizer dan R. Gilgen-Ammann, “Wrist-worn and arm-worn wearables for monitoring heart "
    "rate during sedentary and light-to-vigorous physical activities: device validation study,” JMIR "
    "Cardio, vol. 9, art. e67110, 2025, doi: 10.2196/67110.")

doc.save(str(ROOT / "Draft_Naskah_HR_BLE_SINTA2.docx"))
print("OK -> Draft_Naskah_HR_BLE_SINTA2.docx")

# ============================ MANUSCRIPT EN ============================
doc = Document(str(ROOT / "Draft_Manuscript_HR_BLE_EN.docx"))

replace_in_paragraphs(doc, [
    ("physiological data [[add citation: IoMT survey]]", "physiological data [7]"),
    ("the receiver application is stopped by the system [[add citation: related wearable BLE/mHealth "
     "study]]",
     "the receiver application is stopped by the system [8], [10]"),
    ("bounded by the Maximum Transmission Unit (MTU)",
     "bounded by the Maximum Transmission Unit (MTU) [11]"),
    ("the sensor on a consumer smartwatch is not clinically validated",
     "the sensor on a consumer smartwatch is not clinically validated and its accuracy is known to "
     "vary across activity conditions [12]"),
])

contrib = find(doc, "The contributions of this work are")
insert_before(contrib,
    "Several recent studies have examined aspects adjacent to this work. He et al. [7] mapped the "
    "architecture, applications, and challenges of the IoMT, including the need for reliable data "
    "delivery in remote patient monitoring. At the protocol level, Balas et al. [8] observed the "
    "behaviour of BLE in a heart-rate monitoring application and showed that connection parameters "
    "affect the quality of data delivery, while Xu et al. [11] modelled BLE performance and "
    "demonstrated the influence of link-layer parameters on throughput. At the network level, the "
    "survey by Majumdar et al. [10] summarised reliability, robustness, and energy-efficiency "
    "approaches in wireless body area networks (WBANs), and Xiong and Jiang [9] developed a "
    "delay-tolerant network (DTN) routing protocol based on the store-carry-forward paradigm for "
    "intermittently connected networks. At the device level, the validation study by Schweizer and "
    "Gilgen-Ammann [12] confirmed that the accuracy of wrist-worn heart-rate sensors varies across "
    "activity conditions, so sensor contact quality should be reported alongside the data.")
insert_before(contrib,
    "Nevertheless, these studies largely focus on link-layer behaviour, routing protocols, or sensor "
    "accuracy, and do not address application-level data-completeness guarantees on a consumer device "
    "pair (a Wear OS smartwatch and an Android smartphone) under disconnect–reconnect conditions — "
    "including empirical verification that every recorded reading actually arrives at the receiver "
    "without duplication. This work fills that gap through the combination of store-and-forward, "
    "application-level acknowledgement (ACK), an idempotent receiver, and completeness evaluation "
    "based on timestamp matching.")

set_text(find(doc, "[7]"),
    "[7] P. He, D. Huang, D. Wu, H. He, Y. Wei, Y. Cui, R. Wang, and L. Peng, “A survey of Internet of "
    "medical things: technology, application and future directions,” Digital Communications and "
    "Networks, vol. 12, no. 5, pp. 717–742, 2026, doi: 10.1016/j.dcan.2024.11.013.")
set_text(find(doc, "[8]"),
    "[8] Z. Balas, K. Tokarz, B. Zieliński, and T. Guźniczak, “Research on the behaviour of Bluetooth "
    "Low Energy protocol in the heart rate monitoring application,” Procedia Computer Science, "
    "vol. 225, pp. 63–69, 2023, doi: 10.1016/j.procs.2023.09.092.")
set_text(find(doc, "[9]"),
    "[9] Y. Xiong and S. Jiang, “Multi-decision dynamic intelligent routing protocol for "
    "delay-tolerant networks,” Electronics, vol. 12, no. 21, art. 4528, 2023, "
    "doi: 10.3390/electronics12214528.")
p10 = find(doc, "[10]")
set_text(p10,
    "[10] P. Majumdar, S. Roy, S. Sikdar, P. Ghosh, and N. Ghosh, “A survey on data-driven approaches "
    "for reliability, robustness, and energy efficiency in wireless body area networks,” Sensors, "
    "vol. 24, no. 20, art. 6531, 2024, doi: 10.3390/s24206531.")
p11 = insert_after(p10,
    "[11] H. Xu, Z. Yan, B. Li, and M. Yang, “Modeling and analysis of the performance for Bluetooth "
    "Low Energy,” IEEE Communications Letters, vol. 28, no. 3, pp. 732–736, 2024, "
    "doi: 10.1109/LCOMM.2024.3352545.")
insert_after(p11,
    "[12] T. Schweizer and R. Gilgen-Ammann, “Wrist-worn and arm-worn wearables for monitoring heart "
    "rate during sedentary and light-to-vigorous physical activities: device validation study,” JMIR "
    "Cardio, vol. 9, art. e67110, 2025, doi: 10.2196/67110.")

doc.save(str(ROOT / "Draft_Manuscript_HR_BLE_EN.docx"))
print("OK -> Draft_Manuscript_HR_BLE_EN.docx")
