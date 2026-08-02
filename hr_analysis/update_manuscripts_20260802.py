#!/usr/bin/env python3
"""Revisi naskah 2 Agu 2026 — penyelarasan klaim & penambahan pembahasan.

Sasaran: kedua naskah siap disubmit ke jurnal SINTA 2 sebagai paper
RANCANG BANGUN, bukan paper klaim keandalan.

Perubahan:
1. Judul & abstrak diubah ke bingkai rancang-bangun; angka utama memakai
   hasil versi final (99,90%), bukan agregat dua versi (89,04%).
2. Tabel 2 disusun ulang menjadi dua blok versi perangkat lunak dengan
   subtotal masing-masing; baris Total lintas-versi DIHAPUS karena tidak
   dapat ditafsirkan. Sesi 2 (53 record, smartphone tidak pernah terhubung)
   dikeluarkan dari agregat dan dilaporkan terpisah sebagai sesi gagal.
       versi awal (Sesi 1+3+4) : 386 mnt, 23.162 / 20.455 / 2.707 / 88,31%
       versi revisi (Sesi 5)   :  33 mnt,  1.976 /  1.974 /     2 / 99,90%
3. §3.5 dihitung ulang tanpa Sesi 2: 2.707 hilang = 2.702 false-sent + 5 pending
   (sebelumnya 2.760 = 2.755 + 5, termasuk 53 false-sent milik Sesi 2).
4. §3.3 diseragamkan cakupannya: Tabel 3 dan Gambar 8–9 kini sama-sama lima
   sesi (make_figs.py telah diperluas), sehingga 91,6/0,1/8,4% dan n=23.063
   konsisten antara tabel, gambar, dan teks.
5. §3.4 menyatakan throughput dalam satuan yang sama (KiB/detik).
6. §2.5 memuat rekonsiliasi inklusi data: 47.422 baris mentah → 25.191
   periode pengukuran (22.231 baris ekor off-wrist dikecualikan).
7. PEMBAHASAN ditambahkan pada §3.2, §3.3, §3.4, §3.5, dan §3.6 — interpretasi
   hasil, keterkaitan dengan literatur, dan batas penafsiran.
8. Bagian "Ketersediaan Data dan Kode" ditambahkan.

Angka diverifikasi ulang dari CSV mentah memakai pipeline hr_analysis/make_figs.py.
"""
import copy
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[1]


# --------------------------------------------------------------------------
# util
# --------------------------------------------------------------------------
def set_text(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        p.add_run(text)


def find(doc, prefix):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    raise SystemExit(f"PARAGRAF TIDAK DITEMUKAN: {prefix!r}")


def replace_para(doc, prefix, text):
    set_text(find(doc, prefix), text)


def insert_after(ref, *texts):
    """Sisipkan paragraf baru setelah `ref`, mewarisi format `ref`."""
    last = ref
    for t in texts:
        el = copy.deepcopy(ref._p)
        last._p.addnext(el)
        new = Paragraph(el, ref._parent)
        set_text(new, t)
        last = new
    return last


def rebuild_table(tbl, items):
    """Bangun ulang isi tabel (baris header dipertahankan).

    item str        -> baris judul kelompok (sel digabung, miring)
    item (tuple)    -> baris data; tandai tebal lewat item[-1] is True
    """
    for r in list(tbl.rows)[1:]:
        tbl._tbl.remove(r._tr)
    for item in items:
        row = tbl.add_row()
        if isinstance(item, str):
            cell = row.cells[0].merge(row.cells[-1])
            cell.text = item
            for p in cell.paragraphs:
                for r in p.runs:
                    r.italic = True
            continue
        bold = item[-1] is True
        values = item[:-1] if isinstance(item[-1], bool) else item
        for i, v in enumerate(values):
            row.cells[i].text = v
            if bold:
                for p in row.cells[i].paragraphs:
                    for r in p.runs:
                        r.bold = True


# --------------------------------------------------------------------------
# isi Tabel 2 (identik struktur untuk kedua naskah)
# --------------------------------------------------------------------------
TABEL2_ID = [
    "Versi awal — markSynced dijalankan tanpa menunggu ACK",
    ("Sesi 1", "23/06 06:38", "36 mnt", "2.177", "2.164", "13", "99,40%"),
    ("Sesi 3", "23/06 14:15", "345 mnt", "20.702", "18.013", "2.689", "87,01%"),
    ("Sesi 4", "28/06 21:29", "5 mnt", "283", "278", "5", "98,23%"),
    ("Subtotal versi awal", "—", "386 mnt", "23.162", "20.455", "2.707", "88,31%", True),
    "Versi revisi — markSynced dijalankan hanya setelah ACK diterima",
    ("Sesi 5", "05/07 23:07", "33 mnt", "1.976", "1.974", "2", "99,90%"),
    ("Subtotal versi revisi", "—", "33 mnt", "1.976", "1.974", "2", "99,90%", True),
]

TABEL2_EN = [
    "Initial version — markSynced executed without waiting for an ACK",
    ("Session 1", "23/06 06:38", "36 min", "2,177", "2,164", "13", "99.40%"),
    ("Session 3", "23/06 14:15", "345 min", "20,702", "18,013", "2,689", "87.01%"),
    ("Session 4", "28/06 21:29", "5 min", "283", "278", "5", "98.23%"),
    ("Subtotal, initial version", "—", "386 min", "23,162", "20,455", "2,707", "88.31%", True),
    "Revised version — markSynced executed only after an ACK is received",
    ("Session 5", "05/07 23:07", "33 min", "1,976", "1,974", "2", "99.90%"),
    ("Subtotal, revised version", "—", "33 min", "1,976", "1,974", "2", "99.90%", True),
]


# ==========================================================================
# NASKAH INDONESIA
# ==========================================================================
doc = Document(str(ROOT / "Draft_Naskah_HR_BLE_SINTA2.docx"))

# --- judul ---------------------------------------------------------------
replace_para(doc, "Pengiriman Data Detak Jantung yang Andal",
    "Rancang Bangun Sistem Pengiriman Data Detak Jantung Smartwatch–Smartphone "
    "Berbasis Bluetooth Low Energy dengan Skema Store-and-Forward dan Konfirmasi Penerimaan")

# --- abstrak -------------------------------------------------------------
replace_para(doc, "Pemantauan detak jantung berkelanjutan menggunakan wearable",
    "Pemantauan detak jantung berkelanjutan menggunakan wearable memerlukan pengiriman yang "
    "menjaga integritas data meskipun ukuran notifikasi BLE terbatas dan koneksi dapat terputus. "
    "Penelitian ini merancang dan mengimplementasikan sistem dua aplikasi: smartwatch Wear OS "
    "merekam satu record per detik ke SQLite dan mengirim batch ber-opcode sebagai peripheral/GATT "
    "server, sedangkan smartphone Android merangkai frame, menyimpan record secara idempoten, dan "
    "mengirim konfirmasi penerimaan (ACK) tingkat aplikasi. Sistem diuji pada sepasang perangkat "
    "konsumer dalam lima sesi (23 Juni–5 Juli 2026; 25.191 pembacaan periode pengukuran) yang "
    "mencakup dua versi perangkat lunak, sehingga hasil dilaporkan terpisah per versi. Versi awal, "
    "yang menandai record sebagai terkirim tanpa menunggu konfirmasi, mencapai delivery 88,31% "
    "(20.455 dari 23.162 record); analisis silang mengungkap 2.702 dari 2.707 record yang hilang "
    "(99,8%) ternyata sudah bertanda synced sehingga tidak pernah dijadwalkan ulang. Setelah "
    "markSynced direvisi agar dijalankan hanya setelah ACK diterima, sesi versi revisi mencapai "
    "99,90% (1.974 dari 1.976 record) tanpa record keliru-tanda dan seluruh tujuh batch "
    "terkonfirmasi. Pada kedua versi, setiap record yang diterima memiliki bpm dan accuracy "
    "identik dengan sumbernya (fidelitas nilai 100%) tanpa duplikat basis data. Pada MTU 512, "
    "batch 180 record (8.461 byte, 19 frame) terkirim dalam 121,5 ± 25,5 ms, setara 0,07% dari "
    "interval kirim 3 menit. Kontribusi utama penelitian ini adalah rancang bangun sistem beserta "
    "bukti empiris bahwa penandaan status kirim yang mendahului konfirmasi merupakan sumber "
    "kehilangan data yang dapat dihilangkan; ketahanan versi revisi pada skenario putus–sambung "
    "terkendali masih memerlukan pengujian lanjutan.")

replace_para(doc, "Kata kunci:",
    "Kata kunci: rancang bangun; IoMT; Bluetooth Low Energy; wearable; detak jantung; "
    "store-and-forward; konfirmasi penerimaan; Wear OS")

# --- abstract (Inggris, di naskah ID) ------------------------------------
ABSTRACT_EN = (
    "Continuous wearable heart-rate monitoring requires data delivery that preserves integrity "
    "despite the limited BLE notification size and intermittent connectivity. This study designs "
    "and implements a two-application system: a Wear OS smartwatch records one SQLite row per "
    "second and transmits opcode-framed batches as a peripheral/GATT server, while an Android "
    "smartphone reassembles the frames, persists the records idempotently, and returns an "
    "application-level acknowledgement (ACK). The system was tested on a consumer device pair "
    "across five sessions (23 June–5 July 2026; 25,191 measurement-period readings) spanning two "
    "software versions, so results are reported per version. The initial version, which marked "
    "records as sent without waiting for confirmation, achieved 88.31% delivery (20,455 of 23,162 "
    "records); a cross-analysis revealed that 2,702 of the 2,707 lost records (99.8%) were already "
    "flagged as synced and were therefore never rescheduled. After markSynced was revised to run "
    "only upon receiving an ACK, the revised-version session reached 99.90% (1,974 of 1,976 "
    "records) with no falsely flagged records and all seven batches acknowledged. In both "
    "versions, every received record had bpm and accuracy values identical to the source (100% "
    "value fidelity) with no database duplicates. At MTU 512, a 180-record batch (8,461 bytes, 19 "
    "frames) was transferred in 121.5 ± 25.5 ms, equal to 0.07% of the 3-minute sending interval. "
    "The main contribution is the system design together with empirical evidence that marking "
    "send status ahead of confirmation is a removable source of data loss; the resilience of the "
    "revised version under controlled disconnect–reconnect scenarios still requires further testing."
)
replace_para(doc, "Continuous wearable heart-rate monitoring", ABSTRACT_EN)
replace_para(doc, "Keywords:",
    "Keywords: system design; IoMT; Bluetooth Low Energy; wearable; heart rate; store-and-forward; "
    "acknowledgement; Wear OS")

# --- §2.5 perangkat & skenario ------------------------------------------
replace_para(doc, "Pengujian dilakukan pada perangkat fisik agar perilaku radio",
    "Pengujian dilakukan pada perangkat fisik agar perilaku radio, sensor, dan eksekusi latar "
    "belakang yang tidak sepenuhnya tercakup oleh emulator dapat diamati. Spesifikasi perangkat "
    "uji ditunjukkan pada Tabel 1. Aplikasi dipasang dalam mode rilis agar pengukuran kinerja "
    "tidak terbias oleh mode debug. Pengujian mencakup empat sesi pada 23–28 Juni 2026 yang "
    "menggunakan implementasi awal dan satu sesi pada 5 Juli 2026 yang menggunakan revisi "
    "penantian ACK sebelum markSynced. Karena kedua rangkaian tersebut menjalankan perangkat "
    "lunak yang berbeda, hasilnya dilaporkan terpisah per versi dan tidak dijumlahkan menjadi satu "
    "angka keberhasilan tunggal.")

insert_after(find(doc, "Pengujian dilakukan pada perangkat fisik agar perilaku radio"),
    "Inklusi data dilakukan sebagai berikut. Smartwatch menghasilkan 47.422 baris rekaman pada "
    "kelima sesi. Sebanyak 22.231 baris pada ekor rekaman Juni dikecualikan karena bernilai beku "
    "(BPM konstan) dengan akurasi ≤ 0 akibat smartwatch dilepas dari pergelangan tangan; batas "
    "periode pengukuran ditetapkan pada pembacaan berakurasi tertinggi yang terakhir pada tiap "
    "sesi. Sisanya, 25.191 pembacaan, membentuk periode pengukuran yang dipakai pada seluruh "
    "metrik di Bagian 3. Satu sesi tambahan pada 23 Juni pukul 08.41 (53 pembacaan, 1 menit) "
    "dikeluarkan dari agregat pengiriman karena smartphone tidak pernah berhasil terhubung "
    "sehingga sesi tersebut tidak menguji jalur pengiriman; sesi ini tetap disertakan pada "
    "statistik kualitas sensor yang tidak bergantung pada koneksi.")

# --- Tabel 1: baris dataset ---------------------------------------------
t1 = doc.tables[0]
for row in t1.rows:
    if row.cells[0].text.strip().startswith("Dataset uji"):
        row.cells[1].text = ("23 Jun–5 Jul 2026; 5 sesi (4 versi awal, 1 versi revisi ACK); "
                             "25.191 pembacaan periode pengukuran dari 47.422 baris rekaman")

replace_para(doc, "Tabel tersebut merangkum lingkungan pengujian",
    "Tabel tersebut merangkum lingkungan pengujian yang dijaga tetap: sepasang perangkat konsumer "
    "(Samsung Galaxy Watch SM-R860 sebagai peripheral dan Xiaomi Redmi Note 10 Pro sebagai "
    "central), kerangka kerja Flutter dengan aplikasi terpasang dalam mode rilis, MTU yang diminta "
    "512 byte, serta interval pengiriman yang dapat dipilih 3/5 menit. Dataset mencakup lima sesi "
    "pada 23 Juni–5 Juli 2026 dengan 25.191 pembacaan periode pengukuran, disaring dari 47.422 "
    "baris rekaman mentah.")

# --- §3.2 delivery -------------------------------------------------------
replace_para(doc, "Evaluasi mencakup lima sesi",
    "Hasil pengiriman dilaporkan terpisah menurut versi perangkat lunak pada Tabel 2. Versi awal "
    "diuji pada tiga sesi valid berdurasi total 386 menit dan mencatat 20.455 dari 23.162 record "
    "(88,31%). Versi revisi, yang menunda penandaan synced sampai ACK diterima, diuji pada satu "
    "sesi berdurasi 33 menit dan mencatat 1.974 dari 1.976 record (99,90%); dua record yang "
    "tersisa masih berstatus belum terkirim karena direkam setelah jadwal pengiriman terakhir, "
    "bukan karena hilang. Seluruh tujuh batch pada sesi tersebut memperoleh ACK. Pada kedua versi "
    "tidak ditemukan satu pun duplikat di basis data penerima, dan seluruh record yang berhasil "
    "dicocokkan memiliki nilai bpm dan accuracy identik dengan catatan smartwatch (fidelitas nilai "
    "100%).")

rebuild_table(doc.tables[1], TABEL2_ID)
replace_para(doc, "Tabel 2. Hasil pengiriman",
    "Tabel 2. Hasil pengiriman per sesi, dikelompokkan menurut versi perangkat lunak.")

replace_para(doc, "Tabel menunjukkan rentang delivery",
    "Subtotal kedua kelompok sengaja tidak dijumlahkan menjadi satu angka lintas-versi karena "
    "keduanya menjalankan perangkat lunak yang berbeda; angka gabungan semacam itu tidak "
    "menggambarkan kinerja versi mana pun. Sesi pada 23 Juni pukul 08.41 (53 pembacaan, 1 menit) "
    "tidak dimasukkan ke dalam tabel karena smartphone tidak pernah berhasil terhubung sepanjang "
    "sesi, sehingga tidak ada jalur pengiriman yang diuji; sesi tersebut dilaporkan sebagai sesi "
    "gagal, bukan sebagai pengukuran delivery 0%.")

# pembahasan §3.2 (disisipkan setelah penjelasan Gambar 7)
insert_after(find(doc, "Gambar tersebut menampilkan sinyal detak jantung sesi utama"),
    "Pemisahan hasil menurut versi memperlihatkan pola yang tertutup oleh angka agregat. Pada "
    "versi awal, kehilangan tidak tersebar merata sepanjang sesi melainkan menumpuk pada periode "
    "sebelum smartphone terhubung: Sesi 3 kehilangan 2.689 record yang hampir seluruhnya berada "
    "sebelum koneksi tercatat sekitar pukul 15.00 (Gambar 7), sedangkan Sesi 1 dan Sesi 4 yang "
    "smartphone-nya terhubung sejak awal hanya kehilangan 13 dan 5 record. Fidelitas nilai 100% "
    "pada seluruh record yang diterima memperkuat pembacaan tersebut: tidak satu pun record sampai "
    "dalam keadaan rusak, sehingga kehilangan pada versi awal merupakan kegagalan penjadwalan "
    "ulang pengiriman, bukan kegagalan integritas kanal. Pengamatan ini sejalan dengan kajian "
    "perilaku BLE pada aplikasi pemantauan detak jantung yang menempatkan kondisi tautan, bukan "
    "kerusakan payload, sebagai penentu kualitas pengiriman [8].",
    "Selisih antara 88,31% dan 99,90% karena itu tidak dapat dibaca sebagai perbaikan kinerja "
    "radio, melainkan sebagai akibat perubahan satu keputusan perangkat lunak: kapan sebuah record "
    "boleh dinyatakan terkirim. Penafsiran ini tetap memiliki batas. Kedua angka berasal dari "
    "kondisi pengujian yang tidak identik — versi awal diuji pada sesi yang jauh lebih panjang dan "
    "sempat berjalan tanpa koneksi, sedangkan versi revisi diuji selama 33 menit dengan koneksi "
    "kontinu — sehingga perbandingan langsung keduanya bersifat indikatif, bukan pembuktian "
    "kausal. Pembuktian kausal menuntut eksperimen ablasi dengan kedua versi dijalankan pada "
    "skenario gangguan yang sama dan direplikasi beberapa kali, yang menjadi agenda pengujian "
    "berikutnya.")

# --- Gambar 6 & 8 & 9: cakupan lima sesi --------------------------------
replace_para(doc, "Gambar 6. Kelengkapan data",
    "Gambar 6. Kelengkapan data pada lima sesi evaluasi (diterima vs hilang).")
replace_para(doc, "Diagram tersebut membandingkan jumlah record yang diterima",
    "Diagram tersebut membandingkan jumlah record yang diterima dan yang hilang pada kelima sesi. "
    "Kehilangan terkonsentrasi pada sesi terpanjang (Sesi 3, 13,0% hilang) dan pada sesi 23 Juni "
    "pukul 08.41 yang smartphone-nya tidak pernah terhubung, sedangkan tiga sesi lainnya hampir "
    "lengkap (99,40%, 98,23%, dan 99,90%). Perbedaan panjang batang juga memperlihatkan bahwa "
    "sebagian besar volume data berasal dari satu sesi panjang, sehingga rata-rata tak-berbobot "
    "antarsesi akan menyesatkan.")

# --- §3.3 kualitas data sensor ------------------------------------------
replace_para(doc, "Distribusi akurasi sensor pada periode pengukuran",
    "Distribusi akurasi sensor sepanjang periode pengukuran kelima sesi ditunjukkan pada Tabel 3 "
    "dan Gambar 8. Sebanyak 91,6% pembacaan (23.063 pembacaan) berada pada akurasi tertinggi "
    "(nilai 3), 8,4% (2.113 pembacaan) pada kondisi tanpa kontak (nilai -1) yang lazim terjadi "
    "saat sensor sesaat kehilangan kontak dengan kulit [12], dan 0,1% (15 pembacaan) pada nilai 0; "
    "pada Sesi 5 seluruh pembacaan berkontak baik. Statistik nilai detak jantung pada pembacaan "
    "berakurasi tertinggi (n = 23.063, divisualkan pada Gambar 9) adalah minimum 60 bpm, maksimum "
    "123 bpm, rata-rata 83,5 bpm, dan simpangan baku 9,6 bpm. Tabel 3 beserta Gambar 8 dan Gambar "
    "9 memakai cakupan yang sama, yaitu periode pengukuran kelima sesi, sehingga angkanya dapat "
    "dibandingkan langsung. Sebanyak 22.231 baris pada ekor rekaman — bernilai beku dengan akurasi "
    "≤ 0 akibat smartwatch tidak terpasang — telah dikecualikan sebagaimana diuraikan pada "
    "Bagian 2.5.")

replace_para(doc, "Tabel 3. Distribusi nilai akurasi sensor",
    "Tabel 3. Distribusi nilai akurasi sensor (periode pengukuran, lima sesi).")

replace_para(doc, "Dari tabel terlihat kualitas kontak didominasi",
    "Dari tabel terlihat kualitas kontak didominasi tingkat tertinggi (91,6%), dengan porsi tanpa "
    "kontak 8,4% dan status tidak dapat dipercaya hanya 0,1%. Komposisi tersebut divisualkan pada "
    "Gambar 8, sedangkan sebaran nilai detak jantung pada pembacaan berakurasi tertinggi "
    "ditunjukkan pada Gambar 9.")

replace_para(doc, "Gambar 8. Distribusi kualitas kontak sensor",
    "Gambar 8. Distribusi kualitas kontak sensor pada lima sesi evaluasi.")
replace_para(doc, "Diagram tersebut merangkum kualitas kontak sensor",
    "Diagram tersebut merangkum kualitas kontak sensor pada kelima sesi: mayoritas pembacaan "
    "(91,6%) berada pada tingkat akurasi tertinggi, 8,4% tanpa kontak, dan 0,1% pada status tidak "
    "dapat dipercaya — mencerminkan lepas-kontak sesaat yang lazim pada sensor optik di "
    "pergelangan tangan.")

replace_para(doc, "Histogram tersebut menunjukkan sebaran nilai detak jantung",
    "Histogram tersebut menunjukkan sebaran nilai detak jantung pada pembacaan berakurasi "
    "tertinggi kelima sesi (n = 23.063): terpusat di sekitar rata-rata 83,5 bpm dengan simpangan "
    "baku 9,6 bpm dan rentang 60–123 bpm. Profil ini wajar untuk aktivitas ringan dan "
    "mengindikasikan data fisiologis yang masuk akal, bukan nilai beku akibat sensor tidak "
    "terpasang.")

# pembahasan §3.3
insert_after(find(doc, "Histogram tersebut menunjukkan sebaran nilai detak jantung"),
    "Angka 8,4% tanpa kontak sebaiknya dibaca sebagai karakteristik sensor optik di pergelangan "
    "tangan, bukan sebagai cacat sistem pengiriman. Studi validasi perangkat wearable melaporkan "
    "bahwa akurasi sensor pergelangan tangan memang bervariasi mengikuti kondisi aktivitas dan "
    "kualitas kontak kulit [12], sehingga proporsi pembacaan tanpa kontak pada kisaran satu digit "
    "merupakan hal wajar untuk pemakaian sehari-hari. Yang relevan bagi rancangan ini adalah bahwa "
    "status akurasi tersebut ikut dikirim dan tersimpan utuh di penerima: penapisan kualitas data "
    "dapat dilakukan di sisi analisis tanpa kehilangan informasi, karena sistem tidak membuang "
    "pembacaan berakurasi rendah di sisi sumber. Perbedaan menyolok antara Sesi 5 — yang seluruh "
    "pembacaannya berkontak baik — dan sesi-sesi Juni juga menunjukkan bahwa kualitas kontak lebih "
    "ditentukan oleh cara pemakaian daripada oleh perangkat lunak, sehingga proporsi ini tidak "
    "dapat dijadikan tolok ukur kinerja sistem antarsesi.")

# --- §3.4 kinerja transfer ----------------------------------------------
replace_para(doc, "Kinerja transfer diukur melalui instrumentasi HR-METRIC",
    "Kinerja transfer diukur melalui instrumentasi HR-METRIC pada smartwatch. Pada Sesi 5 "
    "(interval 3 menit, tujuh batch), seluruh batch berukuran seragam — 180 pembacaan, payload "
    "8.461 byte (47,0 byte per pembacaan), dan 19 frame pada MTU 512 — dengan durasi transfer "
    "rata-rata 121,5 ± 25,5 ms (rentang 92,8–166,4 ms) dan throughput rata-rata 70,3 ± 13,2 "
    "KiB/detik; setiap batch dikonfirmasi ACK penuh oleh smartphone. Jumlah frame terukur sesuai "
    "perhitungan ukuran chunk (MTU − 4 = 508 byte): 8.461 ÷ 508 = 16,7, dibulatkan ke atas menjadi "
    "17 frame DATA, ditambah START dan END menjadi 19 frame. Sebagai ilustrasi batch besar, pada "
    "evaluasi Juni satu batch berisi 228 pembacaan (payload 10.717 byte; 24 frame, konsisten "
    "dengan perhitungan yang sama) terkirim dalam ±0,32 detik, setara ±32,4 KiB/detik. Negosiasi "
    "MTU pada pasangan perangkat uji secara konsisten menghasilkan nilai penuh 512 byte pada kedua "
    "rangkaian pengujian.")

insert_after(find(doc, "Kinerja transfer diukur melalui instrumentasi HR-METRIC"),
    "Ditinjau dari penggunaan kanal, durasi 121,5 ms terhadap interval kirim 180 detik berarti "
    "radio hanya aktif mengirim data sekitar 0,07% waktu sesi. Rasio inilah yang memungkinkan "
    "skema batch periodik digabungkan dengan konfirmasi penerimaan tanpa beban kanal yang berarti: "
    "konfirmasi terjadi tujuh kali sepanjang sesi, bukan sekali untuk setiap pembacaan. Perbedaan "
    "tersebut penting karena penambahan konfirmasi ketat pada aliran waktu-nyata berkelanjutan "
    "justru dikenal menambah kepadatan lalu lintas dan menurunkan keandalan pada jaringan area "
    "tubuh [10]. Dengan kata lain, penundaan pengiriman selama beberapa menit — yang dapat "
    "diterima untuk pemantauan tren detak jantung, meskipun tidak untuk peringatan waktu-nyata — "
    "adalah harga yang dibayar untuk memperoleh kelengkapan data dengan biaya kanal yang rendah.",
    "Perlu dicatat bahwa throughput kedua rangkaian pengujian berbeda cukup jauh, yaitu 70,3 "
    "KiB/detik pada Sesi 5 berbanding sekitar 32,4 KiB/detik pada batch Juni, meskipun MTU yang "
    "dinegosiasikan sama-sama 512 byte. Karena parameter koneksi BLE seperti connection interval "
    "tidak dicatat pada rangkaian Juni, selisih tersebut belum dapat diatribusikan pada satu "
    "sebab; parameter koneksi telah dilaporkan berpengaruh nyata terhadap kualitas pengiriman pada "
    "aplikasi pemantauan detak jantung [8] dan terhadap kinerja BLE secara umum [11]. Pencatatan "
    "parameter koneksi karena itu ditambahkan sebagai bagian instrumentasi pada pengujian "
    "berikutnya. Pengukuran lintas variasi kondisi lain — interval 5 menit, jarak antar-perangkat, "
    "dan skenario gangguan — juga merupakan bagian dari pekerjaan lanjutan.")

# --- §3.5 temuan ---------------------------------------------------------
replace_para(doc, "Analisis silang empat sesi Juni",
    "Analisis silang tiga sesi valid versi awal menemukan bahwa dari 2.707 record yang tidak "
    "ditemukan pada smartphone, 2.702 (99,8%) sudah berstatus synced = 1 dan hanya 5 yang masih "
    "berstatus pending. Dengan kata lain, hampir seluruh data yang hilang bukan sedang menunggu "
    "giliran kirim, melainkan sudah dianggap selesai oleh smartwatch. Pada versi revisi, Sesi 5 "
    "tidak menghasilkan satu pun record keliru-tanda: tujuh batch masing-masing menerima ACK 180 "
    "record dan dua record terakhir tetap berstatus synced = 0 sebagaimana seharusnya. Namun Sesi "
    "5 tidak mencakup skenario putus–sambung, sehingga ketahanan versi revisi terhadap pemutusan "
    "tautan masih harus diuji secara khusus.")

replace_para(doc, "Bila dihitung atas seluruh rekaman smartwatch pada evaluasi Juni",
    "Bila dihitung atas seluruh rekaman smartwatch pada evaluasi Juni (45.446 baris, termasuk "
    "periode sensor tidak terpasang), jumlah record bertanda terkirim yang tidak ditemukan di "
    "smartphone adalah 2.879. Angka 2.702 di atas adalah bagian dari jumlah tersebut yang berada "
    "pada periode pengukuran tiga sesi valid, sehingga konsisten dengan Tabel 2; sisanya berasal "
    "dari ekor rekaman off-wrist dan dari sesi 08.41 yang dikeluarkan dari agregat.")

insert_after(find(doc, "Bila dihitung atas seluruh rekaman smartwatch pada evaluasi Juni"),
    "Akar masalahnya terletak pada urutan operasi. Pada versi awal, penandaan synced = 1 "
    "dieksekusi segera setelah panggilan pengiriman dikembalikan oleh lapisan BLE, bukan setelah "
    "penerima mengonfirmasi penyimpanan. Karena notifikasi BLE bersifat best-effort dan tidak "
    "menyediakan konfirmasi pada lapisan atribut [1], keberhasilan panggilan pengiriman hanya "
    "berarti frame telah diserahkan ke tumpukan protokol lokal — bukan bahwa frame sampai dan "
    "tersimpan di penerima. Akibatnya record yang tidak pernah terkirim tetap dianggap selesai dan "
    "tidak pernah masuk kembali ke antrean, sehingga penyangga store-and-forward kehilangan "
    "fungsinya justru pada saat paling dibutuhkan, yaitu ketika tautan sedang terputus.",
    "Temuan ini bersifat umum bagi setiap rancangan penyangga berbasis penanda status: penanda "
    "kirim hanya boleh dimutakhirkan oleh peristiwa yang berasal dari penerima, tidak boleh oleh "
    "peristiwa lokal di sisi pengirim. Prinsip yang sama mendasari skema store-carry-forward pada "
    "jaringan toleran-tunda, tempat sebuah pesan baru boleh dilepas dari penyangga setelah "
    "penerusan berikutnya terkonfirmasi [9]. Nilai praktis dari temuan ini adalah bahwa kegagalan "
    "tersebut tidak tampak pada pengujian fungsional biasa — sistem terlihat berjalan normal, "
    "antarmuka melaporkan seluruh data terkirim, dan kehilangan baru terungkap ketika basis data "
    "kedua sisi dibandingkan record per record. Karena itu, verifikasi kelengkapan berbasis "
    "pencocokan timestamp seperti yang dipakai di sini sebaiknya menjadi bagian baku pengujian "
    "sistem pemantauan berkelanjutan, bukan sekadar analisis tambahan.")

# --- §3.6 pembahasan & keterbatasan -------------------------------------
replace_para(doc, "Hasil menunjukkan bahwa framing ber-opcode dan kendali aliran",
    "Secara keseluruhan, hasil pengujian menunjukkan bahwa rancangan protokol — framing ber-opcode "
    "dengan kendali aliran, penerima idempoten, dan konfirmasi tingkat aplikasi — bekerja "
    "sebagaimana dirancang pada perangkat nyata: seluruh record yang diterima memiliki fidelitas "
    "nilai 100% dan tidak ada duplikat pada kedua versi, sedangkan transfer batch berlangsung dua "
    "hingga tiga orde besaran lebih cepat daripada interval pengirimannya. Kontribusi empiris yang "
    "paling bernilai justru datang dari kegagalan versi awal: kegagalan tersebut menunjukkan bahwa "
    "keberadaan penyangga store-and-forward saja tidak cukup, dan kebenaran mekanisme bergantung "
    "pada satu detail implementasi yang mudah terlewat, yaitu sumber peristiwa yang berhak "
    "memutakhirkan penanda status kirim.")

insert_after(find(doc, "Secara keseluruhan, hasil pengujian menunjukkan bahwa rancangan protokol"),
    "Dibandingkan dengan pekerjaan terdahulu, penelitian ini menempati posisi yang berbeda. Kajian "
    "pada tataran protokol menyoroti pengaruh parameter koneksi terhadap kualitas pengiriman BLE "
    "[8] dan memodelkan kinerjanya secara analitis [11], sedangkan survei pada jaringan area tubuh "
    "merangkum pendekatan keandalan dan efisiensi energi pada tataran arsitektur jaringan [10]; "
    "ketiganya tidak menyertakan implementasi ujung-ke-ujung pada sepasang perangkat konsumer "
    "beserta verifikasi kelengkapan per record. Sebaliknya, skema store-carry-forward pada "
    "jaringan toleran-tunda [9] menyediakan gagasan penyanggaan yang serupa, tetapi diterapkan "
    "pada tataran perutean antar-simpul, bukan pada satu tautan wearable–gateway yang dibatasi MTU "
    "dan pembatasan eksekusi latar belakang sistem operasi seluler [3]. Kontribusi penelitian ini "
    "adalah menurunkan gagasan tersebut ke tataran aplikasi pada perangkat nyata, lengkap dengan "
    "bukti kuantitatif atas satu mode kegagalan yang tidak muncul dalam analisis tataran protokol "
    "maupun simulasi jaringan.")

replace_para(doc, "Penelitian ini memiliki beberapa keterbatasan.",
    "Penelitian ini memiliki beberapa keterbatasan yang membatasi generalisasi hasilnya. Pertama, "
    "pengujian hanya menggunakan satu pasangan perangkat pada satu lingkungan radio, sehingga "
    "angka latensi dan throughput tidak dapat digeneralisasi lintas perangkat. Kedua, versi revisi "
    "baru diuji pada satu sesi tanpa replikasi dan tanpa skenario putus–sambung terkendali, "
    "sehingga klaim ketahanannya terbatas pada kondisi koneksi kontinu yang telah diuji. Ketiga, "
    "ACK saat ini belum memuat batchId dan belum memvalidasi jumlah record yang dikonfirmasi "
    "terhadap batch yang tertunda, sehingga mekanisme ini meningkatkan keandalan tetapi belum "
    "memberi jaminan formal bebas kegagalan. Keempat, tidak ada variasi terkendali pada interval "
    "kirim, jarak, maupun MTU, dan konsumsi energi belum diukur. Kelima, operasi latar belakang "
    "belum menjamin kelanjutan setelah force-close atau reboot. Keenam, komunikasi belum "
    "menambahkan autentikasi maupun enkripsi tingkat aplikasi. Ketujuh, sensor smartwatch konsumer "
    "tidak dinilai sebagai perangkat klinis; kontribusi penelitian ini berfokus pada keandalan "
    "pengiriman data, bukan pada validitas medis nilai BPM.")

# --- kesimpulan ----------------------------------------------------------
replace_para(doc, "Penelitian ini merancang dan mengimplementasikan sistem pengiriman data",
    "Penelitian ini merancang dan mengimplementasikan sistem pengiriman data detak jantung "
    "berbasis BLE dengan framing batch ber-opcode, penyimpanan lokal, konfirmasi penerimaan "
    "tingkat aplikasi, penerima idempoten, dan foreground service, lalu mengujinya pada sepasang "
    "perangkat konsumer selama lima sesi. Versi awal mencapai delivery 88,31% dan mengungkap bahwa "
    "2.702 dari 2.707 record yang hilang sudah keliru bertanda terkirim sehingga tidak pernah "
    "dijadwalkan ulang; versi revisi yang menunda penandaan sampai ACK diterima mencapai 99,90% "
    "tanpa record keliru-tanda. Pada kedua versi, fidelitas nilai mencapai 100% tanpa duplikat, "
    "dan transfer tujuh batch 180 record pada MTU 512 memerlukan rata-rata 121,5 ms atau 0,07% "
    "dari interval kirimnya. Temuan utamanya adalah bahwa penanda status kirim harus dimutakhirkan "
    "oleh konfirmasi penerima, bukan oleh keberhasilan panggilan pengiriman lokal. Pengujian "
    "berikutnya perlu mereplikasi versi revisi pada skenario putus–sambung terkendali, menambahkan "
    "identitas batch pada ACK, mengukur konsumsi energi, serta memperluas pengujian lintas "
    "perangkat dan kondisi radio.")

# --- ketersediaan data & kode -------------------------------------------
ack = find(doc, "Ucapan Terima Kasih")
head = copy.deepcopy(ack._p)
ack._p.addprevious(head)
hp = Paragraph(head, ack._parent)
set_text(hp, "Ketersediaan Data dan Kode")
insert_after(hp,
    "Kode sumber kedua aplikasi, skrip analisis, serta berkas CSV hasil ekspor kelima sesi "
    "pengujian tersedia pada repositori publik penelitian ini di [[URL repositori]], sehingga "
    "seluruh angka pada Tabel 2, Tabel 3, dan Gambar 6–9 dapat direproduksi.")

replace_para(doc, "Catatan: dokumen ini adalah draf.",
    "Catatan: dokumen ini adalah draf. Sesuaikan dengan template resmi jurnal SINTA 2 tujuan "
    "(format kolom, gaya sitasi, batas halaman). Bagian bertanda [[...]] perlu dilengkapi penulis, "
    "termasuk URL repositori pada bagian Ketersediaan Data dan Kode.")

doc.save(str(ROOT / "Draft_Naskah_HR_BLE_SINTA2.docx"))
print("OK — Draft_Naskah_HR_BLE_SINTA2.docx")


# ==========================================================================
# NASKAH INGGRIS
# ==========================================================================
doc = Document(str(ROOT / "Draft_Manuscript_HR_BLE_EN_v2.docx"))

replace_para(doc, "Reliable Heart-Rate Data Delivery",
    "Design and Implementation of a Smartwatch-to-Smartphone Heart-Rate Delivery System over "
    "Bluetooth Low Energy with a Store-and-Forward and Acknowledgement Scheme")

replace_para(doc, "Continuous wearable heart-rate monitoring", ABSTRACT_EN)
replace_para(doc, "Keywords:",
    "Keywords: system design; IoMT; Bluetooth Low Energy; wearable; heart rate; store-and-forward; "
    "acknowledgement; Wear OS")

# --- contributions -------------------------------------------------------
replace_para(doc, "The contributions of this work are:",
    "The contributions of this work are: (1) the design and implementation of a two-application "
    "Wear OS–Android system for heart-rate acquisition and delivery; (2) an opcode-framed batch "
    "protocol over BLE notifications with flow control; (3) store-and-forward, application-level "
    "ACK, and idempotent receiver mechanisms; (4) background execution using foreground services; "
    "and (5) an empirical evaluation on physical devices that quantifies a concrete failure mode — "
    "marking send status ahead of receiver confirmation — together with the effect of removing it, "
    "supported by reproducible analysis tools.")

# --- §2.5 ----------------------------------------------------------------
replace_para(doc, "Testing was conducted on physical devices.",
    "Testing was conducted on physical devices. Four sessions on 23–28 June 2026 used the initial "
    "implementation, while one session on 5 July 2026 used the revised implementation that waits "
    "for an ACK before markSynced. Because the two campaigns ran different software, results are "
    "reported per version and are not summed into a single success figure.")

insert_after(find(doc, "Testing was conducted on physical devices."),
    "Data inclusion was handled as follows. The smartwatch produced 47,422 recorded rows across "
    "the five sessions. Of these, 22,231 rows in the tail of the June recording were excluded "
    "because they held frozen (constant) values with accuracy ≤ 0 while the smartwatch was off the "
    "wrist; the measurement period of each session ends at its last highest-accuracy reading. The "
    "remaining 25,191 readings constitute the measurement period used for every metric in Section "
    "3. One additional session on 23 June at 08:41 (53 readings, 1 minute) is excluded from the "
    "delivery aggregate because the smartphone never connected, so no delivery path was exercised; "
    "it is retained in the sensor-quality statistics, which do not depend on connectivity.")

t1 = doc.tables[0]
for row in t1.rows:
    if row.cells[0].text.strip().startswith("Test dataset"):
        row.cells[1].text = ("23 Jun–5 Jul 2026; 5 sessions (4 initial version, 1 ACK-revised); "
                             "25,191 measurement-period readings out of 47,422 recorded rows")

replace_para(doc, "The table summarises the controlled test environment",
    "The table summarises the controlled test environment: a consumer device pair (Samsung Galaxy "
    "Watch SM-R860 as the peripheral and Xiaomi Redmi Note 10 Pro as the central), the Flutter "
    "framework with the applications installed in release mode, a requested MTU of 512 bytes, and "
    "a selectable 3/5-minute sending interval. The dataset covers five sessions during 23 June–5 "
    "July 2026 with 25,191 measurement-period readings filtered from 47,422 raw recorded rows.")

# --- §3.2 ----------------------------------------------------------------
replace_para(doc, "The evaluation covers four initial-version sessions",
    "Delivery results are reported separately by software version in Table 2. The initial version "
    "was evaluated over three valid sessions totalling 386 minutes and recorded 20,455 of 23,162 "
    "records (88.31%). The revised version, which defers the synced flag until an ACK arrives, was "
    "evaluated over one 33-minute session and recorded 1,974 of 1,976 records (99.90%); the two "
    "remaining records were still unsent because they were recorded after the last scheduled "
    "flush, not because they were lost. All seven batches in that session were acknowledged. "
    "Neither version produced a single duplicate in the receiver database, and every matched "
    "record carried bpm and accuracy values identical to the smartwatch record (100% value "
    "fidelity).")

rebuild_table(doc.tables[1], TABEL2_EN)
replace_para(doc, "Table 2. Delivery results",
    "Table 2. Delivery results per session, grouped by software version.")

replace_para(doc, "The table shows delivery ranging",
    "The two subtotals are deliberately not combined into a single cross-version figure: the two "
    "groups ran different software, so any combined number would describe neither version. The "
    "session on 23 June at 08:41 (53 readings, 1 minute) is omitted from the table because the "
    "smartphone never connected during it, so no delivery path was exercised; it is reported as a "
    "failed session rather than as a 0% delivery measurement.")

insert_after(find(doc, "The figure plots the main-session heart-rate signal"),
    "Separating the results by version exposes a pattern that the aggregate figure conceals. Under "
    "the initial version, losses were not spread evenly across a session but clustered in the "
    "period before the smartphone connected: Session 3 lost 2,689 records, almost all of them "
    "before the connection was observed at approximately 15:00 (Figure 7), whereas Sessions 1 and "
    "4 — whose smartphones were connected from the start — lost only 13 and 5 records "
    "respectively. The 100% value fidelity across all received records reinforces this reading: "
    "not one record arrived corrupted, so the loss under the initial version was a failure to "
    "reschedule transmission rather than a failure of channel integrity. This is consistent with "
    "studies of BLE behaviour in heart-rate monitoring applications, which identify link "
    "conditions rather than payload corruption as the determinant of delivery quality [20], [25].",
    "The gap between 88.31% and 99.90% therefore cannot be read as an improvement in radio "
    "performance; it follows from changing a single software decision, namely when a record may be "
    "declared sent. This interpretation has limits. The two figures come from non-identical test "
    "conditions — the initial version was evaluated over far longer sessions that included "
    "unconnected periods, while the revised version was evaluated over 33 minutes of continuous "
    "connectivity — so a direct comparison is indicative rather than causal. Establishing "
    "causality requires an ablation experiment in which both versions run under the same "
    "disruption scenario with several replications, which is the next item on the testing agenda.")

replace_para(doc, "Figure 6. Data completeness",
    "Figure 6. Data completeness across the five evaluation sessions (received vs lost).")
replace_para(doc, "The chart compares the numbers of received and lost records",
    "The chart compares the numbers of received and lost records across all five sessions. Losses "
    "concentrate in the longest session (Session 3, 13.0% lost) and in the 23 June 08:41 session "
    "whose smartphone never connected, while the other three sessions are nearly complete (99.40%, "
    "98.23%, and 99.90%). The differing bar lengths also show that most of the data volume comes "
    "from a single long session, so an unweighted average across sessions would be misleading.")

# --- §3.3 ----------------------------------------------------------------
replace_para(doc, "The sensor-accuracy distribution during the measurement period",
    "The sensor-accuracy distribution across the measurement period of all five sessions is shown "
    "in Table 3 and Figure 8. 91.6% of readings (23,063) were at the highest accuracy (value 3), "
    "8.4% (2,113 readings) at the no-contact condition (value -1) — which commonly occurs when the "
    "sensor momentarily loses skin contact [11], [23] — and 0.1% (15 readings) at value 0; in "
    "Session 5 every reading had good contact. The heart-rate statistics for the highest-accuracy "
    "readings (n = 23,063, visualised in Figure 9) were: minimum 60 bpm, maximum 123 bpm, mean "
    "83.5 bpm, and standard deviation 9.6 bpm. Table 3, Figure 8, and Figure 9 all use the same "
    "scope — the measurement period of the five sessions — so their numbers are directly "
    "comparable. The 22,231 tail rows with frozen values and accuracy ≤ 0, recorded while the "
    "smartwatch was off the wrist, were excluded as described in Section 2.5.")

replace_para(doc, "Table 3. Distribution of sensor-accuracy values",
    "Table 3. Distribution of sensor-accuracy values (measurement period, five sessions).")

replace_para(doc, "The table shows that contact quality is dominated",
    "The table shows that contact quality is dominated by the highest level (91.6%), with 8.4% at "
    "no-contact and only 0.1% at the unreliable status. That composition is visualised in Figure "
    "8, while the distribution of heart-rate values for the highest-accuracy readings is shown in "
    "Figure 9.")

replace_para(doc, "Figure 8. Sensor contact-quality distribution",
    "Figure 8. Sensor contact-quality distribution across the five evaluation sessions.")
replace_para(doc, "The chart summarises the sensor contact quality",
    "The chart summarises the sensor contact quality across the five sessions: the majority of "
    "readings (91.6%) are at the highest accuracy level, 8.4% at no-contact, and 0.1% at the "
    "unreliable status — reflecting the momentary contact losses that are common for wrist-worn "
    "optical sensors.")

replace_para(doc, "The histogram shows the distribution of heart-rate values",
    "The histogram shows the distribution of heart-rate values for the highest-accuracy readings "
    "of all five sessions (n = 23,063): centred on a mean of 83.5 bpm with a standard deviation of "
    "9.6 bpm and a range of 60–123 bpm. This profile is reasonable for light activity and "
    "indicates physiologically plausible data rather than frozen values from an unworn sensor.")

insert_after(find(doc, "The histogram shows the distribution of heart-rate values"),
    "The 8.4% no-contact share is best read as a characteristic of wrist-worn optical sensing "
    "rather than as a defect of the delivery system. Device-validation studies report that "
    "wrist-worn heart-rate accuracy varies with activity condition and skin-contact quality [23], "
    "so a single-digit proportion of no-contact readings is unremarkable for everyday wear. What "
    "matters for this design is that the accuracy status is itself transmitted and stored intact "
    "at the receiver: quality filtering can be applied at analysis time without information loss, "
    "because the system does not discard low-accuracy readings at the source — an approach "
    "consistent with recommended practice for ambulatory wearable data [11]. The marked contrast "
    "between Session 5, in which every reading had good contact, and the June sessions also shows "
    "that contact quality is governed by how the device is worn rather than by the software, so "
    "this proportion cannot serve as a cross-session performance indicator.")

# --- §3.4 ----------------------------------------------------------------
replace_para(doc, "Transfer performance was measured through the HR-METRIC",
    "Transfer performance was measured through the HR-METRIC instrumentation on the smartwatch. In "
    "Session 5 (3-minute interval, seven batches), every batch was uniform — 180 readings, an "
    "8,461-byte payload (47.0 bytes per reading), and 19 frames at MTU 512 — with a mean transfer "
    "duration of 121.5 ± 25.5 ms (range 92.8–166.4 ms) and a mean throughput of 70.3 ± 13.2 KiB/s; "
    "every batch was fully acknowledged by the smartphone. The measured frame count matches the "
    "chunk-size calculation (MTU − 4 = 508 bytes): 8,461 ÷ 508 = 16.7, rounded up to 17 DATA "
    "frames, plus START and END = 19 frames. As an illustration of a large batch, in the June "
    "evaluation one batch of 228 readings (10,717-byte payload; 24 frames, consistent with the "
    "same calculation) was delivered in ~0.32 s, equivalent to ~32.4 KiB/s. MTU negotiation on the "
    "tested device pair consistently yielded the full 512 bytes in both test campaigns.")

insert_after(find(doc, "Transfer performance was measured through the HR-METRIC"),
    "In terms of channel occupancy, 121.5 ms against a 180-second sending interval means the radio "
    "is actively transmitting data for roughly 0.07% of the session. That ratio is what allows a "
    "periodic batch scheme to be combined with delivery acknowledgement without meaningful channel "
    "cost: acknowledgement happens seven times per session rather than once per reading. The "
    "distinction matters, because adding strict acknowledgement to a continuous real-time stream "
    "is known to aggravate congestion and degrade delay and loss when several patients are "
    "monitored [18], and reliability–energy trade-offs of this kind are a recurring theme in body "
    "area networks [22]. In other words, a delivery delay of a few minutes — acceptable for "
    "heart-rate trend monitoring, though not for real-time alerting — is the price paid for data "
    "completeness at low channel cost.",
    "It should be noted that throughput differed considerably between the two campaigns: 70.3 "
    "KiB/s in Session 5 against approximately 32.4 KiB/s for the June batch, even though the "
    "negotiated MTU was 512 bytes in both. Because BLE connection parameters such as the "
    "connection interval were not logged during the June campaign, this difference cannot yet be "
    "attributed to a single cause; connection parameters have been reported to affect delivery "
    "quality in heart-rate monitoring applications [20] and BLE performance generally [8]. Logging "
    "connection parameters has therefore been added to the instrumentation for subsequent testing. "
    "Measurements across other conditions — the 5-minute interval, device distance, and disruption "
    "scenarios — likewise remain future work.")

# --- §3.5 ----------------------------------------------------------------
replace_para(doc, "A cross-analysis of the four June sessions",
    "A cross-analysis of the three valid initial-version sessions found that of the 2,707 records "
    "absent from the smartphone, 2,702 (99.8%) were already marked synced = 1 and only 5 were "
    "legitimately pending. In other words, almost all of the missing data was not waiting its turn "
    "to be sent; the smartwatch had already considered it finished. Under the revised version, "
    "Session 5 produced no falsely flagged records at all: seven batches each received an ACK for "
    "180 records, and the final two records correctly remained at synced = 0. Session 5 did not, "
    "however, include a disconnect–reconnect scenario, so the revised version's resilience to link "
    "loss still requires dedicated testing.")

replace_para(doc, "Across the full June recording",
    "Across the full June recording (45,446 rows including the not-worn period), 2,879 records "
    "marked synced = 1 were absent from the smartphone. The figure of 2,702 above is the subset "
    "falling within the measurement period of the three valid sessions and is therefore consistent "
    "with Table 2; the remainder comes from the off-wrist tail and from the 08:41 session excluded "
    "from the aggregate.")

insert_after(find(doc, "Across the full June recording"),
    "The root cause lies in the ordering of operations. In the initial version, the synced = 1 "
    "flag was written as soon as the send call returned from the BLE layer, rather than after the "
    "receiver confirmed persistence. Because BLE notifications are best-effort and provide no "
    "attribute-layer acknowledgement [9], a successful send call means only that the frame was "
    "handed to the local protocol stack — not that it arrived and was stored. Records that were "
    "never delivered were consequently treated as complete and never re-entered the send queue, so "
    "the store-and-forward buffer lost its function precisely when it was most needed, that is, "
    "while the link was down.",
    "The lesson generalises to any buffer design driven by a status flag: a sent flag may be "
    "advanced only by an event originating at the receiver, never by a local event at the sender. "
    "The same principle underlies store-carry-forward schemes in delay-tolerant networks, where a "
    "message may be released from the buffer only after the next hop is confirmed [21], and it "
    "mirrors retransmission handling after network disconnection in other protocol stacks [24]. "
    "The practical value of the finding is that this failure is invisible to ordinary functional "
    "testing: the system appears to run normally, the user interface reports all data as sent, and "
    "the loss surfaces only when the two databases are compared record by record. Completeness "
    "verification through timestamp matching, as used here, should therefore be a standard part of "
    "testing continuous monitoring systems rather than a supplementary analysis.")

# --- §3.6 ----------------------------------------------------------------
replace_para(doc, "The results show that opcode framing and flow control",
    "Taken together, the results show that the protocol design — opcode framing with flow control, "
    "an idempotent receiver, and application-level acknowledgement — behaves as intended on real "
    "devices: every received record had 100% value fidelity with no duplicates under either "
    "version, and batch transfer completed two to three orders of magnitude faster than its "
    "sending interval. The most valuable empirical contribution, however, comes from the failure "
    "of the initial version: it demonstrates that the presence of a store-and-forward buffer is "
    "not by itself sufficient, and that the correctness of the mechanism hinges on one "
    "easily-overlooked implementation detail — which event is entitled to advance the send-status "
    "flag.")

insert_after(find(doc, "Taken together, the results show that the protocol design"),
    "Relative to prior work, this study occupies a distinct position. Protocol-level studies "
    "examine how connection parameters affect BLE delivery quality [20] and model its performance "
    "analytically [8], while surveys of wireless body area networks summarise reliability and "
    "energy-efficiency approaches at the network-architecture level [22]; none couples this with "
    "an end-to-end implementation on a consumer device pair verified record by record. Conversely, "
    "store-carry-forward schemes in delay-tolerant networking [21] supply a comparable buffering "
    "idea, but applied to routing between nodes rather than to a single wearable–gateway link "
    "constrained by MTU and by mobile-OS background-execution limits [10], [26]. Direct-"
    "transmission wearable architectures [14], [15] optimise for real-time delivery and therefore "
    "do not address backlog recovery at all. The contribution here is to bring the buffering idea "
    "down to the application level on real devices, together with quantitative evidence of a "
    "failure mode that neither protocol-level analysis nor network simulation surfaces.")

replace_para(doc, "This study has several limitations.",
    "This study has several limitations that bound the generality of its results. First, testing "
    "used a single device pair in a single radio environment, so the latency and throughput figures "
    "cannot be generalised across devices. Second, the revised version was evaluated in only one "
    "session, without replication and without a controlled disconnect–reconnect scenario, so its "
    "resilience claim is confined to the continuously connected condition actually tested. Third, "
    "the current ACK carries no batchId and does not validate the acknowledged record count against "
    "the pending batch, so the mechanism improves reliability without providing a formal "
    "failure-free guarantee. Fourth, no controlled variation of sending interval, distance, or MTU "
    "was performed, and energy consumption was not measured. Fifth, background execution does not "
    "guarantee continuation after force-close or reboot. Sixth, the communication adds no "
    "application-level authentication or encryption. Seventh, the consumer smartwatch sensor was "
    "not assessed as a clinical device; the contribution of this work concerns delivery "
    "reliability, not the medical validity of the BPM values.")

# --- conclusion ----------------------------------------------------------
replace_para(doc, "This study designed and implemented a BLE heart-rate delivery system",
    "This study designed and implemented a BLE heart-rate delivery system using opcode batch "
    "framing, local persistence, application-level ACKs, an idempotent receiver, and foreground "
    "services, and evaluated it on a consumer device pair across five sessions. The initial "
    "version achieved 88.31% delivery and revealed that 2,702 of its 2,707 lost records had "
    "already been falsely marked as sent and were therefore never rescheduled; the revised version, "
    "which defers that marking until an ACK arrives, reached 99.90% with no falsely marked "
    "records. Under both versions value fidelity was 100% with no duplicates, and seven 180-record "
    "batches at MTU 512 transferred in 121.5 ms on average, or 0.07% of their sending interval. "
    "The central finding is that a send-status flag must be advanced by receiver confirmation "
    "rather than by the success of a local send call. Future work should replicate the revised "
    "version under controlled disconnect–reconnect scenarios, add batch identity to the ACK, "
    "measure energy consumption, and extend testing across devices and radio conditions.")

ack = find(doc, "Acknowledgement")
head = copy.deepcopy(ack._p)
ack._p.addprevious(head)
hp = Paragraph(head, ack._parent)
set_text(hp, "Data and Code Availability")
insert_after(hp,
    "The source code of both applications, the analysis scripts, and the exported CSV files of all "
    "five test sessions are available in the public repository of this study at [[repository URL]], "
    "so that every figure in Table 2, Table 3, and Figures 6–9 can be reproduced.")

replace_para(doc, "Note: this is a draft.",
    "Note: this is a draft. Adapt it to the target journal template (column format, citation style, "
    "page limit). References are numbered in IEEE order of first citation; reference [19] is an "
    "SSRN preprint whose year and venue must be completed. Items marked [[...]] must be completed "
    "by the author, including the repository URL in the Data and Code Availability section.")

doc.save(str(ROOT / "Draft_Manuscript_HR_BLE_EN_v2.docx"))
print("OK — Draft_Manuscript_HR_BLE_EN_v2.docx")
