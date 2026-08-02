#!/usr/bin/env python3
"""Perbarui kedua manuskrip dengan hasil pengujian 5 Juli 2026 (Sesi 5):
- Agregat lima sesi: 25.191 direkam, 22.429 diterima, delivery 89,04%.
- Tabel hasil: baris Sesi 5 (1.976/1.974/2/99,90%) + Total baru.
- §3.4 ditulis ulang: statistik 7 batch (180 record/8.461 B/19 frame; durasi
  121,5±25,5 ms; throughput 70,3±13,2 KB/s; semua ber-ACK), validasi rumus
  chunk = MTU−4, MTU 512 konsisten; placeholder replikasi DIHAPUS karena
  replikasinya kini ada.
- §3.5: dekomposisi lima sesi (2.762 = 2.755 false-sent + 7 pending) +
  konfirmasi Sesi 5 (false-sent = 0 saat tersambung penuh).
- §3.3/Tabel 3: distribusi akurasi gabungan (23.063/15/2.113; 91,6/0,1/8,4%),
  statistik bpm gabungan (n=23.063; 83,5; SD 9,6).
- Abstrak, §2.5/Tabel 1, §3.6, dan Kesimpulan disesuaikan.
- Figur 6–9 tetap menampilkan evaluasi Juni; teks penjelasnya diberi cakupan.
Angka bersumber dari analisis 2026-07-05_*.csv + watch_log (7 baris HR-METRIC).
"""
from docx import Document
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]

def set_text(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        p.add_run(text)

def find(doc, prefix):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    raise SystemExit(f"PARAGRAF TIDAK DITEMUKAN: {prefix!r}")

def replace_in_paragraphs(doc, replacements):
    done = {old: False for old, _ in replacements}
    for p in doc.paragraphs:
        t = p.text
        hit = False
        for old, new in replacements:
            if old in t:
                t = t.replace(old, new); done[old] = True; hit = True
        if hit:
            set_text(p, t)
    missing = [k for k, v in done.items() if not v]
    if missing:
        raise SystemExit(f"TEKS TIDAK DITEMUKAN: {missing}")

def replace_in_tables(doc, replacements):
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    t = p.text
                    hit = False
                    for old, new in replacements:
                        if old in t:
                            t = t.replace(old, new); hit = True
                    if hit:
                        set_text(p, t)

def insert_row_before_total(tbl, values):
    row = tbl.add_row()
    for i, v in enumerate(values):
        row.cells[i].text = v
    total_tr = tbl.rows[-2]._tr          # baris Total (sebelum baris baru)
    total_tr.addprevious(tbl.rows[-1]._tr)

# ============================ NASKAH ID ============================
doc = Document(str(ROOT / "Draft_Naskah_HR_BLE_SINTA2.docx"))

replace_in_paragraphs(doc, [
    # --- Abstrak ID ---
    ("Evaluasi pada perangkat fisik mencakup empat sesi (total 23.215 pembacaan pengukuran, di luar "
     "22.231 sampel saat sensor tidak terpasang).",
     "Evaluasi pada perangkat fisik mencakup lima sesi (total 25.191 pembacaan pengukuran, di luar "
     "22.231 sampel saat sensor tidak terpasang)."),
    ("sehingga rasio keberhasilan agregat lintas seluruh sesi sebesar 88,1%.",
     "sehingga rasio keberhasilan agregat lintas seluruh sesi sebesar 89,0%."),
    ("Sebanyak 90,8% pembacaan berada pada akurasi sensor tertinggi.",
     "Sebanyak 91,6% pembacaan berada pada akurasi sensor tertinggi."),
    ("Satu paket berisi 228 pembacaan (10.717 byte, 24 frame) terkirim dalam ±0,32 detik pada MTU 512.",
     "Pada MTU 512, batch tipikal berisi 180 pembacaan (8.461 byte, 19 frame) terkirim rata-rata dalam "
     "121,5 milidetik (tujuh batch), dan batch terbesar berisi 228 pembacaan (10.717 byte, 24 frame) "
     "terkirim dalam ±0,32 detik."),
    # --- Abstract EN (di dokumen ID) ---
    ("Evaluation on physical devices spanned four sessions (23,215 measurement readings in total, "
     "excluding 22,231 samples recorded while the sensor was not worn).",
     "Evaluation on physical devices spanned five sessions (25,191 measurement readings in total, "
     "excluding 22,231 samples recorded while the sensor was not worn)."),
    ("yielding an aggregate delivery ratio of 88.1% across all sessions.",
     "yielding an aggregate delivery ratio of 89.0% across all sessions."),
    ("90.8% of readings were at the highest sensor-accuracy level.",
     "91.6% of readings were at the highest sensor-accuracy level."),
    ("One batch of 228 readings (10,717 bytes, 24 frames) was delivered in ~0.32 s at MTU 512.",
     "At MTU 512, a typical batch of 180 readings (8,461 bytes, 19 frames) was delivered in 121.5 ms "
     "on average (seven batches), and the largest batch of 228 readings (10,717 bytes, 24 frames) in "
     "~0.32 s."),
    # --- §2.5 ---
    ("Pengujian dijalankan dalam empat sesi pada rentang 23–28 Juni 2026; pada tiap run",
     "Pengujian dijalankan dalam lima sesi pada rentang 23 Juni–5 Juli 2026; pada tiap run"),
    ("Total terkumpul 23.215 pembacaan pada periode pengukuran (sensor terpasang).",
     "Total terkumpul 25.191 pembacaan pada periode pengukuran (sensor terpasang)."),
    ("Dataset akhir mencakup empat sesi pada 23–28 Juni 2026 dengan total 23.215 pembacaan pengukuran.",
     "Dataset akhir mencakup lima sesi pada 23 Juni–5 Juli 2026 dengan total 25.191 pembacaan "
     "pengukuran."),
    # --- §3.2 ---
    ("Evaluasi dilakukan pada empat sesi pengujian dengan hasil per sesi pada Tabel 2 dan profil "
     "temporalnya pada Gambar 6 dan Gambar 7.",
     "Evaluasi dilakukan pada lima sesi pengujian — empat sesi pada 23–28 Juni 2026 dan satu sesi "
     "tambahan pada 5 Juli 2026 — dengan hasil per sesi pada Tabel 2 dan profil temporal empat sesi "
     "Juni pada Gambar 6 dan Gambar 7."),
    ("Secara agregat, smartwatch merekam 23.215 pembacaan pada periode pengukuran dan 20.455 di "
     "antaranya tercatat di smartphone (rasio keberhasilan 88,1%) tanpa duplikat.",
     "Secara agregat lima sesi, smartwatch merekam 25.191 pembacaan pada periode pengukuran dan 22.429 "
     "di antaranya tercatat di smartphone (rasio keberhasilan 89,0%) tanpa duplikat."),
    ("Dari tabel terlihat rentang delivery antar-sesi sangat lebar (0,00–99,40%): dua sesi dengan "
     "koneksi penuh hampir lengkap, Sesi 3 (terpanjang) turun ke 87,01% akibat periode awal tanpa "
     "koneksi, dan Sesi 2 kehilangan seluruh 53 pembacaannya karena smartphone tidak pernah terhubung. "
     "Profil temporal kehilangan tersebut divisualkan pada Gambar 6 dan Gambar 7.",
     "Dari tabel terlihat rentang delivery antar-sesi sangat lebar (0,00–99,90%): sesi-sesi dengan "
     "koneksi penuh hampir lengkap — termasuk Sesi 5 (99,90%) yang hanya menyisakan dua pembacaan "
     "terakhir berstatus menunggu kirim — Sesi 3 (terpanjang) turun ke 87,01% akibat periode awal "
     "tanpa koneksi, dan Sesi 2 kehilangan seluruh 53 pembacaannya karena smartphone tidak pernah "
     "terhubung. Profil temporal kehilangan empat sesi Juni divisualkan pada Gambar 6 dan Gambar 7."),
    ("Diagram tersebut membandingkan jumlah record yang diterima dan yang hilang pada tiap sesi.",
     "Diagram tersebut membandingkan jumlah record yang diterima dan yang hilang pada empat sesi "
     "evaluasi Juni (Sesi 5 tercantum pada Tabel 2)."),
    # --- §3.3 ---
    ("Sebanyak 90,8% pembacaan berada pada akurasi tertinggi (nilai 3), 9,1% (2.113 pembacaan)",
     "Sebanyak 91,6% pembacaan (23.063 pembacaan) berada pada akurasi tertinggi (nilai 3), 8,4% "
     "(2.113 pembacaan)"),
    ("dan 0,1% (15 pembacaan) pada nilai 0.",
     "dan 0,1% (15 pembacaan) pada nilai 0; pada Sesi 5 seluruh pembacaan berkontak baik."),
    ("(n=21.087, Gambar 9) adalah minimum 60 bpm, maksimum 123 bpm, rata-rata 83,4 bpm, dan simpangan "
     "baku 9,9 bpm.",
     "(n=23.063; sebaran empat sesi Juni divisualkan pada Gambar 9) adalah minimum 60 bpm, maksimum "
     "123 bpm, rata-rata 83,5 bpm, dan simpangan baku 9,6 bpm."),
    ("Dari tabel terlihat kualitas kontak didominasi tingkat tertinggi (90,8%), dengan porsi tanpa "
     "kontak 9,1% dan tingkat sedang hanya 0,1%.",
     "Dari tabel terlihat kualitas kontak didominasi tingkat tertinggi (91,6%), dengan porsi tanpa "
     "kontak 8,4% dan tingkat sedang hanya 0,1%."),
    ("Komposisi kualitas kontak tersebut serta sebaran nilai detak jantung pada pembacaan terbaik "
     "divisualkan pada Gambar 8 dan Gambar 9.",
     "Komposisi kualitas kontak tersebut serta sebaran nilai detak jantung pada pembacaan terbaik "
     "(empat sesi Juni) divisualkan pada Gambar 8 dan Gambar 9."),
    ("Diagram tersebut merangkum kualitas kontak sensor selama periode pengukuran: mayoritas pembacaan "
     "(90,8%)",
     "Diagram tersebut merangkum kualitas kontak sensor pada empat sesi evaluasi Juni: mayoritas "
     "pembacaan (90,8%)"),
    ("Histogram tersebut menunjukkan sebaran nilai detak jantung pada pembacaan akurasi tertinggi "
     "(n = 21.087)",
     "Histogram tersebut menunjukkan sebaran nilai detak jantung pada pembacaan akurasi tertinggi "
     "empat sesi Juni (n = 21.087)"),
    # --- §3.5 ---
    ("Dari 2.760 pembacaan yang hilang, 2.755 di antaranya (99,8%)",
     "Dari 2.762 pembacaan yang hilang pada kelima sesi, 2.755 di antaranya (99,7%)"),
    ("hanya 5 pembacaan terakhir yang memang masih berstatus belum terkirim (menunggu interval "
     "berikutnya)",
     "hanya 7 pembacaan yang memang masih berstatus belum terkirim (menunggu interval berikutnya: 5 di "
     "akhir Sesi 4 dan 2 di akhir Sesi 5)"),
    ("ditambah 5 record tertunda yang wajar",
     "ditambah 7 record tertunda yang wajar"),
    ("bukti paling jelas bahwa penandaan dilakukan prematur.",
     "bukti paling jelas bahwa penandaan dilakukan prematur. Sebaliknya, Sesi 5 yang tersambung penuh "
     "tidak menghasilkan satu pun penandaan keliru (false-sent = 0) dan seluruh tujuh batch-nya "
     "dikonfirmasi ACK, menegaskan bahwa penandaan prematur khusus terjadi pada kondisi koneksi "
     "terputus."),
    ("Bila dihitung atas seluruh rekaman smartwatch (45.446 baris",
     "Bila dihitung atas seluruh rekaman smartwatch pada evaluasi Juni (45.446 baris"),
    # --- §3.6 ---
    ("Pertama, evaluasi mencakup empat sesi pada satu pasang perangkat",
     "Pertama, evaluasi mencakup lima sesi pada satu pasang perangkat"),
    # --- Kesimpulan ---
    ("Pengujian pada perangkat fisik lintas empat sesi (total 23.215 pembacaan)",
     "Pengujian pada perangkat fisik lintas lima sesi (total 25.191 pembacaan)"),
    ("dengan 90,8% data berakurasi tertinggi serta transfer batch yang cepat (±0,32 detik untuk 228 "
     "pembacaan pada MTU 512)",
     "dengan 91,6% data berakurasi tertinggi serta transfer batch yang cepat (rata-rata 121,5 ms untuk "
     "180 pembacaan; ±0,32 detik untuk 228 pembacaan pada MTU 512)"),
    ("sehingga rasio keberhasilan agregat menjadi 88,1%",
     "sehingga rasio keberhasilan agregat menjadi 89,0%"),
])

# --- §3.4 ditulis ulang penuh ---
set_text(find(doc, "Sebagai ilustrasi kinerja per batch"),
    "Kinerja transfer diukur melalui instrumentasi HR-METRIC pada smartwatch. Pada Sesi 5 (interval 3 "
    "menit, tujuh batch), seluruh batch berukuran seragam — 180 pembacaan, payload 8.461 byte (47,0 "
    "byte per pembacaan), dan 19 frame pada MTU 512 — dengan durasi transfer rata-rata 121,5 ± 25,5 ms "
    "(rentang 92,8–166,4 ms) dan throughput rata-rata 70,3 ± 13,2 KB/detik; setiap batch dikonfirmasi "
    "ACK penuh oleh smartphone. Jumlah frame terukur sesuai perhitungan ukuran chunk (MTU − 4 = 508 "
    "byte): 8.461 ÷ 508 = 16,7, dibulatkan ke atas menjadi 17 frame DATA, ditambah START dan END "
    "menjadi 19 frame. Sebagai ilustrasi batch besar, pada evaluasi Juni satu batch berisi 228 "
    "pembacaan (payload 10.717 byte; 24 frame, konsisten dengan perhitungan yang sama) terkirim dalam "
    "±0,32 detik pada MTU 512; di sisi penerima, perakitan ulang frame memerlukan ±250 ms dan "
    "penyimpanan batch ke basis data ±56 ms. Negosiasi MTU pada pasangan perangkat uji secara "
    "konsisten menghasilkan nilai penuh 512 byte pada kedua rangkaian pengujian. Nilai-nilai ini "
    "menunjukkan transfer batch berlangsung cepat relatif terhadap interval pengiriman (menit). "
    "Pengukuran lintas variasi kondisi lain (interval 5 menit, jarak antar-perangkat, dan skenario "
    "gangguan) merupakan bagian dari pekerjaan lanjutan.")

# --- tabel: Tabel 1 (dataset), Tabel 2 (+Sesi 5, Total), Tabel 3 (akurasi) ---
replace_in_tables(doc, [
    ("23–28 Jun 2026; 4 sesi; 23.215 pembacaan pengukuran",
     "23 Jun–5 Jul 2026; 5 sesi; 25.191 pembacaan pengukuran"),
    ("387 mnt", "420 mnt"), ("23.215", "25.191"), ("20.455", "22.429"),
    ("2.760", "2.762"), ("88,11%", "89,04%"),
    ("21.087", "23.063"), ("90,8%", "91,6%"), ("9,1%", "8,4%"),
])
insert_row_before_total(doc.tables[1],
    ["Sesi 5", "05/07 23:07", "33 mnt", "1.976", "1.974", "2", "99,90%"])

doc.save(str(ROOT / "Draft_Naskah_HR_BLE_SINTA2.docx"))
print("OK -> Draft_Naskah_HR_BLE_SINTA2.docx")

# ============================ MANUSCRIPT EN ============================
doc = Document(str(ROOT / "Draft_Manuscript_HR_BLE_EN.docx"))

replace_in_paragraphs(doc, [
    ("Evaluation on physical devices spanned four sessions (23,215 measurement readings in total, "
     "excluding 22,231 samples recorded while the sensor was not worn).",
     "Evaluation on physical devices spanned five sessions (25,191 measurement readings in total, "
     "excluding 22,231 samples recorded while the sensor was not worn)."),
    ("yielding an aggregate delivery ratio of 88.1% across all sessions.",
     "yielding an aggregate delivery ratio of 89.0% across all sessions."),
    ("90.8% of readings were at the highest sensor-accuracy level.",
     "91.6% of readings were at the highest sensor-accuracy level."),
    ("One batch of 228 readings (10,717 bytes, 24 frames) was delivered in ~0.32 s at MTU 512.",
     "At MTU 512, a typical batch of 180 readings (8,461 bytes, 19 frames) was delivered in 121.5 ms "
     "on average (seven batches), and the largest batch of 228 readings (10,717 bytes, 24 frames) in "
     "~0.32 s."),
    ("Testing was carried out over four sessions during 23–28 June 2026; in each run",
     "Testing was carried out over five sessions during 23 June–5 July 2026; in each run"),
    ("A total of 23,215 readings were collected during the measurement period (sensor worn).",
     "A total of 25,191 readings were collected during the measurement period (sensor worn)."),
    ("The final dataset covers four sessions during 23–28 June 2026 with 23,215 measurement readings "
     "in total.",
     "The final dataset covers five sessions during 23 June–5 July 2026 with 25,191 measurement "
     "readings in total."),
    ("The evaluation was carried out over four test sessions, with per-session results in Table 2 and "
     "their temporal profiles in Figure 6 and Figure 7.",
     "The evaluation was carried out over five test sessions — four during 23–28 June 2026 and one "
     "additional session on 5 July 2026 — with per-session results in Table 2 and the temporal "
     "profiles of the four June sessions in Figure 6 and Figure 7."),
    ("In aggregate, the smartwatch recorded 23,215 readings during the measurement period and 20,455 "
     "of them were recorded on the smartphone (88.1% delivery ratio) with no duplicates.",
     "In aggregate across the five sessions, the smartwatch recorded 25,191 readings during the "
     "measurement period and 22,429 of them were recorded on the smartphone (89.0% delivery ratio) "
     "with no duplicates."),
    ("The table shows a wide range of per-session delivery (0.00–99.40%): the two fully connected "
     "sessions are nearly complete, Session 3 (the longest) drops to 87.01% due to its initial "
     "disconnected period, and Session 2 lost all of its 53 readings because the smartphone never "
     "connected. The temporal profile of these losses is visualised in Figure 6 and Figure 7.",
     "The table shows a wide range of per-session delivery (0.00–99.90%): the fully connected "
     "sessions are nearly complete — including Session 5 (99.90%), whose only two missing readings "
     "were the final two seconds still awaiting the next send — Session 3 (the longest) drops to "
     "87.01% due to its initial disconnected period, and Session 2 lost all of its 53 readings "
     "because the smartphone never connected. The temporal profile of the June losses is visualised "
     "in Figure 6 and Figure 7."),
    ("The chart compares the numbers of received and lost records in each session.",
     "The chart compares the numbers of received and lost records in the four June sessions "
     "(Session 5 is listed in Table 2)."),
    ("90.8% of readings were at the highest accuracy (value 3), 9.1% (2,113 readings)",
     "91.6% of readings (23,063) were at the highest accuracy (value 3), 8.4% (2,113 readings)"),
    ("and 0.1% (15 readings) at value 0.",
     "and 0.1% (15 readings) at value 0; in Session 5 every reading had good contact."),
    ("(n=21,087, Figure 9) were: minimum 60 bpm, maximum 123 bpm, mean 83.4 bpm, and standard "
     "deviation 9.9 bpm.",
     "(n=23,063; the June distribution is visualised in Figure 9) were: minimum 60 bpm, maximum 123 "
     "bpm, mean 83.5 bpm, and standard deviation 9.6 bpm."),
    ("The table shows that contact quality is dominated by the highest level (90.8%), with 9.1% at "
     "no-contact and only 0.1% at the medium level.",
     "The table shows that contact quality is dominated by the highest level (91.6%), with 8.4% at "
     "no-contact and only 0.1% at the medium level."),
    ("The composition of contact quality and the distribution of heart-rate values for the best "
     "readings are visualised in Figure 8 and Figure 9.",
     "The composition of contact quality and the distribution of heart-rate values for the best "
     "readings (four June sessions) are visualised in Figure 8 and Figure 9."),
    ("The chart summarises the sensor contact quality during the measurement period: the majority of "
     "readings (90.8%)",
     "The chart summarises the sensor contact quality in the four June sessions: the majority of "
     "readings (90.8%)"),
    ("The histogram shows the distribution of heart-rate values for the highest-accuracy readings "
     "(n = 21,087)",
     "The histogram shows the distribution of heart-rate values for the highest-accuracy readings of "
     "the June sessions (n = 21,087)"),
    ("Of the 2,760 lost readings, 2,755 (99.8%)",
     "Of the 2,762 readings lost across the five sessions, 2,755 (99.7%)"),
    ("only the last 5 readings were genuinely still pending (awaiting the next interval)",
     "only 7 readings were genuinely still pending (awaiting the next interval: 5 at the end of "
     "Session 4 and 2 at the end of Session 5)"),
    ("plus 5 legitimately pending ones",
     "plus 7 legitimately pending ones"),
    ("the clearest evidence that the marking is applied prematurely.",
     "the clearest evidence that the marking is applied prematurely. Conversely, the fully connected "
     "Session 5 produced no false marking at all (zero false-sent) and all seven of its batches were "
     "acknowledged, confirming that premature marking occurs specifically under disconnected "
     "conditions."),
    ("When computed over the entire smartwatch recording (45,446 rows",
     "When computed over the entire smartwatch recording of the June evaluation (45,446 rows"),
    ("First, the evaluation spans four sessions on a single device pair",
     "First, the evaluation spans five sessions on a single device pair"),
    ("Testing on physical devices across four sessions (23,215 readings in total)",
     "Testing on physical devices across five sessions (25,191 readings in total)"),
    ("with 90.8% of data at the highest accuracy and fast batch transfer (~0.32 s for 228 readings at "
     "MTU 512)",
     "with 91.6% of data at the highest accuracy and fast batch transfer (121.5 ms on average for 180 "
     "readings; ~0.32 s for 228 readings at MTU 512)"),
    ("so the aggregate delivery ratio was 88.1%",
     "so the aggregate delivery ratio was 89.0%"),
])

set_text(find(doc, "As an illustration of per-batch performance"),
    "Transfer performance was measured through the HR-METRIC instrumentation on the smartwatch. In "
    "Session 5 (3-minute interval, seven batches), every batch was uniform — 180 readings, an "
    "8,461-byte payload (47.0 bytes per reading), and 19 frames at MTU 512 — with a mean transfer "
    "duration of 121.5 ± 25.5 ms (range 92.8–166.4 ms) and a mean throughput of 70.3 ± 13.2 KB/s; "
    "every batch was fully acknowledged by the smartphone. The measured frame count matches the "
    "chunk-size calculation (MTU − 4 = 508 bytes): 8,461 ÷ 508 = 16.7, rounded up to 17 DATA frames, "
    "plus START and END = 19 frames. As an illustration of a large batch, in the June evaluation one "
    "batch of 228 readings (10,717-byte payload; 24 frames, consistent with the same calculation) was "
    "delivered in ~0.32 s at MTU 512; on the receiver side, frame reassembly took ~250 ms and storing "
    "the batch to the database took ~56 ms. MTU negotiation on the tested device pair consistently "
    "yielded the full 512 bytes in both test campaigns. These values show that batch transfer is fast "
    "relative to the sending interval (minutes). Measurements across other conditions (the 5-minute "
    "interval, device distance, and disruption scenarios) remain future work.")

replace_in_tables(doc, [
    ("23–28 Jun 2026; 4 sessions; 23,215 measurement readings",
     "23 Jun–5 Jul 2026; 5 sessions; 25,191 measurement readings"),
    ("387 min", "420 min"), ("23,215", "25,191"), ("20,455", "22,429"),
    ("2,760", "2,762"), ("88.11%", "89.04%"),
    ("21,087", "23,063"), ("90.8%", "91.6%"), ("9.1%", "8.4%"),
])
insert_row_before_total(doc.tables[1],
    ["Session 5", "05/07 23:07", "33 min", "1,976", "1,974", "2", "99.90%"])

doc.save(str(ROOT / "Draft_Manuscript_HR_BLE_EN.docx"))
print("OK -> Draft_Manuscript_HR_BLE_EN.docx")
