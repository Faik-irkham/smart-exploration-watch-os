#!/usr/bin/env python3
"""Tambahkan paragraf penjelasan setelah SETIAP gambar (caption) di kedua
manuskrip, melengkapi pola pengantar -> gambar -> penjelasan. Caption tetap;
paragraf baru bergaya teks isi (tidak miring/tengah)."""
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]

def insert_after(p, text):
    new_p = OxmlElement("w:p")
    p._p.addnext(new_p)
    np = Paragraph(new_p, p._parent)
    np.add_run(text)
    return np

def add_explanations(path, expl):
    doc = Document(str(path))
    found = set()
    for p in list(doc.paragraphs):
        t = p.text.strip()
        for prefix, text in expl.items():
            if t.startswith(prefix) and prefix not in found:
                insert_after(p, text)
                found.add(prefix)
    missing = set(expl) - found
    if missing:
        raise SystemExit(f"CAPTION TIDAK DITEMUKAN di {path.name}: {sorted(missing)}")
    doc.save(str(path))
    print(f"OK -> {path.name} (+{len(found)} penjelasan)")

# ============================ NASKAH ID ============================
expl_id = {
"Gambar 1.":
    "Pada gambar tersebut, kolom kiri menunjukkan lima komponen smartwatch — sensor, kode native, "
    "lapisan aplikasi, SQLite, dan BLE advertiser/GATT server — yang mengalir dari akuisisi ke "
    "pengiriman (langkah 1–5), sedangkan kolom kanan menunjukkan lima komponen smartphone dari "
    "penerimaan BLE hingga ekspor data (langkah 6–10). Jembatan pada baris bawah menggambarkan dua arah "
    "komunikasi: batch data melalui NOTIFY dari smartwatch ke smartphone, dan konfirmasi ACK melalui "
    "WRITE pada arah sebaliknya.",
"Gambar 2.":
    "Gambar tersebut memperlihatkan satu batch (JSON array berisi bpm, accuracy, dan time) dipecah "
    "menjadi rangkaian frame: diawali START, diikuti sejumlah frame DATA yang masing-masing membawa "
    "satu potongan payload, dan ditutup END. Inset di bagian bawah menunjukkan struktur satu frame "
    "DATA — satu byte opcode diikuti potongan data sepanjang maksimum MTU − 4 byte — dan tiap frame "
    "dikirim sebagai satu notifikasi BLE.",
"Gambar 3.":
    "Diagram urutan tersebut dibaca dari atas ke bawah dalam tiga fase. Fase penyiapan koneksi: "
    "smartwatch beriklan, smartphone memindai lalu terhubung, meminta MTU 512, dan mengaktifkan "
    "notifikasi. Fase transfer: batch dikirim sebagai rangkaian START–DATA–END. Fase persistensi dan "
    "konfirmasi: smartphone menyimpan batch ke basis data lalu menulis ACK, yang membuat smartwatch "
    "menandai record sebagai terkirim.",
"Gambar 4.":
    "Alur pada gambar tersebut memuat dua skala waktu yang berbeda: pembacaan sensor disimpan setiap "
    "satu detik (kontinu), sedangkan pengambilan record belum terkirim, pengiriman batch, dan penantian "
    "ACK berjalan per interval (3/5 menit). Cabang keputusan ACK memperlihatkan inti mekanisme "
    "keandalannya: bila ACK diterima, record ditandai terkirim; bila tidak, record tetap berstatus "
    "belum terkirim dan dikirim ulang pada interval berikutnya.",
"Gambar 5.":
    "Gambar tersebut memisahkan dua lapisan aplikasi smartwatch. Lapisan Dart memuat antarmuka, logika "
    "pemantauan dengan pola BLoC (pencuplikan per detik, pengiriman per interval, penantian ACK), dan "
    "akses SQLite; lapisan native Kotlin memuat pendengar sensor, GATT server dengan antrean frame dan "
    "kendali aliran, karakteristik ACK, serta foreground service. Kedua lapisan berkomunikasi melalui "
    "platform channel: perintah mengalir turun melalui MethodChannel dan peristiwa naik melalui "
    "EventChannel.",
"Gambar 6.":
    "Diagram tersebut membandingkan jumlah record yang diterima dan yang hilang pada tiap sesi. "
    "Terlihat kehilangan terkonsentrasi pada sesi terpanjang (Sesi 3) dan pada sesi yang "
    "smartphone-nya tidak pernah terhubung (Sesi 2), sedangkan dua sesi lainnya hampir lengkap "
    "(masing-masing 99,40% dan 98,23%).",
"Gambar 7.":
    "Gambar tersebut menampilkan sinyal detak jantung selama sesi utama beserta penanda record yang "
    "hilang. Seluruh kehilangan berada sebelum smartphone terhubung sekitar pukul 15.00; setelah tautan "
    "aktif, kehilangan praktis nol. Pola ini menegaskan bahwa kehilangan tidak disebabkan oleh kanal "
    "BLE, melainkan oleh periode tanpa koneksi.",
"Gambar 8.":
    "Histogram tersebut menunjukkan sebaran nilai detak jantung pada pembacaan akurasi tertinggi "
    "(n = 21.087): terpusat di sekitar rata-rata 83,4 bpm dengan rentang 60–123 bpm. Profil ini wajar "
    "untuk aktivitas ringan dan mengindikasikan data fisiologis yang masuk akal, bukan nilai beku "
    "akibat sensor tidak terpasang.",
"Gambar 9.":
    "Diagram tersebut merangkum kualitas kontak sensor selama periode pengukuran: mayoritas pembacaan "
    "(90,8%) berada pada tingkat akurasi tertinggi, 9,1% tanpa kontak, dan 0,1% pada tingkat sedang — "
    "mencerminkan lepas-kontak sesaat yang lazim pada sensor optik di pergelangan tangan.",
}
add_explanations(ROOT / "Draft_Naskah_HR_BLE_SINTA2.docx", expl_id)

# ============================ MANUSCRIPT EN ============================
expl_en = {
"Figure 1.":
    "In the figure, the left column shows the five smartwatch components — sensor, native code, "
    "application layer, SQLite, and the BLE advertiser/GATT server — flowing from acquisition to "
    "transmission (steps 1–5), while the right column shows the five smartphone components from BLE "
    "reception to data export (steps 6–10). The bridge on the bottom row depicts the two communication "
    "directions: data batches via NOTIFY from the smartwatch to the smartphone, and the ACK "
    "confirmation via WRITE in the opposite direction.",
"Figure 2.":
    "The figure shows one batch (a JSON array of bpm, accuracy, and time) being split into a sequence "
    "of frames: opened by START, followed by a number of DATA frames each carrying one payload chunk, "
    "and closed by END. The inset at the bottom shows the structure of a single DATA frame — a "
    "one-byte opcode followed by a chunk of at most MTU − 4 bytes — and each frame is sent as one BLE "
    "notification.",
"Figure 3.":
    "The sequence diagram reads top to bottom in three phases. Connection setup: the smartwatch "
    "advertises, the smartphone scans and connects, requests a 512-byte MTU, and enables "
    "notifications. Transfer: the batch is sent as a START–DATA–END sequence. Persistence and "
    "acknowledgement: the smartphone stores the batch in its database and writes the ACK, upon which "
    "the smartwatch marks the records as sent.",
"Figure 4.":
    "The flow in the figure spans two different time scales: sensor readings are stored every second "
    "(continuous), whereas fetching unsent records, sending the batch, and awaiting the ACK run per "
    "interval (3/5 minutes). The ACK decision branch captures the core of the reliability mechanism: "
    "if the ACK is received the records are marked as sent; otherwise they remain unsent and are "
    "retransmitted in the next interval.",
"Figure 5.":
    "The figure separates the two layers of the smartwatch application. The Dart layer contains the "
    "user interface, the monitoring logic in the BLoC pattern (per-second sampling, per-interval "
    "sending, ACK waiting), and SQLite access; the native Kotlin layer contains the sensor listener, "
    "the GATT server with its frame queue and flow control, the ACK characteristic, and the foreground "
    "service. The two layers communicate via platform channels: commands flow down through a "
    "MethodChannel and events flow up through EventChannels.",
"Figure 6.":
    "The chart compares the numbers of received and lost records in each session. Losses are "
    "concentrated in the longest session (Session 3) and in the session whose smartphone never "
    "connected (Session 2), while the other two sessions are nearly complete (99.40% and 98.23%, "
    "respectively).",
"Figure 7.":
    "The figure plots the heart-rate signal of the main session together with markers of lost "
    "records. All losses occur before the smartphone connected at around 15:00; once the link was "
    "active, loss was practically zero. This pattern confirms that the losses were caused by the "
    "disconnected period rather than by the BLE channel.",
"Figure 8.":
    "The histogram shows the distribution of heart-rate values for the highest-accuracy readings "
    "(n = 21,087): centred around the mean of 83.4 bpm with a range of 60–123 bpm. This profile is "
    "reasonable for light activity and indicates physiologically plausible data rather than frozen "
    "values from an unworn sensor.",
"Figure 9.":
    "The chart summarises the sensor contact quality during the measurement period: the majority of "
    "readings (90.8%) are at the highest accuracy level, 9.1% at no-contact, and 0.1% at the medium "
    "level — reflecting the momentary contact losses that are common for wrist-worn optical sensors.",
}
add_explanations(ROOT / "Draft_Manuscript_HR_BLE_EN.docx", expl_en)
