#!/usr/bin/env python3
"""Update Draft_Naskah_HR_BLE_SINTA2.docx dengan data terbaru (4 sesi),
tabel per-sesi, subbab temuan keterbatasan backfill, dan figur hasil."""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH as ALIGN
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
DOC = ROOT / "Draft_Naskah_HR_BLE_SINTA2.docx"
d = Document(str(DOC))
P = d.paragraphs

def set_text(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        p.add_run(text)

# ---------- referensi paragraf SEBELUM penyisipan ----------
p_3_3 = P[48]   # heading "3.3 Kualitas Data Sensor"
p_3_4 = P[52]   # heading "3.4 Kinerja Transfer"
p_3_5 = P[54]   # heading "3.5 Pembahasan dan Keterbatasan"
t_dev, t_deliv, t_acc = d.tables[0], d.tables[1], d.tables[2]

# ====================================================================
# 1) PARAGRAF NARATIF
# ====================================================================
set_text(P[5],
 "Pemantauan detak jantung secara berkelanjutan menggunakan perangkat wearable merupakan "
 "komponen penting dalam Internet of Medical Things (IoMT). Tantangan utamanya adalah menjamin "
 "kelengkapan data saat dikirim dari smartwatch ke smartphone melalui Bluetooth Low Energy (BLE), "
 "mengingat keterbatasan ukuran paket (MTU), sifat notifikasi BLE yang tidak terjamin sampai, serta "
 "pembatasan proses latar belakang oleh sistem operasi. Penelitian ini merancang dan mengimplementasikan "
 "sistem dua aplikasi: aplikasi smartwatch (Wear OS) berperan sebagai peripheral/GATT server yang "
 "merekam detak jantung tiap detik ke basis data lokal, lalu mengirimkannya secara berkala dalam bentuk "
 "paket (batch) ke aplikasi smartphone (Android) yang berperan sebagai central. Untuk keandalan, paket "
 "dipecah menjadi rangkaian frame ber-opcode (START/DATA/END) dengan kendali aliran, dilengkapi pola "
 "store-and-forward, konfirmasi penerimaan tingkat aplikasi (ACK), penerima yang idempoten (anti-duplikat), "
 "dan foreground service agar tetap berjalan saat layar mati. Evaluasi pada perangkat fisik mencakup empat "
 "sesi (total 23.215 pembacaan pengukuran, di luar 22.231 sampel saat sensor tidak terpasang). Saat tautan "
 "BLE aktif, kehilangan data mendekati nol dan seluruh pembacaan yang diterima identik dengan catatan "
 "smartwatch (fidelity nilai 100%) tanpa duplikat. Namun ditemukan keterbatasan: pada periode smartphone "
 "terputus, backlog tidak selalu di-backfill walau telah ditandai terkirim, sehingga rasio keberhasilan "
 "agregat lintas seluruh sesi sebesar 88,1%. Sebanyak 90,8% pembacaan berada pada akurasi sensor tertinggi. "
 "Satu paket berisi 228 pembacaan (10.717 byte, 24 frame) terkirim dalam ±0,32 detik pada MTU 512. Hasil "
 "menunjukkan mekanisme yang diusulkan menjaga integritas dan kelengkapan data saat tautan aktif; penguatan "
 "backfill backlog saat terputus serta pengujian variasi kondisi menjadi pekerjaan lanjutan.")

set_text(P[8],
 "Continuous heart-rate monitoring using wearable devices is a key component of the Internet of Medical "
 "Things (IoMT). A central challenge is guaranteeing data completeness when readings are transferred from a "
 "smartwatch to a smartphone over Bluetooth Low Energy (BLE), given the limited packet size (MTU), the "
 "best-effort nature of BLE notifications, and operating-system restrictions on background execution. This "
 "work designs and implements a two-application system: a smartwatch (Wear OS) application acting as a BLE "
 "peripheral/GATT server that records the heart rate every second into a local database and periodically "
 "sends it as a batch to a smartphone (Android) application acting as the central. For reliability, each "
 "batch is fragmented into opcode-tagged frames (START/DATA/END) with flow control, complemented by a "
 "store-and-forward scheme, application-level acknowledgement (ACK), an idempotent (duplicate-free) "
 "receiver, and a foreground service for background operation. Evaluation on physical devices spanned four "
 "sessions (23,215 measurement readings in total, excluding 22,231 samples recorded while the sensor was "
 "not worn). While the BLE link was active, data loss was near zero and every received reading was identical "
 "to the smartwatch record (100% value fidelity) with no duplicates. However, a limitation emerged: during "
 "smartphone-disconnected periods the backlog was not always backfilled despite being marked as sent, "
 "yielding an aggregate delivery ratio of 88.1% across all sessions. 90.8% of readings were at the highest "
 "sensor-accuracy level. One batch of 228 readings (10,717 bytes, 24 frames) was delivered in ~0.32 s at "
 "MTU 512. The results show the proposed mechanism preserves data integrity and completeness while the link "
 "is active; strengthening backlog backfill during disconnection and evaluating varied conditions are future work.")

set_text(P[14],
 "Kontribusi penelitian ini adalah: (1) rancangan dan implementasi sistem dua aplikasi Wear OS–Android "
 "untuk akuisisi dan pengiriman data detak jantung; (2) protokol pengiriman batch di atas notifikasi BLE "
 "dengan framing ber-opcode dan kendali aliran; (3) mekanisme keandalan berupa store-and-forward, "
 "konfirmasi penerimaan (ACK) tingkat aplikasi, dan penerima idempoten (anti-duplikat); (4) eksekusi latar "
 "belakang melalui foreground service; serta (5) evaluasi empiris lintas empat sesi terhadap rasio "
 "keberhasilan pengiriman, fidelity nilai, kualitas data, dan kinerja transfer—termasuk temuan "
 "keterbatasan backfill store-and-forward saat perangkat terputus—beserta perangkat bantu pengolahan "
 "data yang mendukung reproduksibilitas.")

set_text(P[36],
 "Pengujian dilakukan pada perangkat fisik karena komunikasi BLE tidak dapat diemulasikan. Spesifikasi "
 "perangkat uji ditunjukkan pada Tabel 1. Aplikasi dipasang dalam mode rilis (release) agar pengukuran "
 "kinerja tidak terbias oleh mode debug. Pengujian dijalankan dalam empat sesi pada rentang 23–28 Juni "
 "2026; pada tiap run smartphone menghubungi smartwatch, pemantauan dijalankan beberapa interval, lalu data "
 "dari kedua aplikasi diekspor ke format CSV dan basis data untuk dianalisis. Total terkumpul 23.215 "
 "pembacaan pada periode pengukuran (sensor terpasang).")

set_text(P[40],
 "Metrik yang diukur meliputi: (1) rasio keberhasilan pengiriman (delivery ratio), yaitu jumlah record yang "
 "cocok di smartphone dibagi jumlah record yang direkam smartwatch, dengan pencocokan berdasarkan atribut "
 "waktu (time) sehingga tidak bergantung pada sinkronisasi jam antar-perangkat; (2) fidelity nilai, yaitu "
 "proporsi record diterima yang nilainya (bpm dan akurasi) identik dengan catatan smartwatch; (3) jumlah "
 "duplikat di penerima; (4) distribusi akurasi sensor mengikuti konstanta status akurasi pada Android "
 "SensorManager (rentang -1 hingga 3) [2]; serta (5) latensi transfer dan throughput per batch yang dicatat "
 "otomatis oleh instrumentasi aplikasi. Throughput dihitung sebagai ukuran payload dibagi durasi transfer.")

set_text(P[45],
 "Evaluasi dilakukan pada empat sesi pengujian dengan hasil per sesi pada Tabel 2 dan profil temporalnya "
 "pada Gambar 6 dan Gambar 7. Secara agregat, smartwatch merekam 23.215 pembacaan pada periode pengukuran "
 "dan 20.455 di antaranya tercatat di smartphone (rasio keberhasilan 88,1%) tanpa duplikat. Namun, "
 "distribusi kehilangan tidak merata: saat tautan BLE aktif, kehilangan mendekati nol (mis. setelah "
 "smartphone terhubung pada sesi utama, delivery praktis 100%), sedangkan seluruh kehilangan terkonsentrasi "
 "pada periode ketika smartphone tidak terhubung—termasuk satu sesi singkat yang smartphone-nya tidak "
 "pernah terhubung sehingga seluruh pembacaannya tidak tersimpan di penerima. Penting dicatat, seluruh "
 "pembacaan yang berhasil diterima identik dengan catatan smartwatch berdasarkan pencocokan waktu, baik "
 "nilai bpm maupun akurasi (fidelity nilai 100%), sehingga tidak terjadi korupsi data pada kanal BLE. "
 "Anomali antara penanda terkirim di smartwatch dan keberadaan record di smartphone dibahas pada Subbab 3.5.")

set_text(P[46], "Tabel 2. Hasil pengiriman per sesi pengujian.")

set_text(P[49],
 "Distribusi akurasi sensor pada periode pengukuran ditunjukkan pada Tabel 3 dan Gambar 9. Sebanyak 90,8% "
 "pembacaan berada pada akurasi tertinggi (nilai 3), 9,1% (2.113 pembacaan) pada kondisi tanpa kontak "
 "(nilai -1) yang lazim terjadi saat sensor sesaat kehilangan kontak dengan kulit [2], dan 0,1% (15 "
 "pembacaan) pada nilai 0. Statistik nilai detak jantung pada pembacaan akurasi tertinggi (n=21.087, Gambar 8) "
 "adalah minimum 60 bpm, maksimum 123 bpm, rata-rata 83,4 bpm, dan simpangan baku 9,9 bpm. Perlu dicatat, "
 "22.231 sampel di ekor rekaman—bernilai beku (konstan) dengan akurasi ≤0 akibat smartwatch tidak "
 "terpasang—telah dikecualikan dari statistik di atas agar tidak membiaskan hasil.")

set_text(P[50], "Tabel 3. Distribusi nilai akurasi sensor (periode pengukuran).")

# heading 3.5 -> 3.6 (Pembahasan)
set_text(P[54], "3.6 Pembahasan dan Keterbatasan")

set_text(P[55],
 "Hasil menunjukkan kombinasi store-and-forward, konfirmasi ACK, dan penerima idempoten mampu menjaga "
 "kelengkapan dan integritas data selama tautan BLE aktif—dibuktikan oleh fidelity nilai 100% dan "
 "kehilangan mendekati nol pada kondisi terhubung—sementara framing ber-opcode dengan kendali aliran "
 "mengatasi keterbatasan ukuran notifikasi BLE. Pencocokan berbasis waktu membuat perhitungan keberhasilan "
 "tidak bergantung pada sinkronisasi jam. Sebaliknya, sebagaimana dibahas pada Subbab 3.5, keandalan "
 "menurun pada skenario putus koneksi sehingga backfill backlog perlu diperkuat.")

set_text(P[56],
 "Penelitian ini memiliki beberapa keterbatasan. Pertama, evaluasi mencakup empat sesi pada satu pasang "
 "perangkat dan satu lingkungan, sehingga diperlukan replikasi lintas perangkat, jarak, dan kondisi "
 "gangguan untuk memperoleh rata-rata dan simpangan baku yang representatif. Kedua, ditemukan bahwa backlog "
 "selama putus koneksi belum selalu di-backfill (Subbab 3.5); penguatan ACK dan pemicu backfill menjadi "
 "prioritas perbaikan. Ketiga, eksekusi latar belakang mencakup kondisi aplikasi di latar belakang dan "
 "layar mati, tetapi belum menjamin operasi saat aplikasi ditutup paksa atau setelah perangkat reboot, yang "
 "juga dibatasi kebijakan hemat daya OS. Keempat, komunikasi BLE belum menerapkan enkripsi/otentikasi, "
 "sehingga aspek keamanan data kesehatan menjadi agenda pengembangan. Kelima, sensor pada smartwatch "
 "konsumer belum tervalidasi secara klinis, sehingga fokus kontribusi adalah keandalan komunikasi, bukan "
 "akurasi medis nilai detak jantung.")

set_text(P[58],
 "Penelitian ini merancang dan mengimplementasikan sistem pengiriman data detak jantung dari smartwatch ke "
 "smartphone berbasis BLE dengan framing batch, store-and-forward, konfirmasi penerimaan (ACK), penerima "
 "idempoten, dan eksekusi latar belakang. Pengujian pada perangkat fisik lintas empat sesi (total 23.215 "
 "pembacaan) menunjukkan integritas data yang tinggi—fidelity nilai 100% dan kehilangan mendekati nol "
 "saat tautan BLE aktif—dengan 90,8% data berakurasi tertinggi serta transfer batch yang cepat (±0,32 "
 "detik untuk 228 pembacaan pada MTU 512). Evaluasi juga mengungkap keterbatasan: pada periode putus "
 "koneksi, backlog belum selalu di-backfill sehingga rasio keberhasilan agregat menjadi 88,1%. Pekerjaan "
 "lanjutan meliputi penguatan mekanisme backfill dan ACK agar tahan putus-sambung, pengukuran formal lintas "
 "variasi kondisi dengan replikasi, penguatan eksekusi latar belakang (termasuk auto-start setelah reboot), "
 "serta penambahan enkripsi untuk keamanan data.")

# ====================================================================
# 2) TABEL 1 — tambah baris dataset
# ====================================================================
r = t_dev.add_row()
r.cells[0].text = "Dataset uji"
r.cells[1].text = "23–28 Jun 2026; 4 sesi; 23.215 pembacaan pengukuran"

# ====================================================================
# 3) TABEL 3 — distribusi akurasi (periode pengukuran)
# ====================================================================
t_acc.add_row()
acc_rows = [("Akurasi sensor", "Jumlah", "Persentase"),
            ("3 (tinggi)", "21.087", "90,8%"),
            ("0 (sedang)", "15", "0,1%"),
            ("-1 (tanpa kontak)", "2.113", "9,1%"),
            ("Total", "23.215", "100%")]
for i, row in enumerate(acc_rows):
    for j, v in enumerate(row):
        t_acc.rows[i].cells[j].text = v
for c in t_acc.rows[0].cells:
    for rn in c.paragraphs[0].runs: rn.bold = True

# ====================================================================
# 4) TABEL 2 — bangun ulang jadi per-sesi (7 kolom)
# ====================================================================
deliv_rows = [
    ("Sesi", "Mulai (WIB)", "Durasi", "Direkam", "Diterima", "Hilang", "Delivery"),
    ("Sesi 1", "23/06 06:38", "36 mnt", "2.177", "2.164", "13", "99,40%"),
    ("Sesi 2", "23/06 08:41", "1 mnt", "53", "0", "53", "0,00%"),
    ("Sesi 3", "23/06 14:15", "345 mnt", "20.702", "18.013", "2.689", "87,01%"),
    ("Sesi 4", "28/06 21:29", "5 mnt", "283", "278", "5", "98,23%"),
    ("Total", "—", "387 mnt", "23.215", "20.455", "2.760", "88,11%"),
]
old_tbl = t_deliv._tbl
new = d.add_table(rows=len(deliv_rows), cols=7)
new.style = "Table Grid"
for i, row in enumerate(deliv_rows):
    for j, v in enumerate(row):
        cell = new.rows[i].cells[j]
        cell.text = v
        if i == 0 or row[0] == "Total":
            for rn in cell.paragraphs[0].runs: rn.bold = True
old_tbl.addnext(new._tbl)
old_tbl.getparent().remove(old_tbl)

# ====================================================================
# 5) SISIPKAN FIGUR HASIL (Gambar 6-9)
# ====================================================================
def insert_fig(before, img, caption, width_in):
    pim = before.insert_paragraph_before()
    pim.alignment = ALIGN.CENTER
    pim.add_run().add_picture(str(ROOT / "figures" / img), width=Inches(width_in))
    pc = before.insert_paragraph_before(caption)
    pc.alignment = ALIGN.CENTER
    pc.runs[0].font.size = Pt(10)

insert_fig(p_3_3, "fig_hr_completeness.png", "Gambar 6. Kelengkapan data per sesi pengujian (diterima vs hilang).", 5.6)
insert_fig(p_3_3, "fig_hr_timeline.png", "Gambar 7. Sinyal detak jantung dan kehilangan paket BLE pada sesi utama.", 6.3)
insert_fig(p_3_4, "fig_hr_bpm_dist.png", "Gambar 8. Distribusi nilai detak jantung pada pembacaan akurasi tertinggi.", 4.7)
insert_fig(p_3_4, "fig_hr_contact.png", "Gambar 9. Distribusi kualitas kontak sensor pada periode pengukuran.", 5.0)

# ====================================================================
# 6) SUBBAB BARU 3.5 — Temuan keterbatasan backfill (sebelum 3.6)
# ====================================================================
ph = p_3_5.insert_paragraph_before("3.5 Temuan: Keterbatasan Backfill Store-and-Forward")
ph.runs[0].bold = True
ph.runs[0].font.size = Pt(11)
pb = p_3_5.insert_paragraph_before(
 "Analisis silang antara penanda status terkirim (synced) pada smartwatch dan keberadaan record di "
 "smartphone mengungkap satu temuan penting. Sebanyak 2.879 pembacaan ditandai telah terkirim (synced = 1) "
 "oleh smartwatch, tetapi tidak ditemukan di basis data smartphone, sementara hanya 5 pembacaan terakhir "
 "yang memang masih berstatus belum terkirim. Seluruh selisih 2.879 record tersebut terjadi pada periode "
 "ketika smartphone sedang tidak terhubung. Hal ini menunjukkan bahwa, pada implementasi saat ini, "
 "penandaan terkirim dilakukan pada tingkat pengiriman notifikasi tanpa terkonfirmasi penuh oleh ACK "
 "tingkat aplikasi ketika tautan tidak aktif, sehingga backlog yang terbentuk selama putus koneksi tidak "
 "selalu dikirim ulang (di-backfill) setelah tautan pulih. Akibatnya, mekanisme store-and-forward menjaga "
 "kelengkapan dengan baik selama tautan aktif, tetapi belum sepenuhnya tahan terhadap skenario putus-sambung. "
 "Perbaikan yang disarankan adalah menunda penandaan synced hingga ACK diterima dan menambahkan pemicu "
 "backfill yang memindai record berstatus belum-terkonfirmasi setiap kali koneksi terbentuk kembali. Temuan "
 "ini memperkuat pentingnya konfirmasi tingkat aplikasi dibanding hanya mengandalkan notifikasi BLE.")

d.save(str(DOC))
print("OK — naskah diperbarui:", DOC.name)
print("tables:", len(d.tables), "| paragraphs:", len(d.paragraphs))
