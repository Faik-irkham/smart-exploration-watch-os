#!/usr/bin/env python3
"""Selaraskan daftar pustaka naskah Indonesia dengan naskah Inggris.

Masalah: naskah ID hanya memuat 12 rujukan dan 6 di antaranya dokumentasi
vendor (Bluetooth SIG, Android Developers, pub.dev), sehingga hanya tersisa
6 pustaka primer — terlalu tipis untuk jurnal SINTA 2.

Solusi: naskah ID memakai daftar pustaka dan penomoran yang sama persis
dengan naskah EN (33 rujukan, 31 jurnal/prosiding). Langkahnya:

1. Petakan ulang seluruh sitasi lama ID -> penomoran EN (sekali jalan, agar
   tidak terjadi pemetaan berantai).
2. Tulis ulang Bagian 1 mengikuti kedalaman related work naskah EN.
3. Tambahkan sitasi pada §2.2, §2.3, dan §2.4 yang sebelumnya kosong.
4. Ganti daftar pustaka 12 entri dengan 33 entri bergaya IEEE berbahasa
   Indonesia (dkk., dan, [Daring]. Tersedia:, Diakses:).

Rujukan pub.dev untuk flutter_blue_plus/sqflite/flutter_bloc dihapus dari
daftar pustaka; ketiga pustaka itu tetap disebut pada Tabel 1 sebagai bagian
lingkungan uji, yang merupakan tempat yang semestinya.
"""
import copy
import re
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[1]

# --------------------------------------------------------------------------
# 1. peta sitasi lama (ID) -> baru (penomoran EN)
# --------------------------------------------------------------------------
REMAP = {
    2: 33,   # SensorManager
    3: 10,   # Foreground services
    7: 1,    # He dkk. — survei IoMT
    8: 20,   # Balas dkk. — perilaku BLE pada pemantauan detak jantung
    9: 21,   # Xiong & Jiang — DTN store-carry-forward
    10: 22,  # Majumdar dkk. — survei WBAN
    11: 8,   # Xu dkk. — pemodelan kinerja BLE
    12: 23,  # Schweizer & Gilgen-Ammann — validasi sensor pergelangan
}
# [1] Bluetooth SIG dan [4]-[6] pub.dev ditangani manual (konteksnya berbeda).


def set_text(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        p.add_run(text)


def find(doc, prefix):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    raise SystemExit(f"PARAGRAF TIDAK DITEMUKAN: {prefix!r}")


def remap_text(t):
    return re.sub(r'\[(\d+)\]',
                  lambda m: f"[{REMAP.get(int(m.group(1)), int(m.group(1)))}]", t)


doc = Document(str(ROOT / "Draft_Naskah_HR_BLE_SINTA2.docx"))

# batas: jangan sentuh daftar pustaka lama (akan diganti seluruhnya)
ref_idx = next(i for i, p in enumerate(doc.paragraphs)
               if p.text.strip().startswith("Daftar Pustaka"))

# --- sitasi berkonteks khusus, dikerjakan sebelum pemetaan massal ---------
p_root = find(doc, "Akar masalahnya terletak pada urutan operasi")
set_text(p_root, p_root.text.replace(
    "tidak menyediakan konfirmasi pada lapisan atribut [1]",
    "tidak menyediakan konfirmasi pada lapisan atribut [9]"))

p_arch = find(doc, "Arsitektur sistem terdiri atas dua perangkat")
set_text(p_arch, p_arch.text
         .replace("basis data lokal SQLite (pustaka sqflite) [5]",
                  "basis data lokal SQLite")
         .replace("central/GATT client (pustaka flutter_blue_plus) [4]",
                  "central/GATT client"))

p_bloc = find(doc, "Aplikasi smartwatch dibangun dengan Flutter")
set_text(p_bloc, p_bloc.text.replace("pola BLoC [6]", "pola BLoC [32]"))

# --- pemetaan massal pada seluruh isi selain daftar pustaka ---------------
for p in doc.paragraphs[:ref_idx]:
    if "[" in p.text:
        new = remap_text(p.text)
        if new != p.text:
            set_text(p, new)
for tbl in doc.tables:
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                if "[" in p.text:
                    new = remap_text(p.text)
                    if new != p.text:
                        set_text(p, new)

# --------------------------------------------------------------------------
# 2. Bagian 1 ditulis ulang mengikuti kedalaman naskah EN
# --------------------------------------------------------------------------
set_text(find(doc, "Pemanfaatan perangkat wearable untuk pemantauan kesehatan"),
    "Pemanfaatan perangkat wearable untuk pemantauan kesehatan secara berkelanjutan berkembang "
    "pesat seiring meluasnya konsep Internet of Medical Things (IoMT), yaitu jaringan perangkat "
    "medis dan kesehatan yang mengumpulkan serta mempertukarkan data fisiologis [1], [2]. "
    "Smartwatch yang dilengkapi sensor detak jantung optik dapat berperan sebagai simpul sensor, "
    "sedangkan smartphone berfungsi sebagai gateway yang mengumpulkan, menyimpan, dan meneruskan "
    "data untuk analisis lebih lanjut [3], [4].")

set_text(find(doc, "Bluetooth Low Energy (BLE) menjadi pilihan utama"),
    "Bluetooth Low Energy (BLE) menjadi pilihan utama kanal komunikasi pada skenario ini karena "
    "konsumsi dayanya yang rendah [5], [6], [7]. Namun, pengiriman data kesehatan secara utuh "
    "melalui BLE menghadapi beberapa kendala: ukuran satu notifikasi dibatasi oleh Maximum "
    "Transmission Unit (MTU) [8]; notifikasi BLE bersifat best-effort sehingga tidak menyediakan "
    "konfirmasi pada lapisan atribut [9]; serta sistem operasi — terutama pada perangkat Wear OS "
    "dan ponsel dengan manajemen daya agresif — membatasi proses yang berjalan di latar belakang "
    "[10]. Pada aplikasi kesehatan, kehilangan sebagian data dapat menurunkan kualitas analisis, "
    "sehingga kelengkapan data menjadi kebutuhan yang esensial [11].")

set_text(find(doc, "Banyak implementasi terdahulu menekankan penampilan nilai terbaru"),
    "Banyak implementasi terdahulu menekankan penampilan nilai terbaru secara waktu-nyata [12], "
    "[13]. Wan dkk. [14] mengembangkan sistem IoT wearable yang mengirimkan data kesehatan "
    "langsung ke server awan tanpa perantara smartphone, sedangkan Manjunath dkk. [15] menyajikan "
    "patch multi-sensor nirbaterai yang dapat ditempel di kulit dan mengalirkan biosinyal secara "
    "kontinu melalui BLE. Arsitektur pengiriman langsung semacam itu memprioritaskan ketepatan "
    "waktu di atas penyanggaan dan tidak menerapkan lapisan store-and-forward lokal pada perangkat "
    "wearable [16], sehingga gangguan singkat pada tautan atau sumber daya mengakhiri kesinambungan "
    "pencatatan. Konsumsi daya juga tetap menjadi persoalan kritis bagi simpul sensor yang "
    "bergantung pada koneksi yang harus dipelihara terus-menerus [17]. Lebih jauh, Fadhel dan Hasan "
    "[18] menunjukkan bahwa menambahkan konfirmasi pengiriman yang ketat pada aliran waktu-nyata "
    "yang berkelanjutan justru dapat memperparah kepadatan lalu lintas serta menimbulkan tundaan "
    "dan kehilangan paket yang besar ketika beberapa pasien dipantau sekaligus, sementara kerangka "
    "kerja wearable terkini untuk deteksi dini risiko kardiovaskular juga masih bertumpu pada "
    "aliran berkelanjutan [19].")

set_text(find(doc, "Sejumlah studi terkini telah mengkaji aspek-aspek yang bersinggungan"),
    "Sejumlah studi terkini mengkaji aspek-aspek yang bersinggungan dengan penelitian ini. Pada "
    "tataran protokol, Balas dkk. [20] mengamati perilaku BLE pada aplikasi pemantauan detak "
    "jantung dan menunjukkan bahwa parameter koneksi memengaruhi kualitas pengiriman data. Pada "
    "tataran jaringan, Xiong dan Jiang [21] mengembangkan protokol perutean delay-tolerant "
    "berbasis paradigma store-carry-forward untuk jaringan dengan konektivitas terputus-putus, dan "
    "survei Majumdar dkk. [22] merangkum pendekatan keandalan, ketangguhan, serta efisiensi energi "
    "pada wireless body area network (WBAN). Pada tataran perangkat, studi validasi Schweizer dan "
    "Gilgen-Ammann [23] menegaskan bahwa akurasi sensor detak jantung pergelangan tangan bervariasi "
    "antar-kondisi aktivitas, sehingga kualitas kontak sensor perlu dilaporkan bersama data. Kajian "
    "penanganan pemutusan koneksi pada tumpukan protokol lain [24] dan evaluasi keandalan koneksi "
    "BLE pada aplikasi kesehatan bergerak [25] sama-sama mengindikasikan bahwa terputusnya tautan, "
    "bukan kerusakan payload, merupakan penyebab dominan kehilangan data.")

set_text(find(doc, "Meskipun demikian, studi-studi tersebut umumnya menitikberatkan"),
    "Meskipun demikian, satu keterbatasan tetap berulang pada implementasi-implementasi tersebut: "
    "tidak satu pun menjamin bahwa setiap pembacaan yang direkam benar-benar tersimpan di penerima "
    "ketika koneksi terputus sesaat atau ketika aplikasi penerima dihentikan oleh pembatasan "
    "eksekusi latar belakang sistem operasi seluler [26]. Yang dibutuhkan adalah mekanisme yang "
    "memadukan penyimpanan sementara (store-and-forward), konfirmasi penerimaan, pencegahan "
    "duplikasi, dan operasi latar belakang, yang diverifikasi secara empiris pada sepasang "
    "perangkat konsumer — smartwatch Wear OS dan smartphone Android — termasuk pembuktian bahwa "
    "setiap pembacaan yang direkam benar-benar tiba di penerima tanpa duplikasi. Celah itulah yang "
    "diisi penelitian ini melalui kombinasi store-and-forward, konfirmasi tingkat aplikasi (ACK), "
    "penerima idempoten, dan evaluasi kelengkapan berbasis pencocokan timestamp.")

# --------------------------------------------------------------------------
# 3. sitasi pada bagian metode yang sebelumnya kosong
# --------------------------------------------------------------------------
p22 = find(doc, "Komunikasi memakai satu layanan GATT khusus")
set_text(p22, p22.text
         .replace('karakteristik “record” bertipe notify.',
                  'karakteristik “record” bertipe notify [27].')
         .replace("MTU dikurangi 4 byte (3 byte header ATT dan 1 byte opcode).",
                  "MTU dikurangi 4 byte (3 byte header ATT dan 1 byte opcode) [28].")
         .replace("dikendalikan dengan flow control:",
                  "dikendalikan dengan flow control [29]:"))

p23 = find(doc, "Setiap pembacaan disimpan di smartwatch dengan penanda status")
set_text(p23, p23.text
         .replace("dikirim sebagai satu batch.", "dikirim sebagai satu batch [30].")
         .replace("jumlah record dalam batch yang berhasil diproses.",
                  "jumlah record dalam batch yang berhasil diproses [31]."))

# --------------------------------------------------------------------------
# 4. daftar pustaka baru
# --------------------------------------------------------------------------
REFS = [
 "[1] P. He dkk., “A survey of Internet of medical things: technology, application and future directions,” Digital Communications and Networks, vol. 12, no. 5, hlm. 717–742, Mei 2026, doi: 10.1016/j.dcan.2024.11.013.",
 "[2] C. Huang, J. Wang, S. Wang, dan Y. Zhang, “Internet of medical things: A systematic review,” Neurocomputing, vol. 557, hlm. 126719, Nov. 2023, doi: 10.1016/j.neucom.2023.126719.",
 "[3] K. T. Putra dkk., “A review on the application of Internet of Medical Things in wearable personal health monitoring: A cloud-edge artificial intelligence approach,” IEEE Access, vol. 12, hlm. 21437–21452, 2024, doi: 10.1109/ACCESS.2024.3358827.",
 "[4] S. F. Ahmed, Md. S. Bin Alam, S. Afrin, S. J. Rafa, N. Rafa, dan A. H. Gandomi, “Insights into Internet of Medical Things (IoMT): Data fusion, security issues and potential solutions,” Information Fusion, vol. 102, hlm. 102060, Feb. 2024, doi: 10.1016/j.inffus.2023.102060.",
 "[5] Z. Zhou dan H.-W. Huang, “Closed-loop transmission power control for reliable and low-power BLE communication in dynamic IoT settings,” IEEE Internet of Things Journal, vol. 13, no. 1, hlm. 1216–1228, Jan. 2026, doi: 10.1109/JIOT.2025.3627414.",
 "[6] R. Verma dan S. Kumar, “Insights into BLE 5.x data transfer modes for IoT-empowered remote patient monitoring,” IEEE Internet of Things Magazine, vol. 8, no. 4, hlm. 52–59, Jul. 2025, doi: 10.1109/IOTM.001.2400267.",
 "[7] G. Wandwi dan C. Wandwi, “Integrating IoT and wearable technologies in internal medicine for remote monitoring,” Journal of Sensors, IoT & Health Sciences, vol. 4, no. 1, hlm. 82–97, Mar. 2026, doi: 10.69996/jsihs.2026005.",
 "[8] H. Xu, Z. Yan, B. Li, dan M. Yang, “Modeling and analysis of the performance for Bluetooth Low Energy,” IEEE Communications Letters, vol. 28, no. 3, hlm. 732–736, Mar. 2024, doi: 10.1109/LCOMM.2024.3352545.",
 "[9] F. Battaglia, G. Gugliandolo, G. Campobello, dan N. Donato, “EEG-over-BLE: A low-latency, reliable, and low-power architecture for multichannel EEG monitoring systems,” IEEE Transactions on Instrumentation and Measurement, vol. 72, hlm. 1–10, 2023, doi: 10.1109/TIM.2023.3268471.",
 "[10] Android Developers, “Foreground services,” Android Developers Guide. Diakses: 30 Jul. 2026. [Daring]. Tersedia: https://developer.android.com/develop/background-work/services/foreground-services",
 "[11] J. Van Der Donckt dkk., “Mitigating data quality challenges in ambulatory wrist-worn wearable monitoring through analytical and practical approaches,” Scientific Reports, vol. 14, no. 1, hlm. 17545, Jul. 2024, doi: 10.1038/s41598-024-67767-3.",
 "[12] A. González-Pérez, M. Matey-Sanz, C. Granell, L. Díaz-Sanahuja, J. Bretón-López, dan S. Casteleyn, “AwarNS: A framework for developing context-aware reactive mobile applications for health and mental health,” Journal of Biomedical Informatics, vol. 141, hlm. 104359, Mei 2023, doi: 10.1016/j.jbi.2023.104359.",
 "[13] O. Alruwaili, A. Yousef, dan A. Armghan, “Monitoring the transmission of data from wearable sensors using probabilistic transfer learning,” IEEE Access, vol. 12, hlm. 97460–97475, 2024, doi: 10.1109/ACCESS.2024.3428444.",
 "[14] J. Wan dkk., “Wearable IoT enabled real-time health monitoring system,” EURASIP Journal on Wireless Communications and Networking, vol. 2018, no. 1, Des. 2018, doi: 10.1186/s13638-018-1308-x.",
 "[15] B. B. Manjunath dkk., “Battery-free, wireless, and skin-mountable multi-sensory patch for biosignal monitoring,” Sensors and Actuators A: Physical, vol. 404, Jul. 2026, doi: 10.1016/j.sna.2026.117751.",
 "[16] C. Y. Kim dkk., “Wireless technologies for wearable electronics: A review,” Advanced Electronic Materials, Jul. 2025, doi: 10.1002/aelm.202400884.",
 "[17] S. Gautam dan S. Kumar, “BLE periodic advertising-based energy-efficient sensor node operation for transfer of large data in monitoring applications,” IEEE Internet of Things Journal, vol. 11, no. 24, hlm. 40070–40085, Des. 2024, doi: 10.1109/JIOT.2024.3451698.",
 "[18] A. A. Fadhel dan H. M. Hasan, “Reducing delay and packets loss in IoT-cloud based ECG monitoring by Gaussian modeling,” International Journal of Online and Biomedical Engineering, vol. 19, no. 6, hlm. 97–113, 2023, doi: 10.3991/ijoe.v19i06.38581.",
 "[19] Arivardhini dan Kayalvizhi, “IoT-based wearable health monitoring system for early detection of cardiovascular disease,” pracetak SSRN. [Daring]. Tersedia: https://ssrn.com/abstract=5790382",
 "[20] Z. Balas, K. Tokarz, B. Zieliński, dan T. Guźniczak, “Research on the behaviour of Bluetooth Low Energy protocol in the heart rate monitoring application,” Procedia Computer Science, vol. 225, hlm. 63–69, 2023, doi: 10.1016/j.procs.2023.09.092.",
 "[21] Y. Xiong dan S. Jiang, “Multi-decision dynamic intelligent routing protocol for delay-tolerant networks,” Electronics, vol. 12, no. 21, art. 4528, 2023, doi: 10.3390/electronics12214528.",
 "[22] P. Majumdar, S. Roy, S. Sikdar, P. Ghosh, dan N. Ghosh, “A survey on data-driven approaches for reliability, robustness, and energy efficiency in wireless body area networks,” Sensors, vol. 24, no. 20, art. 6531, 2024, doi: 10.3390/s24206531.",
 "[23] T. Schweizer dan R. Gilgen-Ammann, “Wrist-worn and arm-worn wearables for monitoring heart rate during sedentary and light-to-vigorous physical activities: Device validation study,” JMIR Cardio, vol. 9, art. e67110, 2025, doi: 10.2196/67110.",
 "[24] M. Domingues, J. N. Faria, dan D. Portugal, “Dimensioning payload size for fast retransmission of MQTT packets in the wake of network disconnections,” EURASIP Journal on Wireless Communications and Networking, vol. 2024, no. 1, Des. 2024, doi: 10.1186/s13638-023-02327-3.",
 "[25] F. Franco, L. Lamazzi, F. Poggi, dan L. Bedogni, “Evaluating Bluetooth Low Energy connection reliability for mobile health applications,” dalam Prosiding 2026 IEEE 23rd Consumer Communications & Networking Conference (CCNC), Jan. 2026, hlm. 1–6, doi: 10.1109/CCNC65079.2026.11366362.",
 "[26] C. Slade, Y. Sun, W. C. Chao, C.-C. Chen, R. M. Benzo, dan P. Washington, “Current challenges and opportunities in active and passive data collection for mobile health sensing: A scoping review,” JAMIA Open, vol. 8, no. 4, Jul. 2025, doi: 10.1093/jamiaopen/ooaf025.",
 "[27] C. Hirsch, L. Davoli, R. Grosu, dan G. Ferrari, “DynGATT: A dynamic GATT-based data synchronization protocol for BLE networks,” Computer Networks, vol. 222, Feb. 2023, doi: 10.1016/j.comnet.2023.109560.",
 "[28] M. Baert, B. Moons, J. Pittevils, Y. Song, N. Madhu, dan J. Hoebeke, “Evaluation of BLE-based audio broadcasting under probabilistic interference,” Computer Communications, vol. 222, hlm. 130–140, Jun. 2024, doi: 10.1016/j.comcom.2024.04.034.",
 "[29] D. F. S. Santos, M. M. Bezerra, W. D. P. da Silva, H. O. Almeida, dan A. Perkusich, “CrediBLE: A credit-based adaptive flow control architecture for Bluetooth Low-Energy gateways,” IEEE Access, vol. 14, hlm. 92942–92965, 2026, doi: 10.1109/ACCESS.2026.3705713.",
 "[30] Y. Li, J. Lv, B. Li, dan W. Dong, “RT-BLE: Real-time multi-connection scheduling for Bluetooth Low Energy,” dalam Prosiding IEEE INFOCOM 2023 – IEEE Conference on Computer Communications, Mei 2023, hlm. 1–10, doi: 10.1109/INFOCOM53939.2023.10229006.",
 "[31] N. Landra, D. Demarchi, dan P. M. Ros, “SharkTooth: A scalable real-time algorithm for BLE-based wireless body sensor networks synchronization,” IEEE Internet of Things Journal, vol. 12, no. 22, hlm. 46174–46192, Nov. 2025, doi: 10.1109/JIOT.2025.3602162.",
 "[32] Sanghmitra, “The state management dilemma: BLoC vs. Provider in modern Flutter development,” International Journal of Scientific Research in Computer Science, Engineering and Information Technology, vol. 10, no. 5, hlm. 326–336, Okt. 2024, doi: 10.32628/cseit241051027.",
 "[33] Google, “SensorManager,” Android Developer Reference. Diakses: 27 Jul. 2026. [Daring]. Tersedia: https://developer.android.com/reference/android/hardware/SensorManager",
]

old_refs = [p for p in doc.paragraphs[ref_idx + 1:]
            if re.match(r'^\[\d+\]', p.text.strip())]
template = old_refs[0]
for text, para in zip(REFS, old_refs):
    set_text(para, text)
for extra in REFS[len(old_refs):]:
    el = copy.deepcopy(template._p)
    old_refs[-1]._p.addnext(el)
    new = Paragraph(el, template._parent)
    set_text(new, extra)
    old_refs.append(new)
for leftover in old_refs[len(REFS):]:
    leftover._p.getparent().remove(leftover._p)

set_text(find(doc, "Catatan: dokumen ini adalah draf."),
    "Catatan: dokumen ini adalah draf. Sesuaikan dengan template resmi jurnal SINTA 2 tujuan "
    "(format kolom, gaya sitasi, batas halaman). Daftar pustaka dinomori mengikuti urutan sitasi "
    "pertama sesuai gaya IEEE; rujukan [19] masih berupa pracetak SSRN yang tahun dan venuenya "
    "perlu dilengkapi. Bagian bertanda [[...]] perlu dilengkapi penulis, termasuk URL repositori "
    "pada bagian Ketersediaan Data dan Kode.")

doc.save(str(ROOT / "Draft_Naskah_HR_BLE_SINTA2.docx"))
print("OK — Draft_Naskah_HR_BLE_SINTA2.docx")
