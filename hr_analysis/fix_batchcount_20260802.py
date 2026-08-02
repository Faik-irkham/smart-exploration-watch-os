#!/usr/bin/env python3
"""Koreksi jumlah batch Sesi 5 + tambahkan bukti antrean store-and-forward.

TEMUAN AUDIT (2 Agu 2026)
Naskah menyebut Sesi 5 terdiri atas "tujuh batch". Angka itu berasal dari
watch_log.txt yang memuat 7 baris HR-METRIC, padahal penangkapan logcat baru
dimulai pukul 23.19.04 sedangkan sesi mulai 23.07.11 (baris pertama log
berbunyi "--------- beginning of main").

    7 batch x 180 record = 1.260
    record diterima ponsel = 1.974
    selisih                =   714  -> batch yang tidak tertangkap log

Rekonstruksi jadwal flush dari data ponsel (interval 3 menit, detik :04)
menutup persis pada 1.974 record dengan 11 flush:
    23:10:04 (173) 23:13:04 23:16:04 23:19:04* 23:22:04* 23:25:04*
    23:28:04* 23:31:04* 23:34:04* 23:37:04* 23:40:04      (* = ada di log)

Jadi: sesi punya 11 batch, 7 di antaranya terinstrumentasi. Statistik kinerja
(121,5 +/- 25,5 ms; 70,3 +/- 13,2 KiB/detik) tetap sah, tetapi cakupannya
harus dinyatakan sebagai tujuh batch terinstrumentasi.

BUKTI BARU
Berkas 2026-07-05_watch_hr_20260705_233035.csv adalah ekspor smartwatch di
tengah Sesi 5 (23.30.35): 1.405 record, 151 di antaranya synced = 0. Flush
terakhir sebelumnya 23.28.04, berjarak tepat 151 detik. Antrean menahan persis
pembacaan sejak flush terakhir — bukti langsung penyangga bekerja benar pada
versi revisi. Ditambahkan sebagai paragraf baru di Bagian 3.E.
"""
import copy
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[1]


def set_text(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        p.add_run(text)


def find(doc, prefix):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    raise SystemExit(f"PARAGRAF TIDAK DITEMUKAN: {prefix!r}")


def sub(doc, prefix, pairs):
    p = find(doc, prefix)
    t = p.text
    for old, new in pairs:
        if old not in t:
            raise SystemExit(f"POTONGAN TIDAK DITEMUKAN pada {prefix!r}: {old!r}")
        t = t.replace(old, new)
    set_text(p, t)
    return p


def insert_after(ref, text):
    el = copy.deepcopy(ref._p)
    ref._p.addnext(el)
    new = Paragraph(el, ref._parent)
    set_text(new, text)
    return new


# ==========================================================================
# NASKAH INDONESIA
# ==========================================================================
doc = Document(str(ROOT / "Draft_Naskah_HR_BLE_SINTA2.docx"))

# --- abstrak -------------------------------------------------------------
sub(doc, "Pemantauan detak jantung berkelanjutan", [
    ("tanpa record keliru-tanda dan seluruh tujuh batch terkonfirmasi.",
     "tanpa record keliru-tanda; sesi tersebut mencakup sebelas batch pengiriman dan seluruh "
     "tujuh batch yang terekam instrumentasi memperoleh konfirmasi."),
])

# --- 3.B Keberhasilan Pengiriman -----------------------------------------
sub(doc, "Hasil pengiriman dilaporkan terpisah menurut versi", [
    ("Seluruh tujuh batch pada sesi tersebut memperoleh ACK.",
     "Sesi tersebut mencakup sebelas batch pengiriman pada interval 3 menit; tujuh di antaranya "
     "terekam oleh instrumentasi dan seluruhnya memperoleh ACK, sedangkan keberhasilan empat "
     "batch selebihnya disimpulkan dari kecocokan record karena penangkapan log baru dimulai di "
     "tengah sesi."),
])

# --- 3.D Kinerja Transfer ------------------------------------------------
sub(doc, "Kinerja transfer diukur melalui instrumentasi HR-METRIC", [
    ("Pada Sesi 5 (interval 3 menit, tujuh batch), seluruh batch berukuran seragam",
     "Sesi 5 mencakup sebelas batch pengiriman pada interval 3 menit, tetapi penangkapan log baru "
     "dimulai pada batch keempat sehingga yang terekam instrumentasi adalah tujuh batch terakhir. "
     "Ketujuh batch itu berukuran seragam"),
    ("setiap batch dikonfirmasi ACK penuh oleh smartphone.",
     "ketujuh batch tersebut dikonfirmasi ACK penuh oleh smartphone. Keseragaman ukuran berlaku "
     "bagi batch yang jatuh pada interval penuh; batch pertama sesi berisi sekitar 173 pembacaan "
     "karena sesi dimulai di tengah interval."),
])

# --- 3.F.1 Interpretasi (duty cycle) -------------------------------------
sub(doc, "Ditinjau dari penggunaan kanal", [
    ("konfirmasi terjadi tujuh kali sepanjang sesi, bukan sekali untuk setiap pembacaan.",
     "konfirmasi terjadi sebelas kali sepanjang sesi 33 menit, bukan sekali untuk setiap "
     "pembacaan."),
])

# --- 3.E Temuan ----------------------------------------------------------
p35 = sub(doc, "Analisis silang tiga sesi valid versi awal", [
    ("tujuh batch masing-masing menerima ACK 180 record",
     "ketujuh batch yang terekam instrumentasi masing-masing menerima ACK 180 record"),
])

insert_after(p35,
    "Bukti langsung bahwa antrean bekerja sebagaimana dirancang diperoleh dari ekspor basis data "
    "smartwatch yang diambil di tengah Sesi 5, tepatnya pukul 23.30.35. Pada saat itu basis data "
    "memuat 1.405 record dan 151 di antaranya masih berstatus synced = 0. Pengiriman terakhir "
    "sebelum ekspor berlangsung pukul 23.28.04, sehingga selang sejak pengiriman tersebut adalah "
    "151 detik — tepat sama dengan jumlah record yang menunggu. Dengan kata lain, antrean menahan "
    "persis seluruh pembacaan yang direkam sejak pengiriman terakhir, tidak lebih dan tidak "
    "kurang, lalu dikosongkan pada pengiriman berikutnya. Berbeda dengan versi awal yang menandai "
    "record sebagai terkirim sebelum dikonfirmasi, versi revisi memelihara antrean yang akurat "
    "sepanjang sesi.")

# --- Kesimpulan ----------------------------------------------------------
sub(doc, "Penelitian ini merancang dan mengimplementasikan sistem pengiriman data", [
    ("dan transfer tujuh batch 180 record pada MTU 512 memerlukan rata-rata 121,5 ms",
     "dan transfer batch 180 record pada MTU 512 memerlukan rata-rata 121,5 ms"),
])

# --- abstract Inggris di naskah ID ---------------------------------------
ABS_EN_OLD = "with no falsely flagged records and all seven batches acknowledged."
ABS_EN_NEW = ("with no falsely flagged records; that session comprised eleven transmission "
              "batches and all seven instrumented batches were acknowledged.")
sub(doc, "Continuous wearable heart-rate monitoring", [(ABS_EN_OLD, ABS_EN_NEW)])

doc.save(str(ROOT / "Draft_Naskah_HR_BLE_SINTA2.docx"))
print("OK — Draft_Naskah_HR_BLE_SINTA2.docx")


# ==========================================================================
# NASKAH INGGRIS
# ==========================================================================
doc = Document(str(ROOT / "Draft_Manuscript_HR_BLE_EN_v2.docx"))

sub(doc, "Continuous wearable heart-rate monitoring", [(ABS_EN_OLD, ABS_EN_NEW)])

sub(doc, "Delivery results are reported separately by software version", [
    ("All seven batches in that session were acknowledged.",
     "That session comprised eleven transmission batches at the 3-minute interval; seven of them "
     "were captured by the instrumentation and all seven were acknowledged, while delivery of the "
     "remaining four is inferred from record matching because the log capture began mid-session."),
])

sub(doc, "Transfer performance was measured through the HR-METRIC", [
    ("In Session 5 (3-minute interval, seven batches), every batch was uniform",
     "Session 5 comprised eleven transmission batches at a 3-minute interval, but log capture "
     "began at the fourth batch, so the instrumented set is the last seven batches. Those seven "
     "were uniform"),
    ("every batch was fully acknowledged by the smartphone.",
     "each of the seven was fully acknowledged by the smartphone. The uniform size applies to "
     "batches falling on a full interval; the first batch of the session held about 173 readings "
     "because the session started mid-interval."),
])

sub(doc, "In terms of channel occupancy", [
    ("acknowledgement happens seven times per session rather than once per reading.",
     "acknowledgement happens eleven times across a 33-minute session rather than once per "
     "reading."),
])

p35 = sub(doc, "A cross-analysis of the three valid initial-version sessions", [
    ("seven batches each received an ACK for 180 records",
     "the seven instrumented batches each received an ACK for 180 records"),
])

insert_after(p35,
    "Direct evidence that the queue behaves as designed comes from a smartwatch database export "
    "taken mid-session during Session 5, at 23:30:35. At that moment the database held 1,405 "
    "records, 151 of which were still at synced = 0. The last transmission before the export "
    "occurred at 23:28:04, so the elapsed interval was 151 seconds — exactly the number of "
    "records waiting. The queue therefore held precisely the readings recorded since the last "
    "transmission, no more and no less, and was drained at the next one. Unlike the initial "
    "version, which marked records as sent before they were confirmed, the revised version "
    "maintained an accurate queue throughout the session.")

sub(doc, "This study designed and implemented a BLE heart-rate delivery system", [
    ("and seven 180-record batches at MTU 512 transferred in 121.5 ms on average",
     "and 180-record batches at MTU 512 transferred in 121.5 ms on average"),
])

doc.save(str(ROOT / "Draft_Manuscript_HR_BLE_EN_v2.docx"))
print("OK — Draft_Manuscript_HR_BLE_EN_v2.docx")
