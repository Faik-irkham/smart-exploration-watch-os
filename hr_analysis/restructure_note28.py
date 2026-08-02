#!/usr/bin/env python3
"""Susun ulang catatan 28-sederhana dengan pola:
kalimat pengantar -> tabel/gambar -> penjelasan hasilnya.
Gaya & gambar dipertahankan (template = file itu sendiri)."""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH as ALIGN
from docx.oxml.ns import qn
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
F = ROOT / "docs" / "26-06-28-Catatan-Faik-BLE-SmartWatch-sederhana.docx"
FIG = ROOT / "figures"

d = Document(str(F))
for child in list(d.element.body):
    if child.tag != qn('w:sectPr'):
        d.element.body.remove(child)

def H(t, size=12):
    p = d.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.size = Pt(size); return p
def body(t): return d.add_paragraph(t)
def cap(t):
    p = d.add_paragraph(t); p.alignment = ALIGN.CENTER
    p.runs[0].italic = True; p.runs[0].font.size = Pt(10); return p
def figure(img, w):
    p = d.add_paragraph(); p.alignment = ALIGN.CENTER
    p.add_run().add_picture(str(FIG / img), width=Inches(w))
def table(rows):
    t = d.add_table(rows=len(rows), cols=len(rows[0])); t.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            c = t.rows[i].cells[j]; c.text = v
            if i == 0 or row[0] == "Total":
                for rn in c.paragraphs[0].runs: rn.bold = True

DELIV = [
 ("Sesi", "Mulai (WIB)", "Durasi", "Direkam", "Diterima", "Hilang", "Delivery"),
 ("Sesi 1", "23/06 06:38", "36 mnt", "2.177", "2.164", "13", "99,40%"),
 ("Sesi 2", "23/06 08:41", "1 mnt", "53", "0", "53", "0,00%"),
 ("Sesi 3", "23/06 14:15", "345 mnt", "20.702", "18.013", "2.689", "87,01%"),
 ("Sesi 4", "28/06 21:29", "5 mnt", "283", "278", "5", "98,23%"),
 ("Total", "—", "387 mnt", "23.215", "20.455", "2.760", "88,11%"),
]
ACC = [
 ("Akurasi sensor", "Jumlah", "Persentase"),
 ("3 (tinggi)", "21.087", "90,8%"),
 ("0 (sedang)", "15", "0,1%"),
 ("-1 (tanpa kontak)", "2.113", "9,1%"),
 ("Total", "23.215", "100%"),
]

# ===================== ISI =====================
H("CATATAN HASIL UJI — PENGIRIMAN DATA DETAK JANTUNG DAN TEMUAN PENTING", 13)
s = d.add_paragraph("28 Juni 2026 · lanjutan dari catatan 26 Juni 2026 (jalur komunikasi BLE/GATT)")
s.runs[0].italic = True; s.runs[0].font.size = Pt(10)

H("Pengambilan Data")
body("Data diambil langsung dari kedua perangkat melalui koneksi adb (aplikasi versi rilis): berkas "
     "CSV hasil ekspor disalin dari folder Download. Smartwatch yang digunakan adalah Samsung Galaxy "
     "Watch (SM-R860), dan ponselnya Xiaomi Redmi Note 10 Pro. Berkas mentahnya bernama "
     "watch-2026-06-28.csv (45.446 baris) dan phone-2026-06-28.csv (42.562 baris), mencakup tanggal 23 "
     "sampai 28 Juni 2026. Berdasarkan jeda waktunya, data terbagi menjadi empat sesi.")

H("Membuang Data Saat Smartwatch Tidak Dipakai")
body("Hampir separuh data (22.231 baris) ternyata tidak valid: nilainya beku di angka 106 dengan kontak "
     "buruk (accuracy ≤ 0). Kondisi ini terjadi saat smartwatch dilepas, tetapi aplikasinya masih "
     "merekam. Karena bukan detak jantung yang sebenarnya, data ini dikeluarkan dari perhitungan. "
     "Sisanya, yaitu data pengukuran yang sah, berjumlah 23.215 baris. Patokan pemotongannya: setiap "
     "sesi dipotong sampai data dengan kontak baik (accuracy = 3) yang terakhir.")

H("Hasil Pengiriman Data")
# -- Tabel 1: pengantar -> tabel -> penjelasan --
body("Hasil pengiriman pada setiap sesi dirangkum pada Tabel 1.")
cap("Tabel 1. Hasil pengiriman per sesi (periode pengukuran).")
table(DELIV)
body("Secara keseluruhan, smartwatch merekam 23.215 data dan 20.455 di antaranya sampai ke ponsel, "
     "sehingga tingkat keberhasilannya 88,11% tanpa data ganda. Namun tingkat keberhasilan antar-sesi "
     "berbeda cukup jauh: sesi yang ponselnya terhubung penuh mencapai 98–99%, sedangkan satu sesi "
     "pendek yang ponselnya tidak pernah terhubung (Sesi 2) kehilangan seluruh datanya.")
# -- Gambar 1: pengantar -> gambar -> penjelasan --
body("Perbandingan jumlah data yang sampai dan yang hilang pada tiap sesi ditampilkan pada Gambar 1.")
figure("fig_hr_completeness.png", 5.6)
cap("Gambar 1. Kelengkapan data tiap sesi (data yang sampai dibanding yang hilang).")
body("Dari gambar terlihat bahwa kehilangan paling banyak terjadi pada sesi besar (Sesi 3) dan sesi "
     "yang gagal terhubung (Sesi 2), sementara sesi lainnya hampir lengkap.")
# -- Gambar 2: pengantar -> gambar -> penjelasan --
body("Pola kehilangan data terhadap waktu pada sesi besar ditunjukkan pada Gambar 2.")
figure("fig_hr_timeline.png", 6.3)
cap("Gambar 2. Grafik detak jantung dan titik data yang hilang pada sesi besar.")
body("Grafik ini memperjelas bahwa kehilangan hanya terjadi di awal, yaitu sebelum ponsel terhubung "
     "sekitar pukul 15.00. Setelah terhubung, hampir tidak ada data yang hilang. Selain itu, semua "
     "data yang sampai memiliki isi yang sama persis dengan catatan smartwatch (nilai bpm dan akurasi), "
     "sehingga tidak ada data yang rusak selama pengiriman.")

H("Kualitas Data Sensor")
# -- Tabel 2: pengantar -> tabel -> penjelasan --
body("Sebaran kualitas kontak sensor pada data pengukuran dirangkum pada Tabel 2.")
cap("Tabel 2. Distribusi nilai akurasi sensor (periode pengukuran).")
table(ACC)
body("Sebagian besar data (90,8%) memiliki kontak baik (accuracy 3); sisanya 9,1% tanpa kontak (-1) dan "
     "0,1% kontak sedang (0). Untuk data dengan kontak baik (21.087 baris), detak jantung terendah 60 "
     "bpm, tertinggi 123 bpm, dan rata-rata 83,4 bpm.")

H("TEMUAN PENTING — Data Tidak Dikirim Ulang Saat Koneksi Putus-Sambung")
body("Temuan ini melengkapi catatan 26 Juni (bagian ACK dan store-and-forward). Setelah dibandingkan, "
     "terdapat 2.879 data yang ditandai “sudah terkirim” (synced = 1) oleh smartwatch, tetapi "
     "sebenarnya tidak ada di ponsel. Hanya 5 data terakhir yang memang masih dalam antrean, dan itu "
     "wajar. Seluruh 2.879 data tersebut hilang ketika ponsel sedang terputus.")
body("Artinya, untuk saat ini smartwatch menandai data “terkirim” begitu data dikirim, belum benar-benar "
     "menunggu balasan (ACK) dari ponsel ketika koneksi terputus. Akibatnya, data yang menumpuk selama "
     "koneksi terputus tidak otomatis dikirim ulang saat tersambung kembali. Dengan kata lain, "
     "mekanisme ini bekerja baik selama koneksi tersambung, tetapi belum tahan terhadap kondisi "
     "putus-sambung.")
body("Saran perbaikan: (1) jangan menandai data “terkirim” sebelum balasan ACK benar-benar diterima; "
     "(2) tambahkan proses kirim ulang otomatis yang memeriksa data yang belum terkonfirmasi setiap "
     "kali koneksi tersambung kembali; (3) gunakan tiga status, yaitu belum dikirim, sudah dikirim "
     "tetapi belum dibalas, dan sudah dipastikan. Intinya, konfirmasi pada tingkat aplikasi (ACK) lebih "
     "dapat diandalkan daripada hanya mengandalkan notifikasi BLE.")

H("Perubahan pada Naskah")
body("Kedua draf naskah (versi Indonesia dan Inggris) telah diperbarui dengan data ini: Tabel 2 menjadi "
     "hasil per sesi, Tabel 3 berisi distribusi akurasi, ditambahkan ukuran kesamaan nilai (fidelity), "
     "serta subbab baru 3.5 mengenai temuan kirim ulang ini, dilengkapi Gambar 6–9. Angka utama "
     "berubah: dari delivery 99,86% (satu sesi lama) menjadi isi data 100% sama dan hampir tidak ada "
     "yang hilang saat koneksi tersambung, dengan angka gabungan 88,1% yang justru memperlihatkan "
     "kelemahan kirim ulang tersebut. Skrip analisisnya: hr_analysis/make_figs.py.")

d.save(str(F))
print("OK — disusun ulang:", F.name)
print("tabel:", len(d.tables), "gambar:", len(d.element.body.findall('.//'+qn('pic:pic'))))
