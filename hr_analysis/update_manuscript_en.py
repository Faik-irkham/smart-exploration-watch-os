#!/usr/bin/env python3
"""Apply the same latest-data updates to the English manuscript."""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH as ALIGN
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
DOC = ROOT / "Draft_Manuscript_HR_BLE_EN.docx"
d = Document(str(DOC))
P = d.paragraphs

def set_text(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        p.add_run(text)

p_3_3 = P[45]   # "3.3 Sensor Data Quality"
p_3_4 = P[49]   # "3.4 Transfer Performance"
p_3_5 = P[51]   # "3.5 Discussion and Limitations"
t_dev, t_deliv, t_acc = d.tables[0], d.tables[1], d.tables[2]

# ---------- narrative ----------
set_text(P[5],
 "Continuous heart-rate monitoring using wearable devices is a key component of the Internet of Medical "
 "Things (IoMT). A central challenge is guaranteeing data completeness when readings are transferred from a "
 "smartwatch to a smartphone over Bluetooth Low Energy (BLE), given the limited packet size (MTU), the "
 "best-effort nature of BLE notifications, and operating-system restrictions on background execution. This "
 "work designs and implements a two-application system: a smartwatch (Wear OS) application acting as a BLE "
 "peripheral/GATT server that records the heart rate every second into a local database and periodically "
 "sends it as a batch to a smartphone (Android) application acting as the central. For reliability, each "
 "batch is fragmented into opcode-tagged frames (START/DATA/END) with flow control, complemented by a "
 "store-and-forward scheme, application-level acknowledgement (ACK), an idempotent (duplicate-free) "
 "receiver, and a foreground service for background operation. Evaluation on physical devices spanned four "
 "sessions (23,215 measurement readings in total, excluding 22,231 samples recorded while the sensor was "
 "not worn). While the BLE link was active, data loss was near zero and every received reading was identical "
 "to the smartwatch record (100% value fidelity) with no duplicates. However, a limitation emerged: during "
 "smartphone-disconnected periods the backlog was not always backfilled despite being marked as sent, "
 "yielding an aggregate delivery ratio of 88.1% across all sessions. 90.8% of readings were at the highest "
 "sensor-accuracy level. One batch of 228 readings (10,717 bytes, 24 frames) was delivered in ~0.32 s at "
 "MTU 512. The results show the proposed mechanism preserves data integrity and completeness while the link "
 "is active; strengthening backlog backfill during disconnection and evaluating varied conditions are future work.")

set_text(P[11],
 "The contributions of this work are: (1) the design and implementation of a two-application Wear OS–Android "
 "system for acquiring and delivering heart-rate data; (2) a batch delivery protocol over BLE notifications "
 "using opcode-based framing with flow control; (3) reliability mechanisms consisting of store-and-forward, "
 "application-level acknowledgement (ACK), and an idempotent (duplicate-free) receiver; (4) background "
 "execution via a foreground service; and (5) an empirical evaluation across four sessions of delivery "
 "ratio, value fidelity, data quality, and transfer performance—including the finding of a store-and-forward "
 "backfill limitation during disconnection—together with supporting data tools for reproducibility.")

set_text(P[33],
 "Testing was conducted on physical devices because BLE communication cannot be emulated. The device "
 "specifications are listed in Table 1. The applications were installed in release mode so that performance "
 "measurements are not biased by debug mode. Testing was carried out over four sessions during 23–28 June "
 "2026; in each run the smartphone connects to the smartwatch, monitoring is started and left running for "
 "several intervals, and the data from both applications is then exported to CSV and database formats for "
 "analysis. A total of 23,215 readings were collected during the measurement period (sensor worn).")

set_text(P[37],
 "The measured metrics are: (1) delivery ratio, i.e., the number of records matched at the smartphone "
 "divided by the number of records recorded by the smartwatch, with matching based on the time attribute so "
 "it does not depend on clock synchronization between devices; (2) value fidelity, i.e., the proportion of "
 "received records whose values (bpm and accuracy) are identical to the smartwatch records; (3) the number "
 "of duplicates at the receiver; (4) the sensor-accuracy distribution following the accuracy-status "
 "constants of the Android SensorManager (range -1 to 3) [2]; and (5) per-batch transfer latency and "
 "throughput recorded automatically by the application instrumentation. Throughput is computed as the "
 "payload size divided by the transfer duration.")

set_text(P[42],
 "The evaluation was carried out over four test sessions, with per-session results in Table 2 and their "
 "temporal profiles in Figure 6 and Figure 7. In aggregate, the smartwatch recorded 23,215 readings during "
 "the measurement period and 20,455 of them were recorded on the smartphone (88.1% delivery ratio) with no "
 "duplicates. However, the loss was not evenly distributed: while the BLE link was active, loss was near "
 "zero (e.g., after the smartphone connected in the main session, delivery was practically 100%), whereas "
 "all losses were concentrated in periods when the smartphone was disconnected—including one short session "
 "whose smartphone never connected, so none of its readings were stored at the receiver. Importantly, every "
 "successfully received reading was identical to the smartwatch record based on time matching, both bpm and "
 "accuracy values (100% value fidelity), so no data corruption occurred on the BLE channel. The discrepancy "
 "between the sent marker on the smartwatch and the presence of records on the smartphone is discussed in "
 "Section 3.5.")

set_text(P[43], "Table 2. Delivery results per test session.")

set_text(P[46],
 "The sensor-accuracy distribution during the measurement period is shown in Table 3 and Figure 9. 90.8% of "
 "readings were at the highest accuracy (value 3), 9.1% (2,113 readings) at the no-contact condition "
 "(value -1)—which commonly occurs when the sensor momentarily loses skin contact [2]—and 0.1% (15 "
 "readings) at value 0. The heart-rate statistics for the highest-accuracy readings (n=21,087, Figure 8) "
 "were: minimum 60 bpm, maximum 123 bpm, mean 83.4 bpm, and standard deviation 9.9 bpm. Note that 22,231 "
 "samples in the tail of the recording—with frozen (constant) values and accuracy ≤0 because the smartwatch "
 "was not worn—were excluded from the statistics above to avoid biasing the results.")

set_text(P[47], "Table 3. Distribution of sensor-accuracy values (measurement period).")

set_text(P[51], "3.6 Discussion and Limitations")

set_text(P[52],
 "The results show that the combination of store-and-forward, ACK, and an idempotent receiver can preserve "
 "data completeness and integrity while the BLE link is active—evidenced by 100% value fidelity and "
 "near-zero loss when connected—while opcode-based framing with flow control overcomes the BLE notification "
 "size limit. Time-based matching makes the delivery computation independent of clock synchronization. "
 "Conversely, as discussed in Section 3.5, reliability degrades in disconnect scenarios, so backlog backfill "
 "needs to be strengthened.")

set_text(P[53],
 "This study has several limitations. First, the evaluation spans four sessions on a single device pair and "
 "a single environment, so replication across devices, distances, and interference conditions is needed to "
 "obtain representative means and standard deviations. Second, it was found that the backlog during "
 "disconnection is not always backfilled (Section 3.5); strengthening the ACK and backfill trigger is a "
 "priority improvement. Third, background execution covers the application being in the background and the "
 "screen off, but does not yet guarantee operation when the application is force-closed or after a device "
 "reboot, which is also constrained by OS power-saving policies. Fourth, the BLE communication does not yet "
 "apply encryption/authentication, so the security of health data is a development agenda. Fifth, the sensor "
 "on a consumer smartwatch is not clinically validated, so the contribution focuses on communication "
 "reliability rather than the medical accuracy of the heart-rate values.")

set_text(P[55],
 "This study designed and implemented a BLE-based heart-rate data delivery system from a smartwatch to a "
 "smartphone using batch framing, store-and-forward, delivery acknowledgement (ACK), an idempotent receiver, "
 "and background execution. Testing on physical devices across four sessions (23,215 readings in total) "
 "showed high data integrity—100% value fidelity and near-zero loss while the BLE link was active—with "
 "90.8% of data at the highest accuracy and fast batch transfer (~0.32 s for 228 readings at MTU 512). The "
 "evaluation also revealed a limitation: during disconnection periods the backlog was not always "
 "backfilled, so the aggregate delivery ratio was 88.1%. Future work includes strengthening the backfill "
 "and ACK mechanisms for disconnect-reconnect resilience, formal measurement across varied conditions with "
 "replication, strengthening background execution (including auto-start after reboot), and adding encryption "
 "for data security.")

# ---------- Table 1: add dataset row ----------
r = t_dev.add_row()
r.cells[0].text = "Test dataset"
r.cells[1].text = "23–28 Jun 2026; 4 sessions; 23,215 measurement readings"

# ---------- Table 3: accuracy (measurement period) ----------
t_acc.add_row()
acc_rows = [("Sensor accuracy", "Count", "Percentage"),
            ("3 (high)", "21,087", "90.8%"),
            ("0 (medium)", "15", "0.1%"),
            ("-1 (no contact)", "2,113", "9.1%"),
            ("Total", "23,215", "100%")]
for i, row in enumerate(acc_rows):
    for j, v in enumerate(row):
        t_acc.rows[i].cells[j].text = v
for c in t_acc.rows[0].cells:
    for rn in c.paragraphs[0].runs: rn.bold = True

# ---------- Table 2: rebuild per session ----------
deliv_rows = [
    ("Session", "Start (WIB)", "Duration", "Recorded", "Received", "Lost", "Delivery"),
    ("Session 1", "23/06 06:38", "36 min", "2,177", "2,164", "13", "99.40%"),
    ("Session 2", "23/06 08:41", "1 min", "53", "0", "53", "0.00%"),
    ("Session 3", "23/06 14:15", "345 min", "20,702", "18,013", "2,689", "87.01%"),
    ("Session 4", "28/06 21:29", "5 min", "283", "278", "5", "98.23%"),
    ("Total", "—", "387 min", "23,215", "20,455", "2,760", "88.11%"),
]
old_tbl = t_deliv._tbl
new = d.add_table(rows=len(deliv_rows), cols=7)
new.style = "Table Grid"
for i, row in enumerate(deliv_rows):
    for j, v in enumerate(row):
        cell = new.rows[i].cells[j]
        cell.text = v
        if i == 0 or row[0] == "Total":
            for rn in cell.paragraphs[0].runs: rn.bold = True
old_tbl.addnext(new._tbl)
old_tbl.getparent().remove(old_tbl)

# ---------- figures ----------
def insert_fig(before, img, caption, width_in):
    pim = before.insert_paragraph_before()
    pim.alignment = ALIGN.CENTER
    pim.add_run().add_picture(str(ROOT / "figures" / img), width=Inches(width_in))
    pc = before.insert_paragraph_before(caption)
    pc.alignment = ALIGN.CENTER
    pc.runs[0].font.size = Pt(10)

insert_fig(p_3_3, "fig_hr_completeness.png", "Figure 6. Data completeness per test session (received vs lost).", 5.6)
insert_fig(p_3_3, "fig_hr_timeline.png", "Figure 7. Heart-rate signal and BLE packet loss in the main session.", 6.3)
insert_fig(p_3_4, "fig_hr_bpm_dist.png", "Figure 8. Distribution of heart-rate values for the highest-accuracy readings.", 4.7)
insert_fig(p_3_4, "fig_hr_contact.png", "Figure 9. Sensor contact-quality distribution during the measurement period.", 5.0)

# ---------- new subsection 3.5 ----------
ph = p_3_5.insert_paragraph_before("3.5 Finding: Store-and-Forward Backfill Limitation")
ph.runs[0].bold = True
ph.runs[0].font.size = Pt(11)
p_3_5.insert_paragraph_before(
 "A cross-analysis between the sent-status marker (synced) on the smartwatch and the presence of records on "
 "the smartphone revealed an important finding. A total of 2,879 readings were marked as sent (synced = 1) "
 "by the smartwatch but were not found in the smartphone database, whereas only the last 5 readings were "
 "genuinely still pending. All of this 2,879-record difference occurred during periods when the smartphone "
 "was disconnected. This indicates that, in the current implementation, the sent marking is applied at the "
 "notification-send level without being fully confirmed by an application-level ACK when the link is "
 "inactive, so the backlog accumulated during disconnection is not always retransmitted (backfilled) after "
 "the link recovers. Consequently, the store-and-forward mechanism preserves completeness well while the "
 "link is active but is not yet fully resilient to disconnect-reconnect scenarios. The recommended "
 "improvements are to defer the synced marking until an ACK is received and to add a backfill trigger that "
 "scans unconfirmed records whenever a connection is re-established. This finding reinforces the importance "
 "of application-level acknowledgement over relying on BLE notifications alone.")

d.save(str(DOC))
print("OK — English manuscript updated:", DOC.name)
print("tables:", len(d.tables), "| paragraphs:", len(d.paragraphs))
