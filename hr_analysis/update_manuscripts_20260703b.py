#!/usr/bin/env python3
"""Perbarui §3.5 kedua manuskrip: analisis silang backfill dihitung ulang pada
CAKUPAN YANG SAMA dengan Tabel 2 (periode pengukuran valid, n=23.215).

Hasil hitung (hr_analysis, dari watch/phone-2026-06-28.csv, logika trim sama
dengan make_figs.py): dari 2.760 pembacaan hilang, 2.755 (99,8%) ber-synced=1
(ditandai terkirim tapi tak ada di ponsel) dan 5 memang pending wajar —
terurai persis 2.755 + 5 = 2.760. Per sesi: 13/13, 53/53, 2.689/2.689, 0/5.
Angka lama 2.879 (cakupan seluruh rekaman) dipertahankan sebagai catatan
sekunder satu paragraf."""
from docx import Document
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]

def set_text(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        p.add_run(text)

def find(doc, prefix):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    raise SystemExit(f"PARAGRAF TIDAK DITEMUKAN: {prefix!r}")

# ============================ NASKAH ID ============================
doc = Document(str(ROOT / "Draft_Naskah_HR_BLE_SINTA2.docx"))

p1 = find(doc, "Analisis silang antara penanda status terkirim")
set_text(p1,
    "Analisis silang antara penanda status terkirim (synced) pada smartwatch dan keberadaan record di "
    "smartphone — dihitung pada periode pengukuran yang sama dengan Tabel 2 — mengungkap satu temuan "
    "penting. Dari 2.760 pembacaan yang hilang, 2.755 di antaranya (99,8%) ternyata telah ditandai "
    "terkirim (synced = 1) oleh smartwatch tetapi tidak ditemukan di basis data smartphone; hanya 5 "
    "pembacaan terakhir yang memang masih berstatus belum terkirim (menunggu interval berikutnya). "
    "Dengan demikian, kehilangan pada Tabel 2 terurai persis menjadi 2.755 record bertanda terkirim "
    "secara keliru ditambah 5 record tertunda yang wajar. Seluruh 2.755 record tersebut hilang pada "
    "periode ketika smartphone sedang tidak terhubung — termasuk 53 record pada Sesi 2 yang "
    "smartphone-nya tidak pernah terhubung sama sekali, bukti paling jelas bahwa penandaan dilakukan "
    "prematur. Hal ini menunjukkan bahwa, pada implementasi saat ini, penandaan terkirim dilakukan pada "
    "tingkat pengiriman notifikasi tanpa terkonfirmasi penuh oleh ACK tingkat aplikasi ketika tautan "
    "tidak aktif, sehingga backlog yang terbentuk selama putus koneksi tidak selalu dikirim ulang "
    "(di-backfill) setelah tautan pulih. Akibatnya, mekanisme store-and-forward menjaga kelengkapan "
    "dengan baik selama tautan aktif, tetapi belum sepenuhnya tahan terhadap skenario putus-sambung. "
    "Perbaikan yang disarankan adalah menunda penandaan synced hingga ACK diterima dan menambahkan "
    "pemicu backfill yang memindai record berstatus belum-terkonfirmasi setiap kali koneksi terbentuk "
    "kembali. Temuan ini memperkuat pentingnya konfirmasi tingkat aplikasi dibanding hanya mengandalkan "
    "notifikasi BLE.")

p2 = find(doc, "Sebagai catatan cakupan: angka 2.879")
set_text(p2,
    "Bila dihitung atas seluruh rekaman smartwatch (45.446 baris, termasuk periode sensor tidak "
    "terpasang), jumlah record bertanda terkirim yang tidak ditemukan di smartphone adalah 2.879 "
    "(45.446 − 42.562 = 2.884, dikurangi 5 record tertunda); angka 2.755 di atas adalah bagian dari "
    "jumlah tersebut yang berada pada periode pengukuran valid, sehingga konsisten dengan Tabel 2.")

doc.save(str(ROOT / "Draft_Naskah_HR_BLE_SINTA2.docx"))
print("OK -> Draft_Naskah_HR_BLE_SINTA2.docx")

# ============================ MANUSCRIPT EN ============================
doc = Document(str(ROOT / "Draft_Manuscript_HR_BLE_EN.docx"))

p1 = find(doc, "A cross-analysis between the sent-status marker")
set_text(p1,
    "A cross-analysis between the sent-status marker (synced) on the smartwatch and the presence of "
    "records on the smartphone — computed over the same measurement period as Table 2 — revealed an "
    "important finding. Of the 2,760 lost readings, 2,755 (99.8%) had been marked as sent (synced = 1) "
    "by the smartwatch yet were not found in the smartphone database; only the last 5 readings were "
    "genuinely still pending (awaiting the next interval). The loss in Table 2 therefore decomposes "
    "exactly into 2,755 falsely-marked-as-sent records plus 5 legitimately pending ones. All 2,755 "
    "records were lost during periods when the smartphone was disconnected — including the 53 records "
    "of Session 2, whose smartphone never connected at all, the clearest evidence that the marking is "
    "applied prematurely. This indicates that, in the current implementation, the sent marking is "
    "applied at the notification-send level without being fully confirmed by an application-level ACK "
    "when the link is inactive, so the backlog accumulated during disconnection is not always "
    "retransmitted (backfilled) after the link recovers. Consequently, the store-and-forward mechanism "
    "preserves completeness well while the link is active but is not yet fully resilient to "
    "disconnect-reconnect scenarios. The recommended improvements are to defer the synced marking until "
    "an ACK is received and to add a backfill trigger that scans unconfirmed records whenever a "
    "connection is re-established. This finding reinforces the importance of application-level "
    "acknowledgement over relying on BLE notifications alone.")

p2 = find(doc, "As a note on scope: the 2,879 figure")
set_text(p2,
    "When computed over the entire smartwatch recording (45,446 rows, including the not-worn period), "
    "the number of marked-as-sent records missing from the smartphone is 2,879 (45,446 − 42,562 = "
    "2,884, minus 5 pending records); the 2,755 figure above is the portion of these that falls within "
    "the valid measurement period, and is therefore consistent with Table 2.")

doc.save(str(ROOT / "Draft_Manuscript_HR_BLE_EN.docx"))
print("OK -> Draft_Manuscript_HR_BLE_EN.docx")
