#!/usr/bin/env python3
"""Perbarui kedua manuskrip (ID & EN) per 3 Juli 2026, menyelaraskan dengan
catatan 3 Juli:
1. §2.6  — definisi atribut waktu (time): timestamp pencatatan aplikasi ≈1 Hz,
   kunci unik record & dasar pencocokan (bukan waktu fisiologis detak).
2. §3.4  — ukuran data: ±47 byte/pembacaan, ukuran batch per interval, hasil
   negosiasi MTU (512 penuh pada pasangan uji), dan validasi hitungan frame
   (22 DATA + START + END = 24, konsisten log instrumentasi).
3. §3.5  — klarifikasi cakupan angka 2.879 (seluruh rekaman) vs 2.760
   (periode pengukuran valid) agar tidak tampak kontradiktif bagi reviewer.
Semua penanda [[...]] dipertahankan."""
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]

def set_text(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        p.add_run(text)

def insert_after(p, text):
    new_p = OxmlElement("w:p")
    p._p.addnext(new_p)
    np = Paragraph(new_p, p._parent)
    np.style = p.style
    np.add_run(text)
    return np

def find(doc, prefix):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    raise SystemExit(f"PARAGRAF TIDAK DITEMUKAN: {prefix!r}")

# ============================ NASKAH ID ============================
doc = Document(str(ROOT / "Draft_Naskah_HR_BLE_SINTA2.docx"))

# 1. §2.6 — definisi atribut waktu (time)
p26 = find(doc, "Metrik yang diukur meliputi")
insert_after(p26,
    "Atribut waktu (time) pada metrik di atas adalah timestamp pencatatan aplikasi smartwatch, bukan "
    "waktu fisiologis terjadinya detak: setiap detik (≈1 Hz) aplikasi mencatat nilai bpm terakhir dari "
    "sensor beserta waktu perangkat saat pencatatan, disimpan sebagai epoch milidetik. Timestamp ini "
    "unik untuk tiap record sehingga berfungsi sebagai kunci record sekaligus dasar pencocokan "
    "watch–smartphone dan deteksi duplikat.")

# 2. §3.4 — ukuran data, MTU ternegosiasi, validasi hitungan frame
p34 = find(doc, "Sebagai ilustrasi kinerja per batch")
set_text(p34,
    "Sebagai ilustrasi kinerja per batch, satu batch berisi 228 pembacaan (payload 10.717 byte) dipecah "
    "menjadi 24 frame dan terkirim dalam ±0,32 detik (323,1 ms) pada MTU 512, setara throughput ±33 "
    "KB/detik; di sisi penerima, perakitan ulang frame memerlukan ±250 ms dan penyimpanan batch ke "
    "basis data ±56 ms. Dari data yang sama, ukuran satu pembacaan terserialisasi adalah ±47 byte JSON "
    "(10.717 ÷ 228), sehingga batch tipikal berisi ±180 pembacaan (±8,3 KB) untuk interval 3 menit atau "
    "±300 pembacaan (±13,8 KB) untuk interval 5 menit. Negosiasi MTU pada pasangan perangkat uji "
    "menghasilkan nilai penuh 512 byte (chunk 508 byte), dan jumlah frame terukur sesuai perhitungan: "
    "10.717 ÷ 508 = 21,1 dibulatkan ke atas menjadi 22 frame DATA, ditambah START dan END menjadi 24 "
    "frame — konsisten dengan log instrumentasi. Nilai-nilai ini menunjukkan transfer batch berlangsung "
    "cepat relatif terhadap interval pengiriman (menit). Pengukuran agregat lintas variasi kondisi "
    "(ukuran batch berbeda akibat interval 3 vs 5 menit, jarak antar-perangkat, dan skenario gangguan) "
    "merupakan bagian dari pekerjaan lanjutan. [[Lengkapi dengan tabel/grafik hasil replikasi.]]")

# 3. §3.5 — klarifikasi cakupan 2.879 vs 2.760
p35 = find(doc, "Analisis silang antara penanda status terkirim")
insert_after(p35,
    "Sebagai catatan cakupan: angka 2.879 dihitung atas seluruh rekaman smartwatch (45.446 baris, "
    "termasuk periode sensor tidak terpasang; 45.446 − 42.562 = 2.884, dikurangi 5 record yang masih "
    "menunggu pengiriman = 2.879), sedangkan angka kehilangan 2.760 pada Tabel 2 hanya mencakup periode "
    "pengukuran valid (23.215 pembacaan). Keduanya konsisten dan tidak saling bertentangan — "
    "perbedaannya semata pada cakupan data yang dihitung.")

doc.save(str(ROOT / "Draft_Naskah_HR_BLE_SINTA2.docx"))
print("OK -> Draft_Naskah_HR_BLE_SINTA2.docx")

# ============================ MANUSCRIPT EN ============================
doc = Document(str(ROOT / "Draft_Manuscript_HR_BLE_EN.docx"))

p26 = find(doc, "The measured metrics are")
insert_after(p26,
    "The time attribute above is the smartwatch application's recording timestamp, not the "
    "physiological time of a heartbeat: every second (≈1 Hz) the application records the latest sensor "
    "bpm value together with the device time at the moment of recording, stored as epoch milliseconds. "
    "This timestamp is unique per record and therefore serves as the record key as well as the basis "
    "for watch–smartphone matching and duplicate detection.")

p34 = find(doc, "As an illustration of per-batch performance")
set_text(p34,
    "As an illustration of per-batch performance, one batch of 228 readings (10,717-byte payload) was "
    "fragmented into 24 frames and delivered in ~0.32 s (323.1 ms) at MTU 512, equivalent to a "
    "throughput of ~33 KB/s; on the receiver side, frame reassembly took ~250 ms and storing the batch "
    "to the database took ~56 ms. From the same data, one serialized reading is ~47 bytes of JSON "
    "(10,717 ÷ 228), so a typical batch contains ~180 readings (~8.3 KB) for the 3-minute interval or "
    "~300 readings (~13.8 KB) for the 5-minute interval. MTU negotiation on the tested device pair "
    "yielded the full 512 bytes (508-byte chunks), and the measured frame count matches the "
    "calculation: 10,717 ÷ 508 = 21.1, rounded up to 22 DATA frames, plus START and END = 24 frames — "
    "consistent with the instrumentation log. These values show that batch transfer is fast relative "
    "to the sending interval (minutes). Aggregate measurements across varied conditions (different "
    "batch sizes due to 3- vs 5-minute intervals, device distance, and disruption scenarios) are part "
    "of future work. [[Complete with replication results table/charts.]]")

p35 = find(doc, "A cross-analysis between the sent-status marker")
insert_after(p35,
    "As a note on scope: the 2,879 figure is computed over the entire smartwatch recording (45,446 "
    "rows, including the not-worn period; 45,446 − 42,562 = 2,884, minus 5 records still pending = "
    "2,879), whereas the 2,760 losses in Table 2 cover only the valid measurement period (23,215 "
    "readings). The two figures are consistent rather than contradictory — they simply differ in the "
    "scope of data counted.")

doc.save(str(ROOT / "Draft_Manuscript_HR_BLE_EN.docx"))
print("OK -> Draft_Manuscript_HR_BLE_EN.docx")
